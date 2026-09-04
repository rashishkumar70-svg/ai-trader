# ============================================================
# AI TRADER PRO v13.6 — "PRO TERMINAL · 500 LIVE" (single file)
# YOUR v13 app EXACTLY as it was + ONE new separate tab:
#   🔴 LIVE 500-STOCK DASHBOARD — one common board that shows which
#   stocks are in an UPTREND right now, auto-analysing up to 500
#   stocks in fast batches, auto-refreshing while you watch.
# Everything else (Analyze · Scanner · Search · Below-100 · Journal
# · Guide · 4 pivot systems · ML · reports) is unchanged from v13.
# ------------------------------------------------------------
# pip install streamlit yfinance pandas numpy ta plotly scikit-learn scipy python-docx streamlit-autorefresh
# streamlit run ai_trader_v13_500.py
# (streamlit-autorefresh = smooth LIVE auto-refresh; without it you
#  get a fallback auto-refresh + a manual 🔄 button. All optional.)
# ============================================================

import streamlit as st

# ── cross-version layout kwarg ──────────────────────────────────────────────
# New Streamlit (2026) REMOVED `use_container_width`; older versions don't know
# `width="stretch"`. This picks the right one at startup, so the app runs on
# BOTH old and new Streamlit (PC + Streamlit Cloud) without layout errors.
try:
    _sv = tuple(int(x) for x in st.__version__.split(".")[:2] if x.isdigit())
    STRETCH = {"width": "stretch"} if _sv >= (1, 46) else {"use_container_width": True}
except Exception:
    STRETCH = {"use_container_width": True}
import yfinance as yf
import pandas as pd
import numpy as np
import ta
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import warnings
import time
import io
import re
import urllib.request
from datetime import datetime
from datetime import time as dtime
from datetime import timezone as _dts, timedelta as _dtd

# ── IST clock: cloud servers run on UTC — every time in this app is IST ──
TZ_IST = _dts(_dtd(hours=5, minutes=30))


def now_ist():
    try:
        return datetime.now(TZ_IST)
    except Exception:
        return datetime.now()


def _dist(ix):
    """IST calendar-date of a candle timestamp (yfinance sends IST-aware stamps)."""
    try:
        return ix.astimezone(TZ_IST).date() if getattr(ix, "tzinfo", None) else ix.date()
    except Exception:
        return ix.date()
import concurrent.futures

warnings.filterwarnings('ignore')

st.set_page_config(page_title="AI Trader Pro v13.6 Pro Terminal", page_icon="💹",
                   layout="wide", initial_sidebar_state="collapsed")

# ============================================================
# CSS — PREMIUM LIGHT THEME (+ dashboard components)
# ============================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
*{font-family:'Inter',sans-serif;box-sizing:border-box;}
.stApp{background:#f5f7ff;color:#1a1f36;}
.navbar{background:linear-gradient(135deg,#0d1b6e 0%,#1565c0 60%,#0288d1 100%);
  border-radius:20px;padding:18px 28px;margin-bottom:22px;box-shadow:0 8px 32px rgba(13,27,110,0.25);}
.input-row{background:white;border-radius:16px;padding:20px 24px;margin-bottom:18px;
  border:1px solid #e0e7ff;box-shadow:0 2px 12px rgba(0,0,0,0.06);}
.sig-buy{background:linear-gradient(135deg,#f0fdf4,#dcfce7);border:3px solid #16a34a;
  border-radius:20px;padding:28px;box-shadow:0 8px 32px rgba(22,163,74,0.15);}
.sig-sell{background:linear-gradient(135deg,#fff1f2,#ffe4e6);border:3px solid #dc2626;
  border-radius:20px;padding:28px;box-shadow:0 8px 32px rgba(220,38,38,0.12);}
.sig-wait{background:linear-gradient(135deg,#fffbeb,#fef3c7);border:3px solid #d97706;
  border-radius:20px;padding:28px;box-shadow:0 8px 32px rgba(217,119,6,0.12);}
.mc-blue{background:linear-gradient(135deg,#eff6ff,#dbeafe);border:2px solid #3b82f6;
  border-radius:16px;padding:18px;text-align:center;box-shadow:0 4px 16px rgba(59,130,246,0.12);}
.mc-green{background:linear-gradient(135deg,#f0fdf4,#dcfce7);border:2px solid #16a34a;
  border-radius:16px;padding:18px;text-align:center;box-shadow:0 4px 16px rgba(22,163,74,0.12);}
.mc-red{background:linear-gradient(135deg,#fff1f2,#ffe4e6);border:2px solid #dc2626;
  border-radius:16px;padding:18px;text-align:center;box-shadow:0 4px 16px rgba(220,38,38,0.1);}
.mc-purple{background:linear-gradient(135deg,#faf5ff,#f3e8ff);border:2px solid #9333ea;
  border-radius:16px;padding:18px;text-align:center;box-shadow:0 4px 16px rgba(147,51,234,0.1);}
.mc-orange{background:linear-gradient(135deg,#fff7ed,#ffedd5);border:2px solid #ea580c;
  border-radius:16px;padding:18px;text-align:center;box-shadow:0 4px 16px rgba(234,88,12,0.1);}
.mc-white{background:white;border:1px solid #e0e7ff;border-radius:16px;padding:18px;
  text-align:center;box-shadow:0 2px 10px rgba(0,0,0,0.05);}
.lv-r{background:linear-gradient(90deg,#fff1f2,white);border-left:5px solid #ef4444;
  border-radius:0 12px 12px 0;padding:11px 18px;margin:3px 0;box-shadow:0 1px 6px rgba(239,68,68,0.08);}
.lv-s{background:linear-gradient(90deg,#f0fdf4,white);border-left:5px solid #22c55e;
  border-radius:0 12px 12px 0;padding:11px 18px;margin:3px 0;box-shadow:0 1px 6px rgba(34,197,94,0.08);}
.lv-p{background:linear-gradient(90deg,#fffbeb,white);border-left:5px solid #f59e0b;
  border-radius:0 12px 12px 0;padding:11px 18px;margin:3px 0;}
.lv-fib{background:linear-gradient(90deg,#faf5ff,white);border-left:5px solid #a855f7;
  border-radius:0 12px 12px 0;padding:11px 18px;margin:3px 0;box-shadow:0 1px 6px rgba(168,85,247,0.08);}
.lv-cur{background:linear-gradient(90deg,#eff6ff,white);border-left:5px solid #3b82f6;
  border-radius:0 12px 12px 0;padding:11px 18px;margin:3px 0;box-shadow:0 2px 8px rgba(59,130,246,0.12);}
.sh{display:flex;align-items:center;gap:10px;background:linear-gradient(90deg,#eff6ff,transparent);
  border-left:5px solid #1d4ed8;padding:10px 18px;border-radius:0 12px 12px 0;margin:26px 0 14px;
  font-size:15px;font-weight:800;color:#1d4ed8;letter-spacing:0.3px;}
.sc-r{background:white;border:1px solid #e0e7ff;border-radius:14px;padding:16px;margin:6px 0;box-shadow:0 2px 8px rgba(0,0,0,0.04);}
.pbar-container{background:#e0e7ff;border-radius:20px;height:10px;overflow:hidden;}
.pbar-buy{background:linear-gradient(90deg,#16a34a,#22c55e);height:10px;border-radius:20px;}
.pbar-sell{background:linear-gradient(90deg,#dc2626,#ef4444);height:10px;border-radius:20px;}
[data-testid="stMetric"]{background:white !important;border:1px solid #e0e7ff !important;border-radius:14px !important;
  padding:14px !important;box-shadow:0 2px 8px rgba(0,0,0,0.04) !important;}
[data-testid="stMetricValue"]{color:#1a1f36 !important;font-weight:800 !important;font-size:22px !important;}
[data-testid="stMetricLabel"]{color:#6b7280 !important;font-size:11px !important;font-weight:600 !important;}
.stButton>button{background:linear-gradient(135deg,#1d4ed8,#3b82f6) !important;color:white !important;
  border:none !important;border-radius:12px !important;font-weight:700 !important;padding:10px 24px !important;
  font-size:14px !important;box-shadow:0 4px 12px rgba(29,78,216,0.3) !important;}
.stButton>button:hover{background:linear-gradient(135deg,#1e40af,#2563eb) !important;transform:translateY(-1px) !important;}
div[data-testid="stTextInput"]>div>div>input{background:white !important;border:2px solid #e0e7ff !important;
  border-radius:12px !important;color:#1a1f36 !important;font-size:14px !important;font-weight:500 !important;}
[data-testid="stSidebar"]{display:none !important;}
#MainMenu{visibility:hidden;}footer{visibility:hidden;}header{visibility:hidden;}
::-webkit-scrollbar{width:6px;}::-webkit-scrollbar-thumb{background:#3b82f6;border-radius:3px;}
/* ---- dashboard components ---- */
.breadth{display:flex;height:18px;border-radius:10px;overflow:hidden;border:1px solid #e0e7ff;background:white;}
.b-up{background:linear-gradient(90deg,#16a34a,#22c55e);height:100%;}
.b-dn{background:linear-gradient(90deg,#dc2626,#ef4444);height:100%;}
.b-fl{background:#e5e7eb;height:100%;}
.uchip{display:inline-block;background:#f0fdf4;border:1.5px solid #16a34a;color:#15803d;font-weight:800;
  font-size:12px;padding:5px 12px;border-radius:20px;margin:3px 3px 3px 0;}
.dchip{display:inline-block;background:#fff1f2;border:1.5px solid #dc2626;color:#b91c1c;font-weight:800;
  font-size:12px;padding:5px 12px;border-radius:20px;margin:3px 3px 3px 0;}
.nchip{display:inline-block;background:#f9fafb;border:1.5px solid #d1d5db;color:#6b7280;font-weight:700;
  font-size:12px;padding:5px 12px;border-radius:20px;margin:3px 3px 3px 0;}
.achip{display:inline-block;background:#fffbeb;border:1.5px solid #f59e0b;color:#92400e;font-size:11.5px;
  padding:4px 10px;border-radius:9px;margin:3px 3px 3px 0;font-weight:700;}
.ucard{background:white;border:1px solid #bbf7d0;border-left:5px solid #16a34a;border-radius:12px;
  padding:12px 14px;margin:4px 0;box-shadow:0 2px 8px rgba(0,0,0,0.05);}
.ucard .nm{font-weight:900;font-size:14px;color:#1a1f36;}
.ucard .pr{font-weight:900;font-size:18px;color:#1d4ed8;}
.badge{display:inline-block;border-radius:20px;padding:2px 10px;font-size:10.5px;font-weight:800;color:white;}
.dashhead{background:linear-gradient(135deg,#7c2d12,#b45309,#d97706);border-radius:18px;padding:18px 24px;
  margin-bottom:16px;box-shadow:0 8px 28px rgba(180,83,9,0.25);}
</style>
""", unsafe_allow_html=True)

# ============================================================
# CURATED STOCK DATABASE + LIVE FULL NSE UNIVERSE
# ============================================================
NIFTY50 = {
    "RELIANCE":"RELIANCE.NS","TCS":"TCS.NS","HDFCBANK":"HDFCBANK.NS","INFY":"INFY.NS",
    "ICICIBANK":"ICICIBANK.NS","HINDUNILVR":"HINDUNILVR.NS","ITC":"ITC.NS","SBIN":"SBIN.NS",
    "BHARTIARTL":"BHARTIARTL.NS","KOTAKBANK":"KOTAKBANK.NS","LT":"LT.NS","AXISBANK":"AXISBANK.NS",
    "ASIANPAINT":"ASIANPAINT.NS","MARUTI":"MARUTI.NS","SUNPHARMA":"SUNPHARMA.NS","TITAN":"TITAN.NS",
    "BAJFINANCE":"BAJFINANCE.NS","HCLTECH":"HCLTECH.NS","TMPV":"TMPV.NS","TATASTEEL":"TATASTEEL.NS",
    "NTPC":"NTPC.NS","ONGC":"ONGC.NS","POWERGRID":"POWERGRID.NS","COALINDIA":"COALINDIA.NS",
    "ADANIENT":"ADANIENT.NS","ADANIPORTS":"ADANIPORTS.NS","JSWSTEEL":"JSWSTEEL.NS","TECHM":"TECHM.NS",
    "INDUSINDBK":"INDUSINDBK.NS","DRREDDY":"DRREDDY.NS","BAJAJ-AUTO":"BAJAJ-AUTO.NS","NESTLEIND":"NESTLEIND.NS",
    "CIPLA":"CIPLA.NS","BRITANNIA":"BRITANNIA.NS","TATAPOWER":"TATAPOWER.NS","APOLLOHOSP":"APOLLOHOSP.NS",
    "BAJAJFINSV":"BAJAJFINSV.NS","ULTRACEMCO":"ULTRACEMCO.NS","HINDALCO":"HINDALCO.NS","M&M":"M&M.NS",
    "WIPRO":"WIPRO.NS","EICHERMOT":"EICHERMOT.NS","HEROMOTOCO":"HEROMOTOCO.NS","BPCL":"BPCL.NS",
    "TATACONSUM":"TATACONSUM.NS","DIVISLAB":"DIVISLAB.NS","SBILIFE":"SBILIFE.NS","HDFCLIFE":"HDFCLIFE.NS",
    "SHRIRAMFIN":"SHRIRAMFIN.NS","GRASIM":"GRASIM.NS",
}
NEXT50 = {
    "ETERNAL (ZOMATO)":"ETERNAL.NS","JIOFIN":"JIOFIN.NS","ADANIGREEN":"ADANIGREEN.NS","ADANIPOWER":"ADANIPOWER.NS","IRFC":"IRFC.NS",
    "RVNL":"RVNL.NS","HAL":"HAL.NS","BEL":"BEL.NS","DLF":"DLF.NS","VEDL":"VEDL.NS","NHPC":"NHPC.NS",
    "SJVN":"SJVN.NS","IRCTC":"IRCTC.NS","BANKBARODA":"BANKBARODA.NS","PNB":"PNB.NS","CANBK":"CANBK.NS",
    "BHEL":"BHEL.NS","SAIL":"SAIL.NS","GAIL":"GAIL.NS","IOC":"IOC.NS","HINDPETRO":"HINDPETRO.NS",
    "LICI":"LICI.NS","AMBUJACEM":"AMBUJACEM.NS","HDFCAMC":"HDFCAMC.NS","PIDILITIND":"PIDILITIND.NS",
    "SIEMENS":"SIEMENS.NS","ABB":"ABB.NS","POLYCAB":"POLYCAB.NS","DIXON":"DIXON.NS",
    "GODREJPROP":"GODREJPROP.NS","TRENT":"TRENT.NS","VBL":"VBL.NS","MARICO":"MARICO.NS",
    "DABUR":"DABUR.NS","COLPAL":"COLPAL.NS","DMART":"DMART.NS","NAUKRI":"NAUKRI.NS",
    "PAGEIND":"PAGEIND.NS","MRF":"MRF.NS","BOSCHLTD":"BOSCHLTD.NS","HAVELLS":"HAVELLS.NS",
    "MUTHOOTFIN":"MUTHOOTFIN.NS","CHOLAFIN":"CHOLAFIN.NS","TATAELXSI":"TATAELXSI.NS",
    "RECLTD":"RECLTD.NS","PFC":"PFC.NS","HUDCO":"HUDCO.NS","NMDC":"NMDC.NS",
}
MIDCAP = {
    "PERSISTENT":"PERSISTENT.NS","COFORGE":"COFORGE.NS","MPHASIS":"MPHASIS.NS","KPITTECH":"KPITTECH.NS",
    "LTTS":"LTTS.NS","ANGELONE":"ANGELONE.NS","MCX":"MCX.NS","CDSL":"CDSL.NS","COCHINSHIP":"COCHINSHIP.NS",
    "MAZDOCK":"MAZDOCK.NS","GRSE":"GRSE.NS","BDL":"BDL.NS","BEML":"BEML.NS","IRCON":"IRCON.NS",
    "KEI":"KEI.NS","APLAPOLLO":"APLAPOLLO.NS","INOXWIND":"INOXWIND.NS","JSWENERGY":"JSWENERGY.NS",
    "WAAREEENER":"WAAREEENER.NS","DEEPAKNTR":"DEEPAKNTR.NS","SRF":"SRF.NS","KALYANKJIL":"KALYANKJIL.NS",
    "SENCO":"SENCO.NS","INDHOTEL":"INDHOTEL.NS","APOLLOTYRE":"APOLLOTYRE.NS","CEATLTD":"CEATLTD.NS",
    "TVSMOTOR":"TVSMOTOR.NS","BHARATFORG":"BHARATFORG.NS","MOTHERSON":"MOTHERSON.NS","MANAPPURAM":"MANAPPURAM.NS",
    "NBCC":"NBCC.NS","NCC":"NCC.NS","IRB":"IRB.NS","ATGL":"ATGL.NS","PIIND":"PIIND.NS",
    "ZYDUSLIFE":"ZYDUSLIFE.NS","LAURUSLABS":"LAURUSLABS.NS","ZENSARTECH":"ZENSARTECH.NS",
    "FEDERALBNK":"FEDERALBNK.NS","AUBANK":"AUBANK.NS","MAXHEALTH":"MAXHEALTH.NS",
    "METROPOLIS":"METROPOLIS.NS","LALPATHLAB":"LALPATHLAB.NS",
}
BELOW100 = {
    "YESBANK":"YESBANK.NS","IDEA":"IDEA.NS","SUZLON":"SUZLON.NS","NHPC":"NHPC.NS","SJVN":"SJVN.NS",
    "RPOWER":"RPOWER.NS","JPPOWER":"JPPOWER.NS","PNB":"PNB.NS","UCOBANK":"UCOBANK.NS","IOB":"IOB.NS",
    "SAIL":"SAIL.NS","NMDC":"NMDC.NS","TRIDENT":"TRIDENT.NS","NETWORK18":"NETWORK18.NS",
    "TV18BRDCST":"TV18BRDCST.NS","SOUTHBANK":"SOUTHBANK.NS","UJJIVANSFB":"UJJIVANSFB.NS",
    "EQUITASBNK":"EQUITASBNK.NS","ESAFSFB":"ESAFSFB.NS","SPICEJET":"SPICEJET.NS","NFL":"NFL.NS",
    "CANBK":"CANBK.NS","BHEL":"BHEL.NS","IOC":"IOC.NS","HINDPETRO":"HINDPETRO.NS",
}
IT_S    = {"TCS":"TCS.NS","INFY":"INFY.NS","WIPRO":"WIPRO.NS","HCLTECH":"HCLTECH.NS","TECHM":"TECHM.NS",
           "LTTS":"LTTS.NS","PERSISTENT":"PERSISTENT.NS","COFORGE":"COFORGE.NS","ZENSARTECH":"ZENSARTECH.NS","MPHASIS":"MPHASIS.NS"}
BANK_S  = {"HDFCBANK":"HDFCBANK.NS","ICICIBANK":"ICICIBANK.NS","SBIN":"SBIN.NS","KOTAKBANK":"KOTAKBANK.NS",
           "AXISBANK":"AXISBANK.NS","PNB":"PNB.NS","BANKBARODA":"BANKBARODA.NS","FEDERALBNK":"FEDERALBNK.NS","YESBANK":"YESBANK.NS"}
POWER_S = {"NTPC":"NTPC.NS","POWERGRID":"POWERGRID.NS","ADANIGREEN":"ADANIGREEN.NS","ADANIPOWER":"ADANIPOWER.NS",
           "TATAPOWER":"TATAPOWER.NS","SUZLON":"SUZLON.NS","NHPC":"NHPC.NS","SJVN":"SJVN.NS",
           "JSWENERGY":"JSWENERGY.NS","RPOWER":"RPOWER.NS","PFC":"PFC.NS"}
DEF_S   = {"HAL":"HAL.NS","BEL":"BEL.NS","BDL":"BDL.NS","GRSE":"GRSE.NS","MAZDOCK":"MAZDOCK.NS",
           "COCHINSHIP":"COCHINSHIP.NS","BEML":"BEML.NS"}
AUTO_S  = {"MARUTI":"MARUTI.NS","TMPV":"TMPV.NS","M&M":"M&M.NS","BAJAJ-AUTO":"BAJAJ-AUTO.NS",
           "HEROMOTOCO":"HEROMOTOCO.NS","TVSMOTOR":"TVSMOTOR.NS","EICHERMOT":"EICHERMOT.NS","MRF":"MRF.NS","APOLLOTYRE":"APOLLOTYRE.NS"}
PHARMA_S= {"SUNPHARMA":"SUNPHARMA.NS","DRREDDY":"DRREDDY.NS","CIPLA":"CIPLA.NS","DIVISLAB":"DIVISLAB.NS",
           "LUPIN":"LUPIN.NS","AUROPHARMA":"AUROPHARMA.NS","APOLLOHOSP":"APOLLOHOSP.NS","MAXHEALTH":"MAXHEALTH.NS"}

ALL_STOCKS = {}
for _db in [NIFTY50,NEXT50,MIDCAP,BELOW100,IT_S,BANK_S,POWER_S,DEF_S,AUTO_S,PHARMA_S]:
    ALL_STOCKS.update(_db)
_seen=set();_clean={}
for _n,_s in ALL_STOCKS.items():
    if _s not in _seen:_seen.add(_s);_clean[_n]=_s
ALL_STOCKS=_clean

FAMOUS = {
    "RELIANCE": "RELIANCE.NS", "TCS": "TCS.NS", "HDFCBANK": "HDFCBANK.NS", "INFY": "INFY.NS",
    "ICICIBANK": "ICICIBANK.NS", "SBIN": "SBIN.NS", "AXISBANK": "AXISBANK.NS", "KOTAKBANK": "KOTAKBANK.NS",
    "BHARTIARTL": "BHARTIARTL.NS", "ITC": "ITC.NS", "LT": "LT.NS", "HINDUNILVR": "HINDUNILVR.NS",
    "BAJFINANCE": "BAJFINANCE.NS", "MARUTI": "MARUTI.NS", "TMPV": "TMPV.NS",
    "TATASTEEL": "TATASTEEL.NS", "JSWSTEEL": "JSWSTEEL.NS", "SUNPHARMA": "SUNPHARMA.NS",
    "TITAN": "TITAN.NS", "HCLTECH": "HCLTECH.NS", "WIPRO": "WIPRO.NS", "TECHM": "TECHM.NS",
    "NTPC": "NTPC.NS", "POWERGRID": "POWERGRID.NS", "ONGC": "ONGC.NS", "COALINDIA": "COALINDIA.NS",
    "TATAPOWER": "TATAPOWER.NS", "ADANIENT": "ADANIENT.NS", "ADANIPORTS": "ADANIPORTS.NS",
    "ADANIPOWER": "ADANIPOWER.NS", "ADANIGREEN": "ADANIGREEN.NS", "M&M": "M&M.NS",
    "BAJAJ-AUTO": "BAJAJ-AUTO.NS", "HEROMOTOCO": "HEROMOTOCO.NS", "TVSMOTOR": "TVSMOTOR.NS",
    "ULTRACEMCO": "ULTRACEMCO.NS", "ASIANPAINT": "ASIANPAINT.NS", "NESTLEIND": "NESTLEIND.NS",
    "BAJAJFINSV": "BAJAJFINSV.NS", "DRREDDY": "DRREDDY.NS", "CIPLA": "CIPLA.NS",
    "HAL": "HAL.NS", "BEL": "BEL.NS", "IRFC": "IRFC.NS", "RVNL": "RVNL.NS", "ETERNAL": "ETERNAL.NS", "JIOFIN": "JIOFIN.NS",
    "DLF": "DLF.NS", "VEDL": "VEDL.NS", "HINDALCO": "HINDALCO.NS", "GAIL": "GAIL.NS",
    "IOC": "IOC.NS", "BPCL": "BPCL.NS", "PNB": "PNB.NS", "BANKBARODA": "BANKBARODA.NS",
    "CANBK": "CANBK.NS", "INDUSINDBK": "INDUSINDBK.NS", "TRENT": "TRENT.NS", "DMART": "DMART.NS",
    "LICI": "LICI.NS", "SBILIFE": "SBILIFE.NS", "HDFCLIFE": "HDFCLIFE.NS", "SIEMENS": "SIEMENS.NS",
    "SUZLON": "SUZLON.NS", "IREDA": "IREDA.NS", "YESBANK": "YESBANK.NS", "IDEA": "IDEA.NS",
    "SAIL": "SAIL.NS", "NMDC": "NMDC.NS", "PFC": "PFC.NS", "RECLTD": "RECLTD.NS",
}

CAT_MAP = {
    "⭐ FAMOUS / MOST-TRADED": FAMOUS,
    "🌐 ALL NSE (live · every stock)": None,
    "🏅 NIFTY 50": NIFTY50, "📈 NEXT 50": NEXT50, "💎 MIDCAP": MIDCAP,
    "💰 BELOW ₹100": BELOW100, "🏆 ALL CURATED": ALL_STOCKS, "💻 IT": IT_S,
    "🏦 BANKING": BANK_S, "⚡ POWER": POWER_S, "🛡️ DEFENCE": DEF_S,
    "🚗 AUTO": AUTO_S, "💊 PHARMA": PHARMA_S,
}

ALIASES = {
    "SBI":"SBIN.NS","STATE BANK":"SBIN.NS","HDFC BANK":"HDFCBANK.NS","ICICI BANK":"ICICIBANK.NS",
    "INFOSYS":"INFY.NS","TATA MOTORS":"TMPV.NS","ETERNAL (ZOMATO)":"ETERNAL.NS","TATA STEEL":"TATASTEEL.NS","TATA POWER":"TATAPOWER.NS",
    "COAL INDIA":"COALINDIA.NS","MAHINDRA":"M&M.NS","BAJAJ FINANCE":"BAJFINANCE.NS","ASIAN PAINTS":"ASIANPAINT.NS",
    "AIRTEL":"BHARTIARTL.NS","HUL":"HINDUNILVR.NS","L&T":"LT.NS","KOTAK":"KOTAKBANK.NS","AXIS BANK":"AXISBANK.NS",
    "YES BANK":"YESBANK.NS","VODAFONE":"IDEA.NS","NESTLE":"NESTLEIND.NS","LIC":"LICI.NS","HDFC LIFE":"HDFCLIFE.NS",
    "SBI LIFE":"SBILIFE.NS","DR REDDY":"DRREDDY.NS","INOX WIND":"INOXWIND.NS","WAAREE":"WAAREEENER.NS",
    "REC":"RECLTD.NS","ZENSAR":"ZENSARTECH.NS","OIL INDIA":"OIL.NS","VEDANTA":"VEDL.NS",
    "BOSCH":"BOSCHLTD.NS","HERO":"HEROMOTOCO.NS","TVS":"TVSMOTOR.NS","EICHER":"EICHERMOT.NS",
}


@st.cache_data(ttl=86400, show_spinner=False)
def fetch_universe():
    """Full live NSE equity list (~2000 symbols) — searchable by name/ticker."""
    url = "https://archives.nseindia.com/content/equities/EQUITY_L.csv"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            raw = r.read().decode("utf-8", "ignore")
        udf = pd.read_csv(io.StringIO(raw))
        udf.columns = [c.strip() for c in udf.columns]
        name_map, sym_map = {}, {}
        for _, row in udf.iterrows():
            sym = str(row.get("SYMBOL", "")).strip().upper()
            nm = str(row.get("NAME OF COMPANY", "")).strip()
            if sym:
                sym_map[sym] = sym + ".NS"
                name_map[f"{nm} ({sym})"] = sym + ".NS"
        if sym_map:
            return {"name_map": name_map, "sym_map": sym_map,
                    "count": len(sym_map), "live": True}
    except Exception:
        pass
    return {"name_map": {k: v for k, v in ALL_STOCKS.items()},
            "sym_map": {k.upper(): v for k, v in ALL_STOCKS.items()},
            "count": len(ALL_STOCKS), "live": False}


def universe_search(query, limit=40):
    from difflib import SequenceMatcher
    q = query.strip().upper()
    if len(q) < 2:
        return {}
    uni = fetch_universe()["name_map"]
    qtokens = [t for t in re.split(r"[^A-Z0-9&]+", q) if t]

    def word_hit(tok, words):
        for w in words:
            if not w:
                continue
            if w.startswith(tok) or tok.startswith(w):
                return 1.0
            if len(tok) >= 3 and tok in w:
                return 0.92
            if SequenceMatcher(None, tok, w).ratio() >= 0.8:
                return 0.85
        return 0.0

    scored = []
    for disp, sym in uni.items():
        du = disp.upper()
        base = sym.replace(".NS", "").replace(".BO", "")
        if q == base or q == du:
            scored.append((3.0, disp, sym)); continue
        words = [w for w in re.split(r"[^A-Z0-9&]+", du) if w]
        hits = [word_hit(t, words) for t in qtokens]
        if hits and all(h > 0 for h in hits):
            score = sum(hits) / len(hits)
            if q in du:
                score += 0.5
            if base.startswith(q):
                score += 0.3
            scored.append((score, disp, sym))
    scored.sort(key=lambda x: (-x[0], len(x[1])))
    return {d: s for _, d, s in scored[:limit]}


# ============================================================
# UTILITIES
# ============================================================
def mkt_status():
    n = now_ist()
    if n.weekday() >= 5:
        return "closed", "🔴 CLOSED", "Monday 9:15 AM"
    t = n.time()
    if t < dtime(9, 0):    return "pre",  "🌅 PRE-MARKET", "Opens 9:15 AM"
    if t < dtime(9, 15):   return "pre",  "🌅 PRE-OPEN",   "Opens very soon!"
    if t <= dtime(15, 30): return "open", "🟢 MARKET LIVE","Closes 3:30 PM"
    return "closed", "🔴 CLOSED", "Opens 9:15 AM tomorrow"


def session_time_context():
    n = now_ist()
    t = n.time()
    close = n.replace(hour=15, minute=30, second=0, microsecond=0)
    sqoff = n.replace(hour=15, minute=15, second=0, microsecond=0)
    mins_left = max(int((close - n).total_seconds() // 60), 0)
    mins_sq = max(int((sqoff - n).total_seconds() // 60), 0)
    hh, mm = mins_sq // 60, mins_sq % 60
    left_txt = (f"{hh}h {mm}m" if hh else f"{mm}m") + " to square-off (3:15 PM)"
    if t <= dtime(11, 0):
        phase = "early"; can_enter = True; main_t = "T2"
        note = f"Early session · plenty of time ({left_txt}). Fresh entries fine — aim T2/T3."
    elif t <= dtime(13, 0):
        phase = "mid"; can_enter = True; main_t = "T2"
        note = f"Midday · about half the session left ({left_txt}). Enter only on confirmation, aim T1–T2."
    elif t <= dtime(14, 30):
        phase = "late"; can_enter = True; main_t = "T1"
        note = f"Late session · only {left_txt}. Quick trade only — take T1, don't wait for T2/T3. Trail tight."
    elif t <= dtime(15, 15):
        phase = "closing"; can_enter = False; main_t = "T1"
        note = f"⏰ Closing hour · {left_txt}. Too late for a fresh intraday buy — manage/exit what you hold, square off by 3:15."
    else:
        phase = "done"; can_enter = False; main_t = "T1"
        note = "Session basically over — this plan is for TOMORROW. Nothing to enter now."
    return {"phase": phase, "mins_left": mins_left, "mins_sq": mins_sq,
            "left_txt": left_txt, "note": note, "can_enter": can_enter, "main_t": main_t}


def safe(x, d=np.nan):
    try:
        v = float(x)
        return d if (np.isnan(v) or np.isinf(v)) else v
    except Exception:
        return d


def fmt_num(n, prefix="₹"):
    if abs(n) >= 1e7:  return f"{prefix}{n/1e7:.2f}Cr"
    if abs(n) >= 1e5:  return f"{prefix}{n/1e5:.2f}L"
    if abs(n) >= 1000: return f"{prefix}{n:,.0f}"
    return f"{prefix}{n:.2f}"


@st.cache_data(ttl=300, show_spinner=False)
def do_search(q):
    q = q.strip().upper()
    out = {}
    if len(q) < 1:
        return out
    hits = universe_search(q, limit=10)
    cands = list(hits.items())
    for alias, sym in ALIASES.items():
        if q in alias:
            cands.insert(0, (alias, sym))
    for sfx in ['.NS', '.BO', '']:
        cands.append((q, q.replace(' ', '') + sfx))
    seen = set()
    for label, sym in cands:
        if sym in seen:
            continue
        seen.add(sym)
        try:
            info = yf.Ticker(sym).info
            p = (info.get('regularMarketPrice') or info.get('currentPrice')
                 or info.get('previousClose', 0))
            if p and float(p) > 0:
                nm = info.get('longName', info.get('shortName', sym))
                out[f"{nm} [{sym}]"] = {'sym': sym, 'name': nm, 'price': float(p)}
        except Exception:
            pass
        if len(out) >= 8:
            break
    return out


# ============================================================
# PROFIT CALCULATOR (from your v13 — kept as-is)
# ============================================================
def calc_money(price, capital, sl, t1, t2, t3):
    if price <= 0 or capital <= 0:
        return {}
    qty = int(capital / price)
    if qty <= 0:
        return {}
    invested = round(qty * price, 2)
    leftover = round(capital - invested, 2)

    def pnl(t): return round(qty * (t - price), 2)
    def move(t): return round(t - price, 2)
    def pct(t): return round((t - price) / price * 100, 2) if price > 0 else 0
    def roi(t): return round(pnl(t) / capital * 100, 2) if capital > 0 else 0
    def sv(t): return round(qty * t, 2)

    loss = round(qty * (price - sl), 2)
    loss_pct = round((price - sl) / price * 100, 2) if price > 0 else 0
    loss_roi = round(loss / capital * 100, 2) if capital > 0 else 0
    return {
        'qty': qty, 'invested': invested, 'leftover': leftover,
        'p1': pnl(t1), 'p2': pnl(t2), 'p3': pnl(t3), 'loss': loss,
        'mv1': move(t1), 'mv2': move(t2), 'mv3': move(t3), 'mv_sl': round(price - sl, 2),
        'pct1': pct(t1), 'pct2': pct(t2), 'pct3': pct(t3), 'pct_sl': loss_pct,
        'roi1': roi(t1), 'roi2': roi(t2), 'roi3': roi(t3), 'roi_loss': loss_roi,
        'sv1': sv(t1), 'sv2': sv(t2), 'sv3': sv(t3), 'sv_sl': sv(sl),
        't1': t1, 't2': t2, 't3': t3, 'sl': sl,
    }


def target_feasibility(money, target):
    if not money or money.get('qty', 0) == 0:
        return {'label': '❌ SKIP', 'color': '#dc2626', 'bg': '#fff1f2',
                'msg': 'Cannot buy even 1 share. Choose cheaper stock.', 'at': None}
    p1, p2, p3 = money['p1'], money['p2'], money['p3']
    if p1 >= target:
        return {'label': '🟢 EASY', 'color': '#16a34a', 'bg': '#f0fdf4',
                'msg': f"T1 gives ₹{p1:.0f}! ₹{round(p1-target,0):.0f} more than target!", 'at': 'T1', 'ap': p1}
    if p2 >= target:
        return {'label': '🔵 ACHIEVABLE', 'color': '#2563eb', 'bg': '#eff6ff',
                'msg': f"T2 gives ₹{p2:.0f}. Need {money['pct2']:.1f}% move.", 'at': 'T2', 'ap': p2}
    if p3 >= target:
        return {'label': '🟣 STRETCH', 'color': '#9333ea', 'bg': '#faf5ff',
                'msg': f"T3 gives ₹{p3:.0f}. Need {money['pct3']:.1f}% move.", 'at': 'T3', 'ap': p3}
    return {'label': '❌ NOT ENOUGH', 'color': '#dc2626', 'bg': '#fff1f2',
            'msg': f"Max profit ₹{p3:.0f}. Need cheaper stock or more capital.", 'at': None, 'ap': p3}


# ============================================================
# LEVEL CALCULATORS — Fibonacci + Standard/Camarilla/Woodie pivots
# ============================================================
def calc_fibonacci(high, low):
    diff = high - low
    levels = {
        '0% (High)': round(high, 2), '23.6%': round(high - 0.236 * diff, 2),
        '38.2%': round(high - 0.382 * diff, 2), '50%': round(high - 0.5 * diff, 2),
        '61.8%': round(high - 0.618 * diff, 2), '78.6%': round(high - 0.786 * diff, 2),
        '100% (Low)': round(low, 2),
    }
    ext = {'127.2%': round(high + 0.272 * diff, 2), '161.8%': round(high + 0.618 * diff, 2),
           '261.8%': round(high + 1.618 * diff, 2)}
    return levels, ext


def calc_all_pivots(ph, pl, pc):
    d = max(ph - pl, 0.01)
    pp = (ph + pl + pc) / 3
    standard = {
        'R3': round(pp + 2 * d, 2), 'R2': round(pp + d, 2), 'R1': round(2 * pp - pl, 2),
        'PP': round(pp, 2),
        'S1': round(2 * pp - ph, 2), 'S2': round(pp - d, 2), 'S3': round(pp - 2 * d, 2),
    }
    camarilla = {
        'R4': round(pc + d * 1.1 / 2, 2), 'R3': round(pc + d * 1.1 / 4, 2),
        'R2': round(pc + d * 1.1 / 6, 2), 'R1': round(pc + d * 1.1 / 12, 2),
        'PP': round(pp, 2),
        'S1': round(pc - d * 1.1 / 12, 2), 'S2': round(pc - d * 1.1 / 6, 2),
        'S3': round(pc - d * 1.1 / 4, 2), 'S4': round(pc - d * 1.1 / 2, 2),
    }
    wp = (ph + pl + 2 * pc) / 4
    woodie = {
        'R2': round(wp + d, 2), 'R1': round(2 * wp - pl, 2), 'PP': round(wp, 2),
        'S1': round(2 * wp - ph, 2), 'S2': round(wp - d, 2),
    }
    return {'Standard': standard, 'Camarilla': camarilla, 'Woodie': woodie}


def demark_pivot(ph, pl, pc, po):
    try:
        if pc < po:   x = ph + 2 * pl + pc
        elif pc > po: x = 2 * ph + pl + pc
        else:         x = ph + pl + 2 * pc
        pp = x / 4
        return {"PP": round(pp, 2), "R1": round(x / 2 - pl, 2), "S1": round(x / 2 - ph, 2)}
    except Exception:
        return None


# ============================================================
# NSE LIVE EXTRAS — circuit limits · pre-open (IEP) · HTF trend
# ============================================================
def _nse_open():
    import http.cookiejar, urllib.request
    h = {"User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                        "(KHTML, like Gecko) Chrome/120.0 Safari/537.36"),
         "Accept": "application/json, text/plain, */*", "Accept-Language": "en-US,en;q=0.9"}
    cj = http.cookiejar.CookieJar()
    op = urllib.request.build_opener(urllib.request.HTTPCookieProcessor(cj))
    op.open(urllib.request.Request("https://www.nseindia.com", headers=h), timeout=8).read()
    return op, h


def _num(x):
    try:
        s = str(x).replace(",", "").strip()
        return float(s) if s not in ("", "-", "None") else None
    except Exception:
        return None


@st.cache_data(ttl=120, show_spinner=False)
def get_nse_quote(sym):
    if sym.startswith("^"):
        return None
    ns = sym.replace(".NS", "").replace(".BO", "").upper()
    try:
        import json, urllib.request
        op, h = _nse_open()
        h = {**h, "Referer": f"https://www.nseindia.com/get-quotes/equity?symbol={ns}"}
        url = f"https://www.nseindia.com/api/quote-equity?symbol={ns}"
        raw = op.open(urllib.request.Request(url, headers=h), timeout=8).read().decode("utf-8", "ignore")
        d = json.loads(raw)
        pi = d.get("priceInfo", {}) or {}
        pm = d.get("preOpenMarket", {}) or {}
        prev = _num(pi.get("previousClose"))
        ltp = _num(pi.get("lastPrice"))
        uc = _num(pi.get("upperCP")); lc = _num(pi.get("lowerCP"))
        band = str(pi.get("pPriceBand", "")).strip()
        iep = _num(pm.get("IEP"))
        pchg = _num(pm.get("Change"))
        ppct = _num(pm.get("perChange"))
        whl = pi.get("weekHighLow", {}) or {}
        wk_hi = _num(whl.get("max")); wk_lo = _num(whl.get("min"))
        if iep and pchg is None and prev:
            pchg = round(iep - prev, 2); ppct = round((iep - prev) / prev * 100, 2)
        no_band = band.lower().startswith("no") or band == ""
        band_num = None if no_band else _num(band)

        def tick(x):
            return round(round(x / 0.05) * 0.05, 2) if x else x

        upper = lower = None
        if not no_band and prev and band_num:
            upper_calc = tick(prev * (1 + band_num / 100))
            lower_calc = tick(prev * (1 - band_num / 100))
            upper = uc if (uc and abs(uc - upper_calc) / upper_calc < 0.02) else upper_calc
            lower = lc if (lc and abs(lc - lower_calc) / lower_calc < 0.02) else lower_calc
        elif not no_band:
            upper, lower = uc, lc
            if uc is None and lc is None:
                no_band = True
        return {
            "prev": prev, "upper": upper, "lower": lower,
            "band": "No band" if no_band else (f"{band_num:g}%" if band_num else band),
            "no_band": no_band, "iep": iep, "iep_chg": pchg, "iep_pct": ppct,
            "wk_hi": wk_hi, "wk_lo": wk_lo, "ltp": ltp,
        }
    except Exception:
        return None


@st.cache_data(ttl=180, show_spinner=False)
def htf_trend(sym, iv):
    nxt = {"1m": ("15m", "15-min"), "5m": ("1h", "1-hour"), "15m": ("1h", "1-hour"),
           "30m": ("1d", "daily"), "1h": ("1d", "daily"), "1d": ("1wk", "weekly")}
    per = {"15m": "1mo", "1h": "3mo", "1d": "1y", "1wk": "2y"}
    hi, lbl = nxt.get(iv, ("1d", "daily"))
    try:
        d = yf.Ticker(sym).history(period=per.get(hi, "3mo"), interval=hi).dropna()
        if len(d) < 55:
            return {"trend": None, "label": lbl}
        c = d["Close"]
        e20 = ta.trend.ema_indicator(c, 20); e50 = ta.trend.ema_indicator(c, 50)
        p = float(c.iloc[-1]); a = float(e20.iloc[-1]); b = float(e50.iloc[-1])
        t = "UP" if (a > b and p > a) else "DOWN" if (a < b and p < a) else "FLAT"
        return {"trend": t, "label": lbl}
    except Exception:
        return {"trend": None, "label": lbl}


# ============================================================
# DAILY SERIES FOR LEVELS / ML
# ============================================================
@st.cache_data(ttl=1800, show_spinner=False)
def _daily_for_ml(sym, years="2y"):
    try:
        d = yf.Ticker(sym).history(period=years, interval="1d").dropna()
        return d if len(d) >= 120 else None
    except Exception:
        return None


def get_sr_fib(sym, price, atr, circuit=None):
    try:
        daily = _daily_for_ml(sym)
        if daily is None or len(daily) < 20:
            daily = yf.Ticker(sym).history(period="6mo", interval="1d").dropna()
        if daily is None or len(daily) < 5:
            return None
        daily = daily[daily['Low'] > 0].copy()
        rng_ok = (daily['High'] - daily['Low']) < daily['Close'] * 0.5
        if rng_ok.sum() >= len(daily) * 0.7:
            daily = daily[rng_ok]
        ph = float(daily['High'].iloc[-1]); pl = float(daily['Low'].iloc[-1])
        pc = float(daily['Close'].iloc[-1])
        prev_close = float(daily['Close'].iloc[-2]) if len(daily) >= 2 else pc
        gap_pct = (price - prev_close) / prev_close * 100 if prev_close > 0 else 0
        wk = daily.tail(252)
        hi52 = float(wk['High'].max()); lo52 = float(wk['Low'].min())
        sw5_hi = float(daily['High'].tail(5).max());  sw5_lo = float(daily['Low'].tail(5).min())
        sw20_hi = float(daily['High'].tail(20).max()); sw20_lo = float(daily['Low'].tail(20).min())
        sw20_lo = max(sw20_lo, lo52); sw20_hi = min(sw20_hi, hi52)
        pivots = calc_all_pivots(ph, pl, pc)
        pvt = pivots['Standard']
        fib_levels, fib_ext = calc_fibonacci(sw20_hi, sw20_lo)
        raw = [
            ("52W HIGH", hi52, "R"), ("R3", pvt['R3'], "R"), ("R2", pvt['R2'], "R"),
            ("20D HIGH", sw20_hi, "R"), ("PREV HIGH", ph, "R"), ("5D HIGH", sw5_hi, "R"),
            ("R1", pvt['R1'], "R"), ("PIVOT", pvt['PP'], "P"), ("S1", pvt['S1'], "S"),
            ("5D LOW", sw5_lo, "S"), ("PREV LOW", pl, "S"), ("20D LOW", sw20_lo, "S"),
            ("S2", pvt['S2'], "S"), ("S3", pvt['S3'], "S"), ("52W LOW", lo52, "S"),
        ]
        if circuit and circuit.get("upper"):
            raw.append(("UPPER CIRCUIT", circuit["upper"], "R"))
        if circuit and circuit.get("lower"):
            raw.append(("LOWER CIRCUIT", circuit["lower"], "S"))
        levels = []
        for nm, lp, tp0 in raw:
            if lp is None or lp <= 0:
                continue
            if lp < lo52 and tp0 == "S" and ("52W" not in nm and "CIRCUIT" not in nm):
                continue
            tp = "P" if tp0 == "P" else ("R" if lp > price else "S")
            dist = round((lp - price) / price * 100, 2) if price > 0 else 0
            levels.append({'name': nm, 'price': round(float(lp), 2), 'type': tp, 'dist': dist})
        levels.sort(key=lambda x: x['price'], reverse=True)
        deduped = []
        for l in levels:
            if deduped and abs(deduped[-1]['price'] - l['price']) / price < 0.0025:
                deduped[-1]['name'] = deduped[-1]['name'] + " / " + l['name']
                continue
            deduped.append(l)
        levels = deduped
        above = [l for l in levels if l['price'] > price]
        below = [l for l in levels if l['price'] < price]
        nr = min(above, key=lambda x: x['price']) if above else None
        ns = max(below, key=lambda x: x['price']) if below else None
        return {
            'levels': levels, 'pvt': pvt, 'pivots': pivots,
            'prev_close': round(prev_close, 2), 'gap_pct': round(gap_pct, 2),
            'nr': nr, 'ns': ns,
            'ph': ph, 'pl': pl, 'pc': pc, 'sw5_hi': sw5_hi, 'sw5_lo': sw5_lo,
            'sw20_hi': sw20_hi, 'sw20_lo': sw20_lo, 'hi52': round(hi52, 2), 'lo52': round(lo52, 2),
            'hi_all': hi52, 'lo_all': lo52,
            'fib_levels': fib_levels, 'fib_ext': fib_ext, 'daily': daily,
        }
    except Exception:
        return None


# ============================================================
# TRADE PLAN GENERATOR
# ============================================================
def make_plan(price, atr, sr, sig, session='closed', circuit=None):
    prev_close = sr['prev_close'] if sr else price
    gap_pct = sr['gap_pct'] if sr else 0
    ns_price = sr['ns']['price'] if sr and sr['ns'] else round(price - atr * 1.5, 2)
    nr_price = sr['nr']['price'] if sr and sr['nr'] else round(price + atr * 2.0, 2)
    big_gap_up = gap_pct > 2.0
    big_gap_dn = gap_pct < -2.0
    near_resist = sr and sr['nr'] and abs(sr['nr']['dist']) < 0.8
    is_buy = 'BUY' in sig
    is_sell = 'SELL' in sig

    if big_gap_up and is_buy:
        sit = "GAP_UP"; act = "⚠️ GAP UP — WAIT FOR PULLBACK"
        buy_at = round(prev_close * 1.005, 2); sl = round(prev_close * 0.985, 2); ac = "#ea580c"
        msg = f"Gapped {gap_pct:+.1f}% from ₹{prev_close:.2f}. Don't chase! Wait for ₹{buy_at:.2f}"
        timing = "Watch till 10:30 AM. No pullback → SKIP today"
    elif big_gap_dn:
        sit = "GAP_DN"; act = "⛔ GAP DOWN — AVOID"
        buy_at = round(ns_price * 1.01, 2); sl = round(ns_price * 0.98, 2); ac = "#dc2626"
        msg = f"Fell {gap_pct:.1f}%. Risky. Support: ₹{ns_price:.2f}"
        timing = "Wait till 11 AM. If support holds → reconsider"
    elif near_resist and is_buy and 'STRONG' not in sig:
        sit = "NEAR_R"; act = "🚧 NEAR RESISTANCE — WAIT BREAKOUT"
        buy_at = round(nr_price * 1.002, 2); sl = round(nr_price * 0.985, 2); ac = "#d97706"
        msg = f"Resistance ₹{nr_price:.2f} is only {abs(sr['nr']['dist']):.1f}% away"
        timing = f"Buy ONLY if 15-min candle closes ABOVE ₹{nr_price:.2f} with volume"
    elif is_buy:
        sup_dist = (price - ns_price) / price * 100 if price > 0 else 0
        if sup_dist > 2.5:
            sit = "BUY_DIP"; act = "⏳ BUY ON PULLBACK TO SUPPORT"
            buy_at = round(ns_price * 1.003, 2)
            sl = round(ns_price * 0.992, 2)
            ac = "#d97706"
            msg = (f"Price ₹{price:.2f} is {sup_dist:.1f}% above support — buying here needs a wide stop. "
                   f"Place a BUY near support ₹{buy_at:.2f}; it fills only if price dips there. Stop ₹{sl:.2f}.")
            timing = "Set a buy-limit at support. If price never dips there, skip it — don't chase."
        else:
            sit = "BUY_NOW"; act = "✅ STRONG BUY NOW" if 'STRONG' in sig else "✅ BUY NOW"
            buy_at = round(price, 2)
            sl = round(max(ns_price * 0.992, price - atr * 1.5), 2); ac = "#16a34a"
            msg = "Price is near support — buy at market with the stop just below support."
            timing = "Enter 9:30 AM–2:30 PM | Best: 9:45–11:00 AM"
    elif is_sell:
        sit = "SELL"; act = "⛔ DO NOT BUY — SELL SIGNAL"
        buy_at = round(price, 2); sl = round(price - atr * 1.5, 2); ac = "#dc2626"
        msg = "Selling pressure. Wrong time to buy."
        timing = "Wait for reversal. Check tomorrow."
    else:
        sit = "WAIT"; act = "⏸ NEUTRAL — WAIT FOR SIGNAL"
        buy_at = round(price * 0.997, 2); sl = round(ns_price * 0.992, 2); ac = "#d97706"
        msg = "No clear direction yet."
        timing = "Check again in 1-2 hours."

    risk = max(buy_at - sl, buy_at * 0.003)
    t1 = round(buy_at + risk * 1.0, 2)
    t2 = round(buy_at + risk * 2.0, 2)
    t3 = round(buy_at + risk * 3.5, 2)
    if sr and sr.get('levels'):
        ra = sorted([l['price'] for l in sr['levels']
                     if l['type'] == 'R' and l['price'] > buy_at * 1.003])
        if len(ra) >= 1: t1 = round(ra[0], 2)
        if len(ra) >= 2: t2 = round(ra[1], 2)
        if len(ra) >= 3: t3 = round(ra[2], 2)
    t2 = max(t2, t1)
    t3 = max(t3, t2)
    if circuit and circuit.get("upper"):
        t1 = min(t1, circuit["upper"]); t2 = min(t2, circuit["upper"]); t3 = min(t3, circuit["upper"])
    if circuit and circuit.get("lower"):
        sl = max(sl, circuit["lower"])

    if session == 'open':
        now = datetime.now()
        mins_left = max(0, (15 * 60 + 30) - (now.hour * 60 + now.minute))
        hh = mins_left // 60; mm2 = mins_left % 60
        left_txt = f"{hh}h {mm2}m left today"
        t = now.time()
        when = f"TODAY (live · {left_txt})"
        head = f"📈 TODAY'S PLAN — {left_txt}"
        if t < dtime(9, 45):
            phase = "OPENING"
            timing = "First 30 min is volatile — wait for the 9:45 candle before entering."
        elif t < dtime(11, 30):
            phase = "MORNING"
            timing = f"Prime window. Full setup valid — T1→T2→T3 all in play ({left_txt})."
        elif t < dtime(13, 30):
            phase = "MIDDAY"
            timing = f"Midday lull — moves are slower. T1–T2 realistic today ({left_txt})."
        elif t < dtime(14, 45):
            phase = "LATE"
            timing = (f"Only ~{left_txt}. Take it ONLY on strong momentum — aim T1, maybe T2. "
                      f"T3 unlikely today. Tighten stop, square off by 3:15.")
            act = act + " · quick move only"
        else:
            phase = "CLOSING"
            timing = "Too late for a fresh intraday entry — square off open trades by 3:15. Plan this for tomorrow instead."
            if 'BUY' in sig:
                act = "🕒 TOO LATE TODAY — PLAN FOR TOMORROW"
                ac = "#b45309"
                msg = "Signal is fine, but not enough time left for a fresh intraday trade. Watch it at tomorrow's open."
    elif session == 'pre':
        when = "TODAY (pre-open)"
        head = "🌅 TODAY'S PLAN — market opens 9:15 AM"
        phase = "PREOPEN"
        timing = "Market opens at 9:15 AM today. Let the 9:15–9:30 first candle print, then act on the level below — don't buy at the very open."
    else:
        when = "TOMORROW"
        head = "🌙 TOMORROW'S PLAN — for the next session"
        phase = "NEXTDAY"
        if sit == 'BUY_NOW':
            timing = "Tomorrow: wait for the 9:30 first candle, then buy near the level below."

    sl_pct = round((buy_at - sl) / buy_at * 100, 2) if buy_at > 0 else 0
    t1p = round((t1 - buy_at) / buy_at * 100, 2) if buy_at > 0 else 0
    t2p = round((t2 - buy_at) / buy_at * 100, 2) if buy_at > 0 else 0
    t3p = round((t3 - buy_at) / buy_at * 100, 2) if buy_at > 0 else 0
    rr = round((t2 - buy_at) / max(buy_at - sl, 0.01), 1)
    return {
        'sit': sit, 'act': act, 'ac': ac, 'buy_at': buy_at, 'sl': sl,
        't1': t1, 't2': t2, 't3': t3, 'risk': round(risk, 2), 'msg': msg, 'timing': timing,
        'gap_pct': gap_pct, 'prev_close': prev_close, 'nr_price': nr_price, 'ns_price': ns_price,
        'sl_pct': sl_pct, 't1p': t1p, 't2p': t2p, 't3p': t3p, 'rr': rr,
        'when': when, 'head': head, 'phase': phase,
    }


# ============================================================
# CORE ANALYSIS ENGINE — 12 INDICATORS (pure function on a df)
# ============================================================
def compute_signals(df_in):
    """All indicator maths from a raw OHLCV dataframe. Used by BOTH the
    single-stock analysis and the 500-stock dashboard (bulk data)."""
    try:
        if df_in is None or df_in.empty or len(df_in) < 10:
            return None
        df = df_in.copy()
        df.columns = [str(c).title() for c in df.columns]
        if not {'Open', 'High', 'Low', 'Close'}.issubset(df.columns):
            return None
        if 'Volume' not in df.columns:
            df['Volume'] = 0
        df = df.dropna(subset=['Open', 'High', 'Low', 'Close'])
        if len(df) < 10:
            return None
    except Exception:
        return None
    n = len(df)
    try:
        for w, c in [(9, "E9"), (20, "E20"), (50, "E50"), (200, "E200")]:
            df[c] = ta.trend.ema_indicator(df['Close'], w) if n > w else np.nan
        for w, c in [(20, "S20"), (50, "S50")]:
            df[c] = ta.trend.sma_indicator(df['Close'], w) if n > w else np.nan
        df['RSI'] = ta.momentum.rsi(df['Close'], 14) if n >= 15 else np.nan
        if n >= 35:
            mc = ta.trend.MACD(df['Close'])
            df['MACD'] = mc.macd(); df['MS'] = mc.macd_signal(); df['MH'] = mc.macd_diff()
        else:
            df['MACD'] = df['MS'] = df['MH'] = np.nan
        if n >= 20:
            bb = ta.volatility.BollingerBands(df['Close'], 20, 2)
            df['BBU'] = bb.bollinger_hband(); df['BBL'] = bb.bollinger_lband()
            df['BBM'] = bb.bollinger_mavg(); df['BBW'] = bb.bollinger_wband(); df['BBP'] = bb.bollinger_pband()
        else:
            df['BBU'] = df['BBL'] = df['BBM'] = df['BBW'] = df['BBP'] = np.nan
        if n >= 14:
            sto = ta.momentum.StochasticOscillator(df['High'], df['Low'], df['Close'])
            df['SK'] = sto.stoch(); df['SD'] = sto.stoch_signal()
        else:
            df['SK'] = df['SD'] = np.nan
        if n >= 30:
            adx_i = ta.trend.ADXIndicator(df['High'], df['Low'], df['Close'], 14)
            df['ADX'] = adx_i.adx(); df['DMP'] = adx_i.adx_pos(); df['DMN'] = adx_i.adx_neg()
        else:
            df['ADX'] = df['DMP'] = df['DMN'] = np.nan
        df['ATR'] = (ta.volatility.average_true_range(df['High'], df['Low'], df['Close'], 14)
                     if n >= 15 else df['High'] - df['Low'])
        df['MFI'] = (ta.volume.MFIIndicator(df['High'], df['Low'], df['Close'], df['Volume'], 14).money_flow_index()
                     if n >= 14 else np.nan)
        df['CCI'] = (ta.trend.CCIIndicator(df['High'], df['Low'], df['Close'], 20).cci() if n >= 20 else np.nan)
        df['WR'] = (ta.momentum.WilliamsRIndicator(df['High'], df['Low'], df['Close'], 14).williams_r()
                    if n >= 14 else np.nan)
        df['ROC'] = (ta.momentum.ROCIndicator(df['Close'], 12).roc() if n >= 12 else np.nan)
        try:
            df['VWAP'] = ((df['Volume'] * (df['High'] + df['Low'] + df['Close']) / 3).cumsum()
                          / df['Volume'].cumsum())
        except Exception:
            df['VWAP'] = np.nan
        df['STD20'] = df['Close'].rolling(20).std() if n >= 20 else np.nan
        if n >= 15:
            h2 = (df['High'] + df['Low']) / 2
            atr_ = ta.volatility.average_true_range(df['High'], df['Low'], df['Close'], 10)
            ub = h2 + 3 * atr_; lb = h2 - 3 * atr_
            sv = [np.nan] * n; sd = [0] * n
            for i in range(10, n):
                if df['Close'].iloc[i] > ub.iloc[i-1]:
                    sd[i] = 1; sv[i] = float(lb.iloc[i])
                elif df['Close'].iloc[i] < lb.iloc[i-1]:
                    sd[i] = -1; sv[i] = float(ub.iloc[i])
                else:
                    sd[i] = sd[i-1] if i > 10 else 1
                    pv = sv[i-1]
                    if sd[i] == 1:
                        cv = float(lb.iloc[i]); sv[i] = max(cv, pv) if not np.isnan(pv) else cv
                    else:
                        cv = float(ub.iloc[i]); sv[i] = min(cv, pv) if not np.isnan(pv) else cv
            df['ST'] = sv; df['STD'] = sd
        else:
            df['ST'] = np.nan; df['STD'] = 0
        df['OBV'] = ta.volume.OnBalanceVolumeIndicator(df['Close'], df['Volume']).on_balance_volume() if n >= 5 else np.nan
    except Exception:
        pass

    lat = df.iloc[-1]; prev = df.iloc[-2] if n >= 2 else lat
    price = float(lat['Close']); g = lambda c: safe(lat.get(c))
    bs = 0; ss = 0; sigs = []

    rsi = g('RSI')
    if not np.isnan(rsi):
        if rsi < 25: bs += 22; t = "Deep Oversold — Strong BUY"
        elif rsi < 35: bs += 15; t = "Oversold — BUY zone"
        elif rsi < 45: bs += 8; t = "Lean BUY"
        elif rsi < 55: bs += 4; ss += 4; t = "Neutral"
        elif rsi < 65: ss += 8; t = "Lean SELL"
        elif rsi < 75: ss += 15; t = "Overbought — SELL zone"
        else: ss += 22; t = "Extreme Overbought — Strong SELL"
        sigs.append({'n': 'RSI', 'v': f"{rsi:.1f}", 't': t, 'b': rsi < 50, 'cat': 'Momentum'})

    mv = g('MACD'); ms_ = g('MS'); mh = g('MH'); pmh = safe(prev.get('MH'))
    if not np.isnan(mv) and not np.isnan(ms_):
        cross_up = (mv > ms_ and not np.isnan(pmh) and not np.isnan(mh) and pmh < 0 < mh)
        cross_dn = (mv < ms_ and not np.isnan(pmh) and not np.isnan(mh) and pmh > 0 > mh)
        if cross_up: bs += 22; t = "🚀 Fresh Bullish Cross — Strong BUY!"
        elif mv > ms_: bs += 12; t = "MACD above Signal — BUY"
        elif cross_dn: ss += 22; t = "🔻 Fresh Bearish Cross — Strong SELL!"
        else: ss += 12; t = "MACD below Signal — SELL"
        sigs.append({'n': 'MACD', 'v': f"{mv:.3f}", 't': t, 'b': mv > ms_, 'cat': 'Trend'})

    ab = be = tot = 0
    for cn in ['E9', 'E20', 'E50', 'S20', 'S50']:
        v = g(cn)
        if not np.isnan(v):
            tot += 1; ab += 1 if price > v else 0; be += 1 if price <= v else 0
    if tot > 0:
        pct = ab / tot
        if pct >= 0.85: bs += 20; t = f"Above all {tot} MAs — Strong BUY"
        elif pct >= 0.60: bs += 13; t = f"Above {ab}/{tot} MAs — BUY"
        elif pct <= 0.15: ss += 20; t = f"Below all {tot} MAs — Strong SELL"
        elif pct <= 0.40: ss += 13; t = f"Below {be}/{tot} MAs — SELL"
        else: bs += 5; ss += 5; t = f"Mixed {ab}/{tot} MAs — Neutral"
        sigs.append({'n': 'Moving Avg', 'v': f"{ab}/{tot}", 't': t, 'b': ab > be, 'cat': 'Trend'})

    bu = g('BBU'); bl = g('BBL'); bw = g('BBW'); bp_val = g('BBP')
    if not np.isnan(bu) and not np.isnan(bl):
        rng = max(bu - bl, 0.01); pos = (price - bl) / rng
        sq = bw < 1.5 if not np.isnan(bw) else False
        if pos < 0.05: bs += 18; t = "At Lower Band — BUY zone"
        elif pos < 0.25: bs += 10; t = "Near Lower Band — Lean BUY"
        elif pos > 0.95: ss += 18; t = "At Upper Band — SELL zone"
        elif pos > 0.75: ss += 10; t = "Near Upper Band — Lean SELL"
        else: bs += 4; ss += 4; t = "Mid Band — Neutral"
        if sq: t += " ⚡ SQUEEZE!"
        sigs.append({'n': 'Bollinger Bands', 'v': f"{pos:.0%}", 't': t, 'b': pos < 0.5, 'cat': 'Volatility'})

    std_ = safe(lat.get('STD'), 0); stv = g('ST')
    if not np.isnan(stv) and std_ != 0:
        if std_ == 1: bs += 20; t = "✅ Supertrend BUY — Uptrend confirmed"
        else: ss += 20; t = "❌ Supertrend SELL — Downtrend confirmed"
        sigs.append({'n': 'Supertrend', 'v': f"₹{stv:.2f}", 't': t, 'b': std_ == 1, 'cat': 'Trend'})

    vw = g('VWAP')
    if not np.isnan(vw):
        pv = (price - vw) / vw * 100
        if pv > 1.0: bs += 12; t = f"Strong above VWAP +{pv:.1f}%"
        elif pv > 0: bs += 7; t = f"Above VWAP +{pv:.1f}%"
        elif pv < -1: ss += 12; t = f"Strong below VWAP {pv:.1f}%"
        else: ss += 7; t = f"Below VWAP {pv:.1f}%"
        sigs.append({'n': 'VWAP', 'v': f"₹{vw:.2f}", 't': t, 'b': price > vw, 'cat': 'Volume'})

    sk = g('SK')
    if not np.isnan(sk):
        if sk < 15: bs += 14; t = "Deep Oversold — Strong BUY"
        elif sk < 25: bs += 9; t = "Oversold — BUY"
        elif sk > 85: ss += 14; t = "Deep Overbought — Strong SELL"
        elif sk > 75: ss += 9; t = "Overbought — SELL"
        else: bs += 3; ss += 3; t = "Neutral zone"
        sigs.append({'n': 'Stochastic', 'v': f"{sk:.1f}", 't': t, 'b': sk < 50, 'cat': 'Momentum'})

    av = g('ADX'); ap = g('DMP'); an = g('DMN')
    if not np.isnan(av) and not np.isnan(ap) and not np.isnan(an):
        strength = "💪 Strong" if av > 30 else "📊 Moderate" if av > 20 else "💤 Weak"
        if ap > an: bs += 10; t = f"Bullish trend — {strength} ADX:{av:.0f}"
        else: ss += 10; t = f"Bearish trend — {strength} ADX:{av:.0f}"
        sigs.append({'n': 'ADX', 'v': f"{av:.1f}", 't': t, 'b': ap > an, 'cat': 'Trend'})

    cci = g('CCI')
    if not np.isnan(cci):
        if cci < -200: bs += 14; t = "Extreme Oversold — Strong BUY"
        elif cci < -100: bs += 8; t = "Oversold — BUY"
        elif cci > 200: ss += 14; t = "Extreme Overbought — Strong SELL"
        elif cci > 100: ss += 8; t = "Overbought — SELL"
        else: bs += 2; ss += 2; t = "Neutral"
        sigs.append({'n': 'CCI', 'v': f"{cci:.0f}", 't': t, 'b': cci < 0, 'cat': 'Momentum'})

    wr = g('WR')
    if not np.isnan(wr):
        if wr < -85: bs += 12; t = "Extremely Oversold — BUY"
        elif wr < -70: bs += 7; t = "Oversold — BUY"
        elif wr > -10: ss += 12; t = "Extremely Overbought — SELL"
        elif wr > -25: ss += 7; t = "Overbought — SELL"
        else: bs += 3; ss += 3; t = "Neutral"
        sigs.append({'n': 'Williams %R', 'v': f"{wr:.1f}", 't': t, 'b': wr < -50, 'cat': 'Momentum'})

    mf = g('MFI')
    if not np.isnan(mf):
        if mf < 20: bs += 12; t = "Strong money inflow — BUY"
        elif mf > 80: ss += 12; t = "Strong money outflow — SELL"
        else: bs += 3; ss += 3; t = "Neutral money flow"
        sigs.append({'n': 'MFI', 'v': f"{mf:.1f}", 't': t, 'b': mf < 50, 'cat': 'Volume'})

    roc = g('ROC')
    if not np.isnan(roc):
        if roc > 5: bs += 8; t = f"Strong positive momentum +{roc:.1f}%"
        elif roc > 0: bs += 4; t = f"Positive momentum +{roc:.1f}%"
        elif roc < -5: ss += 8; t = f"Strong negative momentum {roc:.1f}%"
        else: ss += 4; t = f"Negative momentum {roc:.1f}%"
        sigs.append({'n': 'ROC', 'v': f"{roc:.1f}%", 't': t, 'b': roc > 0, 'cat': 'Momentum'})

    tot2 = bs + ss
    bp = bs / tot2 * 100 if tot2 > 0 else 50
    sp = ss / tot2 * 100 if tot2 > 0 else 50
    conf = round(max(bp, sp), 1)
    if bp >= 80: sig = "STRONG BUY"; sc = "#15803d"; bg = "sig-buy"
    elif bp >= 68: sig = "BUY"; sc = "#16a34a"; bg = "sig-buy"
    elif bp >= 56: sig = "LEAN BUY"; sc = "#2563eb"; bg = "sig-buy"
    elif sp >= 80: sig = "STRONG SELL"; sc = "#b91c1c"; bg = "sig-sell"
    elif sp >= 68: sig = "SELL"; sc = "#dc2626"; bg = "sig-sell"
    elif sp >= 56: sig = "LEAN SELL"; sc = "#ea580c"; bg = "sig-sell"
    else: sig = "NEUTRAL"; sc = "#b45309"; bg = "sig-wait"

    atr_v = safe(lat.get('ATR'), price * 0.015)
    if atr_v <= 0: atr_v = price * 0.015
    std20 = safe(lat.get('STD20'), price * 0.01)
    vol = float(lat['Volume'])
    avgv = float(df['Volume'].rolling(20).mean().iloc[-1]) if n >= 20 else vol
    vr = round(vol / avgv, 2) if avgv > 0 else 1.0
    hi52 = float(df['High'].max()); lo52 = float(df['Low'].min())
    rng52 = hi52 - lo52
    pos52 = round((price - lo52) / rng52 * 100, 1) if rng52 > 0 else 50

    trend = 'SIDEWAYS'
    if n >= 20:
        r2 = df.tail(20); H = r2['High'].values; L = r2['Low'].values
        sh = []; sl2 = []
        for i in range(2, len(H) - 2):
            if H[i] > max(H[i-1], H[i-2], H[i+1], H[i+2]): sh.append(float(H[i]))
            if L[i] < min(L[i-1], L[i-2], L[i+1], L[i+2]): sl2.append(float(L[i]))
        hh = len(sh) >= 2 and sh[-1] > sh[-2]; lh = len(sh) >= 2 and sh[-1] < sh[-2]
        hl = len(sl2) >= 2 and sl2[-1] > sl2[-2]; ll = len(sl2) >= 2 and sl2[-1] < sl2[-2]
        if hh and hl: trend = 'UPTREND'
        elif lh and ll: trend = 'DOWNTREND'
        elif hh and ll: trend = 'VOLATILE'

    return {
        'df': df, 'price': price, 'atr': atr_v, 'std20': std20,
        'sig': sig, 'sc': sc, 'bg': bg, 'bp': round(bp, 1), 'sp': round(sp, 1), 'conf': conf,
        'sigs': sigs, 'n_sigs': len(sigs), 'vr': vr, 'vol': vol, 'avgv': avgv,
        'trend': trend, 'pos52': pos52, 'hi52': hi52, 'lo52': lo52, 'lat': lat, 'prev': prev,
        'rsi': rsi, 'vw': vw, 'stv': stv, 'std_': std_,
        'bbu': bu, 'bbl': bl, 'bbw': bw, 'bbp': bp_val,
        'adx': av, 'cci': cci, 'wr': wr, 'mfi': mf,
    }


def run_analysis(sym, iv, per):
    try:
        df = yf.Ticker(sym).history(period=per, interval=iv)
        if df is None or df.empty:
            return None
        df = df.dropna()
    except Exception:
        return None
    return compute_signals(df)


# ============================================================
# PREMIUM CHART — 5 PANELS
# ============================================================
def make_chart(df, name, plan, sr):
    C = {'bg': '#ffffff', 'plot': '#fafbff', 'grid': '#e0e7ff', 'up': '#16a34a', 'dn': '#dc2626', 'text': '#1a1f36'}
    fig = make_subplots(rows=5, cols=1, shared_xaxes=True, vertical_spacing=0.018,
                        row_heights=[0.42, 0.15, 0.15, 0.13, 0.15],
                        subplot_titles=[f"📈 {name}  |  Candles + EMAs + BB + Supertrend",
                                        "MACD", "RSI + Stochastic", "ADX + MFI", "Volume + OBV"])
    fig.update_layout(paper_bgcolor=C['bg'], plot_bgcolor=C['plot'], height=1000, showlegend=True,
                      xaxis_rangeslider_visible=False, margin=dict(l=60, r=180, t=55, b=30),
                      font=dict(size=11, color=C['text'], family='Inter'),
                      legend=dict(orientation='h', y=1.03, x=1, xanchor='right',
                                  bgcolor='rgba(255,255,255,0.95)', bordercolor='#e0e7ff', borderwidth=1))
    fig.add_trace(go.Candlestick(x=df.index, open=df['Open'], high=df['High'], low=df['Low'], close=df['Close'],
                                 name='Price', increasing=dict(line=dict(color=C['up'], width=1), fillcolor='rgba(22,163,74,0.8)'),
                                 decreasing=dict(line=dict(color=C['dn'], width=1), fillcolor='rgba(220,38,38,0.8)')), row=1, col=1)
    for c, clr, lbl, w in [('E9', '#f97316', 'EMA9', 1.4), ('E20', '#eab308', 'EMA20', 1.8), ('E50', '#6366f1', 'EMA50', 2.0)]:
        if c in df.columns:
            s = df[c].dropna()
            if len(s): fig.add_trace(go.Scatter(x=s.index, y=s, name=lbl, line=dict(color=clr, width=w)), row=1, col=1)
    if 'VWAP' in df.columns:
        s = df['VWAP'].dropna()
        if len(s): fig.add_trace(go.Scatter(x=s.index, y=s, name='VWAP', line=dict(color='#a855f7', width=1.8, dash='dot')), row=1, col=1)
    if 'BBU' in df.columns and 'BBL' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['BBU'], name='BB Upper', line=dict(color='rgba(220,38,38,0.4)', width=1, dash='dot')), row=1, col=1)
        fig.add_trace(go.Scatter(x=df.index, y=df['BBL'], name='BB Lower', line=dict(color='rgba(22,163,74,0.4)', width=1, dash='dot'),
                                 fill='tonexty', fillcolor='rgba(99,102,241,0.04)'), row=1, col=1)
    if plan:
        for y, clr, txt, dash, lw in [
            (plan['sl'], '#ef4444', f"🛑 SL ₹{plan['sl']:.2f}", 'dash', 2.0),
            (plan['t1'], '#22c55e', f"T1 ₹{plan['t1']:.2f}", 'dot', 1.5),
            (plan['t2'], '#16a34a', f"T2 ₹{plan['t2']:.2f}", 'dash', 2.0),
            (plan['t3'], '#a855f7', f"T3 ₹{plan['t3']:.2f}", 'dashdot', 1.5),
            (plan['buy_at'], '#3b82f6', f"📌 BUY ₹{plan['buy_at']:.2f}", 'solid', 2.2)]:
            fig.add_hline(y=y, line_dash=dash, line_color=clr, line_width=lw, annotation_text=txt,
                          annotation_position='right', annotation_font=dict(color=clr, size=10), row=1, col=1)
    if sr and sr.get('levels'):
        for lvl in sr['levels']:
            if abs(lvl['dist']) < 5:
                clr = 'rgba(220,38,38,0.2)' if lvl['type'] == 'R' else 'rgba(22,163,74,0.2)' if lvl['type'] == 'S' else 'rgba(245,158,11,0.2)'
                fig.add_hline(y=lvl['price'], line_dash='dot', line_color=clr, line_width=0.8, row=1, col=1)
    if 'ST' in df.columns:
        for dv, clr, sm in [(1, '#16a34a', 'triangle-up'), (-1, '#dc2626', 'triangle-down')]:
            pts = df[df['STD'] == dv]['ST'].dropna()
            if len(pts): fig.add_trace(go.Scatter(x=pts.index, y=pts, mode='markers', name=f"ST {'Buy' if dv==1 else 'Sell'}",
                                                  marker=dict(color=clr, size=6, symbol=sm)), row=1, col=1)
    if 'MACD' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['MACD'], name='MACD', line=dict(color='#2563eb', width=1.6)), row=2, col=1)
        fig.add_trace(go.Scatter(x=df.index, y=df['MS'], name='Signal', line=dict(color='#f97316', width=1.6)), row=2, col=1)
        hc = ['#16a34a' if v >= 0 else '#dc2626' for v in df['MH'].fillna(0)]
        fig.add_trace(go.Bar(x=df.index, y=df['MH'], name='Histogram', marker_color=hc, opacity=0.75), row=2, col=1)
        fig.add_hline(y=0, line_dash='dash', line_color='#cbd5e1', row=2, col=1)
    if 'RSI' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['RSI'].fillna(50), name='RSI', line=dict(color='#8b5cf6', width=2.0)), row=3, col=1)
    if 'SK' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['SK'].fillna(50), name='Stoch %K', line=dict(color='#f97316', width=1.3, dash='dot')), row=3, col=1)
    for y, clr in [(70, 'rgba(220,38,38,0.4)'), (30, 'rgba(22,163,74,0.4)'), (50, 'rgba(0,0,0,0.1)')]:
        fig.add_hline(y=y, line_dash='dash', line_color=clr, line_width=0.8, row=3, col=1)
    if 'ADX' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['ADX'].fillna(0), name='ADX', line=dict(color='#0ea5e9', width=1.8)), row=4, col=1)
        fig.add_trace(go.Scatter(x=df.index, y=df['DMP'].fillna(0), name='+DI', line=dict(color='#16a34a', width=1.2, dash='dot')), row=4, col=1)
        fig.add_trace(go.Scatter(x=df.index, y=df['DMN'].fillna(0), name='-DI', line=dict(color='#dc2626', width=1.2, dash='dot')), row=4, col=1)
    if 'MFI' in df.columns:
        fig.add_trace(go.Scatter(x=df.index, y=df['MFI'].fillna(50), name='MFI', line=dict(color='#a855f7', width=1.3, dash='dashdot')), row=4, col=1)
    vc = ['#16a34a' if c >= o else '#dc2626' for c, o in zip(df['Close'], df['Open'])]
    fig.add_trace(go.Bar(x=df.index, y=df['Volume'], name='Volume', marker_color=vc, opacity=0.8), row=5, col=1)
    if 'OBV' in df.columns:
        obv = df['OBV'].dropna()
        if len(obv): fig.add_trace(go.Scatter(x=obv.index, y=obv, name='OBV', line=dict(color='#0ea5e9', width=1.5)), row=5, col=1)
    for i in range(1, 6):
        fig.update_xaxes(gridcolor=C['grid'], row=i, col=1, showgrid=True)
        fig.update_yaxes(gridcolor=C['grid'], row=i, col=1, showgrid=True)
    for i in range(1, 5):
        fig.update_xaxes(showticklabels=False, row=i, col=1)
    return fig


def pivot_table_html(levels_dict, price):
    rows = sorted(levels_dict.items(), key=lambda x: x[1], reverse=True)
    html = ""
    for lbl, lp in rows:
        dist = round((lp - price) / price * 100, 2) if price > 0 else 0
        is_cur = abs(dist) < 0.4
        if is_cur:
            css, clr, tag = "lv-cur", "#3b82f6", "📍 nearest to price"
        elif lbl.startswith('R'):
            css, clr, tag = "lv-r", "#ef4444", f"+{abs(dist):.1f}% above"
        elif lbl.startswith('S'):
            css, clr, tag = "lv-s", "#22c55e", f"{abs(dist):.1f}% below"
        else:
            css, clr, tag = "lv-p", "#f59e0b", f"{dist:+.1f}%"
        typ = "RESISTANCE" if lbl.startswith('R') else "SUPPORT" if lbl.startswith('S') else "PIVOT"
        html += (f"<div class='{css}'><div style='display:flex;justify-content:space-between;align-items:center;'>"
                 f"<div><span style='color:{clr};font-weight:700;font-size:13px;'>{lbl} — {typ}</span>"
                 f"<span style='color:#9ca3af;font-size:11px;margin-left:8px;'>{tag}</span></div>"
                 f"<div style='font-size:20px;font-weight:900;color:#1a1f36;'>₹{lp:,.2f}</div></div></div>")
    return html


# ============================================================
# ADVANCED ANALYTICS — ML · regime · volume profile · patterns
#   · risk metrics · probability  (all fail-safe)
# ============================================================
@st.cache_data(ttl=1800, show_spinner=False)
def ml_predict(sym, price):
    try:
        from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
        from sklearn.linear_model import LinearRegression
    except Exception:
        return {"error": "scikit-learn not installed"}
    d = _daily_for_ml(sym)
    if d is None:
        return None
    try:
        d = d.copy()
        d['ret1'] = d['Close'].pct_change()
        d['ret5'] = d['Close'].pct_change(5)
        d['rsi'] = ta.momentum.rsi(d['Close'], 14)
        d['volat'] = d['ret1'].rolling(10).std()
        d['mom'] = d['Close'] / d['Close'].shift(10) - 1
        d['volr'] = d['Volume'] / d['Volume'].rolling(20).mean()
        d['ma20'] = d['Close'] / ta.trend.sma_indicator(d['Close'], 20) - 1
        d['ma50'] = d['Close'] / ta.trend.sma_indicator(d['Close'], 50) - 1
        d['macd'] = ta.trend.MACD(d['Close']).macd_diff()
        feats = ['ret1', 'ret5', 'rsi', 'volat', 'mom', 'volr', 'ma20', 'ma50', 'macd']
        horizons = [1, 3, 5, 10]
        out = {}
        for h in horizons:
            d['tgt'] = d['Close'].shift(-h) / d['Close'] - 1
            data = d.dropna(subset=feats + ['tgt'])
            if len(data) < 80:
                continue
            X = data[feats].values; y = data['tgt'].values
            sp = int(len(X) * 0.8)
            Xtr, Xte, ytr, yte = X[:sp], X[sp:], y[:sp], y[sp:]
            mdls = {"RF": RandomForestRegressor(n_estimators=60, max_depth=6, random_state=1, n_jobs=-1),
                    "GB": GradientBoostingRegressor(n_estimators=60, max_depth=3, random_state=1),
                    "LR": LinearRegression()}
            last = d[feats].iloc[[-1]].values
            preds, accs, te_preds = {}, {}, {}
            for nm, mdl in mdls.items():
                mdl.fit(Xtr, ytr)
                pt = mdl.predict(Xte)
                te_preds[nm] = pt
                accs[nm] = float(np.mean(np.sign(pt) == np.sign(yte))) if len(yte) else 0.5
                preds[nm] = float(mdl.predict(last)[0])
            wsum = sum(max(a, 0.01) for a in accs.values())
            ens = sum(preds[k] * max(accs[k], 0.01) for k in preds) / wsum
            ens_te = sum(te_preds[k] * max(accs[k], 0.01) for k in te_preds) / wsum
            resid = float(np.std(yte - ens_te)) if len(yte) else 0.02
            acc = float(np.mean(np.sign(ens_te) == np.sign(yte))) if len(yte) else 0.5
            pred_price = price * (1 + ens)
            lo = price * (1 + ens - 1.96 * resid); hi = price * (1 + ens + 1.96 * resid)
            out[h] = {"ret": ens * 100, "price": pred_price, "lo": lo, "hi": hi,
                      "acc": acc * 100, "dir": "UP" if ens > 0 else "DOWN",
                      "model_acc": {k: round(v * 100, 1) for k, v in accs.items()}}
        return out or None
    except Exception:
        return None


@st.cache_data(ttl=1800, show_spinner=False)
def ema200_filter(sym, price):
    d = _daily_for_ml(sym)
    if d is None or len(d) < 60:
        return None
    try:
        c = d['Close']
        e200 = ta.trend.ema_indicator(c, 200) if len(c) >= 200 else ta.trend.ema_indicator(c, min(len(c)-1, 150))
        e50 = ta.trend.ema_indicator(c, 50)
        n200 = float(e200.iloc[-1]); n50 = float(e50.iloc[-1])
        approx = len(c) < 200
        dist = (price - n200) / n200 * 100 if n200 else 0
        above = price > n200
        golden = n50 > n200
        near = abs(dist) < 1.2
        slope_up = float(e200.iloc[-1]) > float(e200.iloc[-6]) if len(e200) >= 6 else True
        if above and golden and slope_up:
            verdict, vc, ok = "BULL STRUCTURE — longs favoured", "#15803d", True
        elif above and not golden:
            verdict, vc, ok = "Above 200 EMA but no golden cross — cautious longs", "#2563eb", True
        elif above:
            verdict, vc, ok = "Above 200 EMA — longs allowed", "#16a34a", True
        else:
            verdict, vc, ok = "BELOW 200 EMA — bear structure, AVOID longs", "#dc2626", False
        return {"ema200": round(n200, 2), "ema50": round(n50, 2), "dist": round(dist, 2),
                "above": above, "golden": golden, "near": near, "slope_up": slope_up,
                "verdict": verdict, "vc": vc, "ok": ok, "approx": approx}
    except Exception:
        return None


def entry_timing(res, df, sr, ema2=None):
    price = res['price']
    g = lambda k: res.get(k, np.nan)
    warns = []; goods = []; score = 0
    rsi = g('rsi')
    if not np.isnan(rsi):
        if rsi > 75: warns.append("RSI very overbought (>75)"); score -= 2
        elif rsi > 68: warns.append("RSI overbought"); score -= 1
        elif rsi < 40: goods.append("RSI has room to rise"); score += 1
    d = df.tail(30)
    hi = float(d['High'].max()); lo = float(d['Low'].min()); rng = max(hi - lo, 1e-9)
    pos = (price - lo) / rng
    if pos > 0.85: warns.append(f"Price at top of recent range ({pos:.0%}) — late"); score -= 2
    elif pos < 0.35: goods.append("Price in lower part of range"); score += 1
    vw = g('vw')
    if not np.isnan(vw) and vw > 0:
        vd = (price - vw) / vw * 100
        if vd > 2.5: warns.append(f"Stretched {vd:.1f}% above VWAP"); score -= 2
        elif vd > 1.2: warns.append(f"{vd:.1f}% above VWAP"); score -= 1
    bbp = g('bbp')
    if not np.isnan(bbp):
        if bbp > 0.95: warns.append("At/above upper Bollinger band"); score -= 2
        elif bbp < 0.15: goods.append("Near lower Bollinger band"); score += 1
    if sr and sr.get('nr') and abs(sr['nr']['dist']) < 1.0:
        warns.append(f"Resistance ₹{sr['nr']['price']:.2f} just above"); score -= 1
    if len(df) >= 15:
        rl = float(df['Low'].tail(15).min())
        run = (price - rl) / rl * 100 if rl > 0 else 0
        if run > 6: warns.append(f"Already ran +{run:.1f}% off recent low — move may be done"); score -= 2
        elif run > 3.5: warns.append(f"Up +{run:.1f}% recently"); score -= 1
    if 'MH' in df.columns:
        mh = df['MH'].dropna()
        if len(mh) >= 3 and mh.iloc[-1] > 0 and mh.iloc[-1] < mh.iloc[-2] < mh.iloc[-3]:
            warns.append("MACD momentum fading (histogram shrinking)"); score -= 1
    try:
        if 'RSI' in df.columns and len(df) >= 20:
            half = df.tail(20)
            a, b = half.iloc[:10], half.iloc[10:]
            if b['High'].max() > a['High'].max() and b['RSI'].max() < a['RSI'].max() - 2:
                warns.append("Bearish divergence (price up, RSI down)"); score -= 2
    except Exception:
        pass
    tr = res['trend']
    if ema2 and not ema2.get('above', True):
        warns.append("Below daily 200 EMA (bear structure)"); score -= 3
    elif ema2 and ema2.get('above') and ema2.get('golden'):
        goods.append("Above 200 EMA + golden cross"); score += 1
    if tr == 'UPTREND':
        if score <= -3:
            stage, q, tone, msg = ("EXTENDED / LATE", "POOR", "warn",
                "Uptrend — but price is overextended. You'd be buying near the top. WAIT for a pullback toward support before entering.")
        elif score <= -1:
            stage, q, tone, msg = ("MID — CAUTION", "CAUTION", "warn",
                "Uptrend is intact but this isn't a fresh entry. Prefer a dip toward support, or size smaller.")
        else:
            stage, q, tone, msg = ("EARLY / FRESH", "GOOD", "buy",
                "Uptrend with room to move — a reasonable entry, ideally near support.")
    elif tr == 'DOWNTREND':
        stage, q, tone, msg = ("DOWNTREND", "AVOID", "sell",
            "Downtrend — don't buy. Wait for a clear bottom and reversal.")
    else:
        if score >= 1:
            stage, q, tone, msg = ("BASING", "WATCH", "wait",
                "No clear trend, but not overextended — watch for a breakout.")
        else:
            stage, q, tone, msg = ("CHOPPY", "WAIT", "wait",
                "No clear trend and stretched — wait for a cleaner setup.")
    return {"stage": stage, "quality": q, "tone": tone, "msg": msg,
            "warns": warns, "goods": goods, "score": score}


# ============================================================
# NEWS + SENTIMENT
# ============================================================
_POS = ["surge", "jump", "gain", "rise", "profit", "beat", "record", "high", "wins",
        "win ", "order", "bag", "bags", "approval", "approved", "upgrade", "buy rating",
        "growth", "rally", "strong", "boost", "acquire", "acquisition", "expansion",
        "dividend", "bonus", "outperform", "target raised", "deal", "contract", "surges", "jumps"]
_NEG = ["fall", "drop", "slump", "loss", "miss", "cut", "fraud", "probe", "penalty",
        "warn", "downgrade", "sell rating", "crash", "decline", "weak", "debt", "default",
        "resign", "layoff", "raid", "ban", "scam", "lawsuit", "recall", "plunge", "slips",
        "falls", "drops", "tumble", "concern", "block", "stake sale", "selloff"]


@st.cache_data(ttl=900, show_spinner=False)
def get_stock_news(sym, limit=7):
    items = []
    try:
        items = yf.Ticker(sym).news or []
    except Exception:
        items = []
    out = []
    for it in items[: limit * 2]:
        try:
            c = it.get("content", it)
            title = c.get("title") or it.get("title")
            if not title:
                continue
            prov = c.get("provider") or {}
            pub = (prov.get("displayName") if isinstance(prov, dict) else None) or it.get("publisher") or "—"
            url = ""
            cu = c.get("canonicalUrl") or c.get("clickThroughUrl")
            if isinstance(cu, dict):
                url = cu.get("url", "")
            url = url or it.get("link", "")
            when = c.get("pubDate") or it.get("providerPublishTime", "")
            t = title.lower()
            pos = sum(1 for w in _POS if w in t)
            neg = sum(1 for w in _NEG if w in t)
            sent = "pos" if pos > neg else "neg" if neg > pos else "neu"
            out.append({"title": title, "pub": pub, "url": url, "when": str(when), "sent": sent})
            if len(out) >= limit:
                break
        except Exception:
            continue
    p = sum(1 for x in out if x["sent"] == "pos")
    n = sum(1 for x in out if x["sent"] == "neg")
    if not out:
        verdict, vc = "No recent news found", "#6b7280"
    elif p > n:
        verdict, vc = "NEWS LEANS POSITIVE", "#16a34a"
    elif n > p:
        verdict, vc = "NEWS LEANS NEGATIVE — be careful", "#dc2626"
    else:
        verdict, vc = "NEWS MIXED / NEUTRAL", "#b45309"
    return {"items": out, "pos": p, "neg": n, "verdict": verdict, "vc": vc}


@st.cache_data(ttl=600, show_spinner=False)
def fetch_news(sym, limit=8):
    """(from your v13 — kept as-is) Recent headlines, keyword-tagged."""
    out = []
    try:
        items = yf.Ticker(sym).news or []
    except Exception:
        items = []
    POS = ["surge", "jump", "gain", "rise", "profit", "beat", "high", "record", "wins",
           "win", "order", "bag", "bags", "approval", "approved", "upgrade", "buy", "rally",
           "soar", "growth", "strong", "acquire", "acquisition", "deal", "dividend", "bonus"]
    NEG = ["fall", "drop", "slump", "loss", "miss", "cut", "low", "fraud", "probe", "penalty",
           "warn", "downgrade", "sell", "crash", "decline", "weak", "block", "ban", "raid",
           "resign", "default", "debt", "lawsuit", "fine", "recall", "delay"]
    for it in items[:limit * 2]:
        try:
            c = it.get("content", it)
            title = c.get("title") or it.get("title")
            if not title:
                continue
            prov = c.get("provider") or {}
            pub = (prov.get("displayName") if isinstance(prov, dict) else None) or it.get("publisher") or "—"
            cu = c.get("canonicalUrl") or c.get("clickThroughUrl") or {}
            url = cu.get("url", "") if isinstance(cu, dict) else (it.get("link", "") or "")
            when = c.get("pubDate") or it.get("providerPublishTime", "")
            t = title.lower()
            pos = sum(w in t for w in POS); neg = sum(w in t for w in NEG)
            sent = "pos" if pos > neg else "neg" if neg > pos else "neu"
            out.append({"title": title, "pub": pub, "url": url, "when": str(when), "sent": sent})
            if len(out) >= limit:
                break
        except Exception:
            continue
    return out


def summarize_news_sentiment(news):
    """(from your v13 — kept as-is)"""
    if not news:
        return {"label": "No recent news found", "tone": "neu", "pos": 0, "neg": 0, "neu": 0}
    pos = sum(1 for n in news if n["sent"] == "pos")
    neg = sum(1 for n in news if n["sent"] == "neg")
    neu = len(news) - pos - neg
    if pos > neg and pos >= 2:
        label, tone = "📈 Mostly POSITIVE news", "pos"
    elif neg > pos and neg >= 2:
        label, tone = "📉 Mostly NEGATIVE news — be careful", "neg"
    elif pos > neg:
        label, tone = "Slightly positive", "pos"
    elif neg > pos:
        label, tone = "Slightly negative", "neg"
    else:
        label, tone = "Mixed / neutral headlines", "neu"
    return {"label": label, "tone": tone, "pos": pos, "neg": neg, "neu": neu}


# ============================================================
# TRADE JOURNAL — save today's plan, verify it next day
# ============================================================
import json as _json
import os as _os
# ── 🧑‍🤝‍🧑 MULTI-USER MEMORY ──────────────────────────────────────────────
#   Everyone who opens the app gets their OWN files, so memories never mix:
#   your key lives in the link (?u=yourname). Default key = "main".
#   Set / change it in the 🌙 EOD Review tab ("Your memory name").
def _ukey():
    try:
        u = (st.query_params.get("u") or "").strip().lower()
        u = "".join(ch for ch in u if ch.isalnum() or ch in "_-")[:16]
        return u or "main"
    except Exception:
        return "main"


def journal_file():
    return f"trade_journal_{_ukey()}.json"


def snap_file():
    return f"board_snapshots_{_ukey()}.json"


def eod_file():
    return f"eod_results_{_ukey()}.json"


# ── ♾️ RUNTIME PERSISTENCE — analysis keeps running across page refreshes.
#    The run state lives in a FILE (per person), so a browser refresh,
#    phone screen-lock or reconnect AUTO-RESUMES the analysis. Only the
#    ⏹ Stop / 🧹 Clear buttons end it.
def runtime_file():
    return f"runtime_{_ukey()}.json"


def _rt_sanitize(o):
    """Make everything JSON-safe (numpy floats, timestamps, etc.)."""
    if isinstance(o, dict):
        return {str(k): _rt_sanitize(v) for k, v in o.items()}
    if isinstance(o, (list, tuple)):
        return [_rt_sanitize(v) for v in o]
    if isinstance(o, (bool, str)) or o is None:
        return o
    try:
        if isinstance(o, (int, float)):
            return o
        f = float(o)
        return int(f) if f.is_integer() and abs(f) < 1e15 else f
    except Exception:
        return str(o)


def rt_load():
    try:
        if _os.path.exists(runtime_file()):
            with open(runtime_file(), "r", encoding="utf-8") as f:
                d = _json.load(f)
            return d if isinstance(d, dict) else {}
    except Exception:
        pass
    return {}


def rt_save(engine, **kw):
    try:
        d = rt_load()
        d[engine] = {**d.get(engine, {}), **kw}
        with open(runtime_file(), "w", encoding="utf-8") as f:
            _json.dump(_rt_sanitize(d), f)
        return True
    except Exception:
        return False


def rt_clear(engine):
    try:
        d = rt_load()
        d.pop(engine, None)
        with open(runtime_file(), "w", encoding="utf-8") as f:
            _json.dump(d, f)
    except Exception:
        pass

_JOURNAL = "trade_journal.json"


def journal_load():
    fn = journal_file()
    try:
        if _os.path.exists(fn):
            with open(fn, "r", encoding="utf-8") as f:
                return _json.load(f)
        if _ukey() == "main" and _os.path.exists(_JOURNAL):   # adopt old file
            with open(_JOURNAL, "r", encoding="utf-8") as f:
                return _json.load(f)
    except Exception:
        pass
    return []


def journal_save(entry):
    try:
        data = journal_load()
        entry["id"] = f"{entry['sym']}_{entry['saved']}"
        data = [d for d in data if d.get("id") != entry["id"]]
        data.append(entry)
        with open(journal_file(), "w", encoding="utf-8") as f:
            _json.dump(data, f, indent=2)
        return True
    except Exception:
        return False


def journal_delete(entry_id):
    try:
        data = [d for d in journal_load() if d.get("id") != entry_id]
        with open(journal_file(), "w", encoding="utf-8") as f:
            _json.dump(data, f, indent=2)
        return True
    except Exception:
        return False


def journal_verify(entry):
    try:
        saved_date = entry["saved"][:10]
        d = yf.Ticker(entry["sym"]).history(period="1mo", interval="1d").dropna()
        if d is None or len(d) == 0:
            return {"status": "NO DATA", "color": "#6b7280", "detail": "Couldn't fetch prices."}
        after = d[d.index.strftime("%Y-%m-%d") > saved_date]
        if len(after) == 0:
            return {"status": "TOO SOON", "color": "#6b7280",
                    "detail": "No trading day has completed since you saved this — check tomorrow."}
        hi = float(after["High"].max()); lo = float(after["Low"].min())
        last = float(after["Close"].iloc[-1])
        buy, sl = entry["buy_at"], entry["sl"]
        t1, t2, t3 = entry["t1"], entry["t2"], entry["t3"]
        hit_sl = lo <= sl
        hit_t1 = hi >= t1; hit_t2 = hi >= t2; hit_t3 = hi >= t3
        moved = (last - entry["price"]) / entry["price"] * 100
        want_up = "UP" in entry.get("trend", "") or "BUY" in entry.get("signal", "")
        trend_right = (moved > 0) if want_up else (moved < 0)
        if entry.get("sit") == "BUY_DIP" and lo > buy:
            status, color = "NO-FILL (never dipped to buy)", "#6b7280"
            detail = f"Price never fell to ₹{buy:.2f}; low since was ₹{lo:.2f}. No trade taken — correct to wait."
        elif hit_t3:
            status, color = "WIN → T3 🎯", "#15803d"; detail = f"Hit all targets (high ₹{hi:.2f})."
        elif hit_t2:
            status, color = "WIN → T2 ✅", "#16a34a"; detail = f"Reached T2 ₹{t2:.2f} (high ₹{hi:.2f})."
        elif hit_t1:
            status, color = "WIN → T1 ✅", "#16a34a"; detail = f"Reached T1 ₹{t1:.2f} (high ₹{hi:.2f})."
        elif hit_sl:
            status, color = "LOSS → hit stop ❌", "#dc2626"; detail = f"Fell to stop ₹{sl:.2f} (low ₹{lo:.2f})."
        else:
            status, color = "OPEN — no target/stop hit yet", "#b45309"
            detail = f"Since save: high ₹{hi:.2f}, low ₹{lo:.2f}, now ₹{last:.2f}."
        return {"status": status, "color": color, "detail": detail,
                "moved": round(moved, 2), "trend_right": trend_right,
                "hi": hi, "lo": lo, "last": last}
    except Exception as e:
        return {"status": "ERROR", "color": "#6b7280", "detail": f"{type(e).__name__}"}


# ============================================================
# 🧠 BOARD MEMORY + 🌙 EOD (END-OF-DAY) VERIFICATION
#   The dashboard auto-saves WHAT IT SHOWED (snapshots). The EOD
#   tab later fetches what ACTUALLY happened and scores every
#   call — "how many stocks went as per our calculation?"
# ============================================================
SNAP_FILE = "board_snapshots.json"
EOD_FILE = "eod_results.json"
_SNAP_FIELDS = ("sym", "name", "price", "chg", "sig", "conf", "bp", "sp", "dtr",
                "above200", "score", "buy_at", "sl", "t1", "t2", "rr", "sit")


def snaps_load():
    fn = snap_file()
    try:
        if _os.path.exists(fn):
            with open(fn, "r", encoding="utf-8") as f:
                d = _json.load(f)
            return d if isinstance(d, list) else []
        if _ukey() == "main" and _os.path.exists(SNAP_FILE):   # adopt old file
            with open(SNAP_FILE, "r", encoding="utf-8") as f:
                d = _json.load(f)
            return d if isinstance(d, list) else []
    except Exception:
        pass
    return []


def snaps_save(rows):
    """ONE snapshot per day = that day's calculation (today's board). While the
    terminal runs, today's snapshot is refreshed — no intraday pile-up. The EOD
    tab verifies it the NEXT day. Keeps the last 40 days."""
    try:
        now = now_ist()
        snap = {"id": now.strftime("%Y-%m-%d"),
                "saved": now.strftime("%Y-%m-%d %H:%M"),
                "rows": [{k: r.get(k) for k in _SNAP_FIELDS} for r in rows.values()]}
        data = [s for s in snaps_load() if s.get("id") != snap["id"]]
        data.append(snap)
        data.sort(key=lambda s: s.get("saved", ""))
        data = data[-40:]
        with open(snap_file(), "w", encoding="utf-8") as f:
            _json.dump(data, f)
        return True
    except Exception:
        return False


def eod_load():
    fn = eod_file()
    try:
        if _os.path.exists(fn):
            with open(fn, "r", encoding="utf-8") as f:
                d = _json.load(f)
            return d if isinstance(d, list) else []
        if _ukey() == "main" and _os.path.exists(EOD_FILE):    # adopt old file
            with open(EOD_FILE, "r", encoding="utf-8") as f:
                d = _json.load(f)
            return d if isinstance(d, list) else []
    except Exception:
        pass
    return []


def eod_save(entry):
    try:
        data = [d for d in eod_load() if d.get("id") != entry.get("id")]
        data.append(entry)
        data.sort(key=lambda d: d.get("date", ""))
        data = data[-100:]
        with open(eod_file(), "w", encoding="utf-8") as f:
            _json.dump(data, f, indent=2)
        return True
    except Exception:
        return False


def verify_snapshot(snap):
    """AFTER: fetch what prices ACTUALLY did after a snapshot and score
    every call. Returns per-stock results + a summary."""
    syms = [r["sym"] for r in snap["rows"] if r.get("sym")]
    # market probe: latest candle Yahoo currently has (shows if data is available at all)
    last_candle = None
    try:
        _p = fetch_chunk(("RELIANCE.NS", "TCS.NS"), "1d", "6mo").get("RELIANCE.NS")
        if _p is not None and len(_p):
            last_candle = str(_p.index[-1].date())
    except Exception:
        pass
    # fetch with RETRIES — Yahoo throttles cloud servers; a single pass can fail completely
    got = {}
    for _attempt in range(3):
        todo = [s for s in syms if got.get(s) is None or len(got[s]) == 0]
        if not todo:
            break
        for i in range(0, len(todo), 50):
            try:
                got.update(fetch_chunk(tuple(todo[i:i + 50]), "1d", "6mo"))
            except Exception:
                pass
        if _attempt < 2:
            time.sleep(4)
    try:
        sdate = datetime.strptime(snap["saved"][:10], "%Y-%m-%d").date()
    except Exception:
        sdate = datetime.now().date()
    results = []
    for r in snap["rows"]:
        row = dict(r)
        row.update({"outcome": "NO DATA", "moved": 0.0, "correct": None,
                    "hi": None, "lo": None, "last": None})
        df = got.get(r.get("sym"))
        if df is not None and len(df):
            after = df[[ix.date() > sdate for ix in df.index]]
            same_day = False
            if len(after) == 0:   # snapshot is from today → use today's candle
                after = df[[ix.date() == sdate for ix in df.index]]
                same_day = True
            if len(after):
                hi = float(after["High"].max()); lo = float(after["Low"].min())
                last = float(after["Close"].iloc[-1])
                row["hi"], row["lo"], row["last"] = round(hi, 2), round(lo, 2), round(last, 2)
                row["moved"] = round((last - r["price"]) / r["price"] * 100, 2) if r["price"] else 0.0
                buy_call = ("BUY" in (r.get("sig") or "")) or r.get("dtr") == "UPTREND"
                hit_t1 = hi >= (r.get("t1") or 1e18)
                hit_t2 = hi >= (r.get("t2") or 1e18)
                hit_sl = lo <= (r.get("sl") or -1e18)
                if buy_call:
                    if hit_t2:                       row["outcome"] = "WIN T2 ✅"
                    elif hit_t1 and not hit_sl:      row["outcome"] = "WIN T1 ✅"
                    elif hit_sl and not hit_t1:      row["outcome"] = "LOSS SL ❌"
                    elif hit_t1 and hit_sl:          row["outcome"] = "BOTH T1&SL ⚠️"
                    else:                            row["outcome"] = "NO MOVE ➖"
                    row["correct"] = (row["moved"] > 0) if row["moved"] != 0 else None
                elif r.get("dtr") == "DOWNTREND":
                    if row["moved"] < 0:
                        row["outcome"], row["correct"] = "FELL ✅", True
                    elif row["moved"] > 0:
                        row["outcome"], row["correct"] = "ROSE ❌", False
                    else:
                        row["outcome"], row["correct"] = "FLAT ➖", None
                else:
                    row["outcome"] = "—"
                if same_day:
                    row["outcome"] += " (same-day)"
        results.append(row)
    ups = [x for x in results if x.get("dtr") == "UPTREND"]
    dns = [x for x in results if x.get("dtr") == "DOWNTREND"]
    bu = [x for x in results if "BUY" in (x.get("sig") or "")]
    up_ok = [x for x in ups if x.get("correct") is True]
    t1w = [x for x in bu if str(x["outcome"]).startswith("WIN")]
    sls = [x for x in bu if str(x["outcome"]).startswith("LOSS")]

    def _avg(lst):
        v = [x["moved"] for x in lst if isinstance(x.get("moved"), (int, float))]
        return round(sum(v) / len(v), 2) if v else 0.0

    nodata = [x for x in results if str(x.get("outcome", "")).startswith("NO DATA")]
    scored_up = [x for x in ups if x.get("correct") is not None]
    summary = {"snap": snap["saved"], "checked": len(results),
               "n_up": len(ups), "up_ok": len(up_ok),
               "up_acc": round(len(up_ok) / len(ups) * 100, 1) if scored_up else None,
               "n_scored_up": len(scored_up),
               "avg_up": _avg(ups), "avg_dn": _avg(dns),
               "n_buy": len(bu), "n_t1": len(t1w), "n_sl": len(sls),
               "buy_wr": round(len(t1w) / (len(t1w) + len(sls)) * 100, 1) if (t1w or sls) else None,
               "n_valid": len(results) - len(nodata), "n_nodata": len(nodata),
               "last_candle": last_candle}
    return results, summary


def detect_regime(res):
    adx = res.get('adx', np.nan); bbw = res.get('bbw', np.nan)
    tr = res['trend']
    trending = (not np.isnan(adx)) and adx > 25
    ranging = (not np.isnan(adx)) and adx < 20
    if trending and tr == 'UPTREND':
        reg, rc = "STRONG UPTREND", "#15803d"
    elif trending and tr == 'DOWNTREND':
        reg, rc = "STRONG DOWNTREND", "#b91c1c"
    elif ranging:
        reg, rc = "RANGING / SIDEWAYS", "#b45309"
    else:
        reg, rc = tr.title() if tr != 'SIDEWAYS' else "TRANSITIONING", "#2563eb"
    if not np.isnan(bbw):
        vol = "HIGH VOLATILITY" if bbw > 6 else "LOW VOLATILITY" if bbw < 2 else "MODERATE VOLATILITY"
    else:
        vol = "—"
    vc = "#dc2626" if "HIGH" in vol else "#16a34a" if "LOW" in vol else "#b45309"
    return {"regime": reg, "rc": rc, "vol": vol, "vc": vc, "adx": adx}


def volume_profile(df, bins=24):
    try:
        d = df.tail(120)
        lo, hi = float(d['Low'].min()), float(d['High'].max())
        if hi <= lo:
            return None
        edges = np.linspace(lo, hi, bins + 1)
        centers = (edges[:-1] + edges[1:]) / 2
        vol = np.zeros(bins)
        for _, r in d.iterrows():
            mid = (r['High'] + r['Low']) / 2
            idx = min(int((mid - lo) / (hi - lo) * bins), bins - 1)
            vol[idx] += r['Volume']
        poc = float(centers[int(np.argmax(vol))])
        order = np.argsort(vol)[::-1]
        total = vol.sum(); acc = 0; chosen = []
        for i in order:
            acc += vol[i]; chosen.append(i)
            if acc >= 0.70 * total:
                break
        va_lo = float(centers[min(chosen)]); va_hi = float(centers[max(chosen)])
        price = float(df['Close'].iloc[-1])
        if price > va_hi: pos = "ABOVE value area — extended up (mean-reversion risk)"
        elif price < va_lo: pos = "BELOW value area — extended down (bounce possible)"
        else: pos = "INSIDE value area — fairly priced"
        return {"poc": poc, "va_lo": va_lo, "va_hi": va_hi, "pos": pos, "price": price}
    except Exception:
        return None


def detect_patterns(df):
    try:
        d = df.tail(60)
        H = d['High'].values; L = d['Low'].values; C = d['Close'].values
        n = len(C)
        if n < 20:
            return []
        peaks, troughs = [], []
        for i in range(2, n - 2):
            if H[i] >= max(H[i-2:i]) and H[i] >= max(H[i+1:i+3]): peaks.append((i, H[i]))
            if L[i] <= min(L[i-2:i]) and L[i] <= min(L[i+1:i+3]): troughs.append((i, L[i]))
        out = []
        tol = np.mean(C) * 0.02
        if len(peaks) >= 2:
            (i1, h1), (i2, h2) = peaks[-2], peaks[-1]
            if abs(h1 - h2) < tol and i2 - i1 >= 3:
                out.append(("Double Top", "Bearish — watch for breakdown", "Medium"))
        if len(troughs) >= 2:
            (i1, l1), (i2, l2) = troughs[-2], troughs[-1]
            if abs(l1 - l2) < tol and i2 - i1 >= 3:
                out.append(("Double Bottom", "Bullish — watch for breakout", "Medium"))
        if len(peaks) >= 3:
            (_, a), (_, b), (_, c) = peaks[-3], peaks[-2], peaks[-1]
            if b > a and b > c and abs(a - c) < tol:
                out.append(("Head & Shoulders", "Bearish reversal", "High"))
        move = (C[-1] - C[-15]) / C[-15] * 100 if n >= 15 else 0
        rng = (np.max(H[-6:]) - np.min(L[-6:])) / C[-1] * 100
        if move > 6 and rng < 4:
            out.append(("Bull Flag", "Bullish continuation likely", "Medium"))
        elif move < -6 and rng < 4:
            out.append(("Bear Flag", "Bearish continuation likely", "Medium"))
        if len(peaks) >= 2 and len(troughs) >= 2:
            if peaks[-1][1] < peaks[-2][1] and troughs[-1][1] > troughs[-2][1]:
                out.append(("Triangle / Consolidation", "Breakout pending — trade the break", "Low"))
        return out
    except Exception:
        return []


@st.cache_data(ttl=1800, show_spinner=False)
def risk_metrics(sym):
    d = _daily_for_ml(sym)
    if d is None:
        return None
    try:
        ret = d['Close'].pct_change().dropna()
        if len(ret) < 60:
            return None
        ann = 252
        sharpe = float(ret.mean() / (ret.std() + 1e-9) * np.sqrt(ann))
        cum = (1 + ret).cumprod()
        dd = float(((cum - cum.cummax()) / cum.cummax()).min() * 100)
        ann_ret = float((cum.iloc[-1] ** (ann / len(ret)) - 1) * 100)
        calmar = float(ann_ret / abs(dd)) if dd != 0 else 0
        win = float((ret > 0).mean() * 100)
        var95 = float(np.percentile(ret, 5) * 100)
        avol = float(ret.std() * np.sqrt(ann) * 100)
        vrating = "HIGH" if avol > 40 else "LOW" if avol < 20 else "MODERATE"
        return {"sharpe": round(sharpe, 2), "maxdd": round(dd, 1), "calmar": round(calmar, 2),
                "win": round(win, 1), "var95": round(var95, 2), "avol": round(avol, 1),
                "vrating": vrating}
    except Exception:
        return None


def target_probability(price, plan, sym):
    d = _daily_for_ml(sym)
    try:
        from scipy.stats import norm
        have_scipy = True
    except Exception:
        have_scipy = False
    if d is None:
        return None
    try:
        ret = d['Close'].pct_change().dropna()
        sig = float(ret.std())
        if sig <= 0:
            return None
        H = 3
        sd = sig * np.sqrt(H)

        def prob_up(tp):
            move = (tp - price) / price
            z = move / (sd + 1e-9)
            if have_scipy:
                return float(1 - norm.cdf(z))
            return float(max(0.0, min(1.0, 0.5 - z * 0.2)))

        def rate(p):
            return ("High", "#16a34a") if p >= 0.55 else ("Moderate", "#b45309") if p >= 0.35 else ("Low", "#dc2626")

        res = {}
        for key, tp in [("T1", plan['t1']), ("T2", plan['t2']), ("T3", plan['t3'])]:
            p = prob_up(tp); r, c = rate(p)
            res[key] = {"p": round(p * 100, 1), "rating": r, "color": c, "price": tp}
        smove = (price - plan['sl']) / price
        zs = smove / (sd + 1e-9)
        psl = float(norm.cdf(-zs)) if have_scipy else float(max(0.0, min(1.0, 0.5 - zs * 0.2)))
        res["SL"] = {"p": round(psl * 100, 1)}
        res["scipy"] = have_scipy
        return res
    except Exception:
        return None


# ============================================================
# REPORTS
# ============================================================
def _report_text(sym, name, res, plan, sr, nse, htf):
    L = []
    L.append(f"{name} [{sym}] — AI Trader Pro analysis")
    L.append(datetime.now().strftime("Generated %d %b %Y, %H:%M"))
    L.append("=" * 52)
    L.append(f"Price: Rs {res['price']:.2f}")
    L.append(f"Trend: {res['trend']} | Signal: {res['sig']} ({res['conf']:.0f}%)")
    L.append(f"Buy%/Sell%: {res['bp']:.0f}/{res['sp']:.0f}")
    if htf and htf.get('trend'):
        L.append(f"Higher timeframe ({htf['label']}): {htf['trend']}")
    if nse and nse.get('upper'):
        L.append(f"Circuit: Upper Rs {nse['upper']:.2f} / Lower Rs {nse['lower']:.2f} ({nse.get('band','')})")
    elif nse and nse.get('no_band'):
        L.append("Circuit: No band (F&O)")
    if nse and nse.get('iep'):
        L.append(f"Pre-open (IEP): Rs {nse['iep']:.2f} ({nse.get('iep_pct',0):+.2f}%)")
    L.append("")
    L.append(f"PLAN ({plan['when']}): {plan['act']}")
    L.append(f"  Timing : {plan['timing']}")
    L.append(f"  Buy at : Rs {plan['buy_at']:.2f}")
    L.append(f"  Stop   : Rs {plan['sl']:.2f} ({plan['sl_pct']:.2f}%)")
    L.append(f"  T1     : Rs {plan['t1']:.2f} (+{plan['t1p']:.2f}%)")
    L.append(f"  T2     : Rs {plan['t2']:.2f} (+{plan['t2p']:.2f}%)")
    L.append(f"  T3     : Rs {plan['t3']:.2f} (+{plan['t3p']:.2f}%)")
    L.append(f"  R:R    : 1:{plan['rr']}")
    if sr and sr.get('pivots'):
        L.append("")
        L.append("PIVOTS")
        for sn, d in sr['pivots'].items():
            L.append(f"  {sn}: " + " ".join(f"{k}={v:.2f}" for k, v in d.items()))
    L.append("")
    L.append(f"INDICATORS ({res['n_sigs']})")
    for s in res['sigs']:
        L.append(f"  {s['n']:<16} {str(s['v']):<10} {'BUY' if s['b'] else 'SELL':<5} {s['t']}")
    L.append("")
    L.append("Educational only. Not financial advice. Always use a stop loss.")
    return "\n".join(L)


def build_report_html(sym, name, iv, per, res, plan, sr, nse, htf, mst_s):
    now = datetime.now().strftime("%d %b %Y, %H:%M")
    tr = res['trend']
    def row(k, v): return f"<tr><td class='k'>{k}</td><td class='v'>{v}</td></tr>"
    if nse and nse.get('upper'):
        circ = f"Upper ₹{nse['upper']:.2f} · Lower ₹{nse['lower']:.2f} · Band {nse.get('band','')} · Prev ₹{nse.get('prev',0):.2f}"
    elif nse and nse.get('no_band'):
        circ = "No price band (F&O stock)"
    else:
        circ = "unavailable"
    preopen = (f"₹{nse['iep']:.2f} ({nse.get('iep_chg',0):+.2f} / {nse.get('iep_pct',0):+.2f}%)"
               if nse and nse.get('iep') else "—")
    htf_txt = f"{htf['label']} trend: {htf['trend']}" if htf and htf.get('trend') else "—"
    sr_rows = ""
    if sr and sr.get('levels'):
        for l in sr['levels']:
            sr_rows += f"<tr><td>{l['name']}</td><td>{l['type']}</td><td>₹{l['price']:.2f}</td><td>{l['dist']:+.2f}%</td></tr>"
    piv_html = ""
    if sr and sr.get('pivots'):
        for sysname, d in sr['pivots'].items():
            cells = " · ".join(f"{k} ₹{v:.2f}" for k, v in d.items())
            piv_html += f"<p><b>{sysname}:</b> {cells}</p>"
    fib_html = ""
    if sr and sr.get('fib_levels'):
        fib_html = " · ".join(f"{k} ₹{v:.2f}" for k, v in sr['fib_levels'].items())
    ind_rows = ""
    for s in res['sigs']:
        ind_rows += f"<tr><td>{s['n']}</td><td>{s['v']}</td><td>{'BUY' if s['b'] else 'SELL'}</td><td>{s['t']}</td></tr>"
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>{name} Analysis</title>
<style>
body{{font-family:Arial,Helvetica,sans-serif;color:#1a1f36;max-width:900px;margin:20px auto;padding:0 20px;}}
h1{{color:#1d4ed8;margin:0;}} h2{{color:#1d4ed8;border-bottom:2px solid #e0e7ff;padding-bottom:4px;margin-top:26px;}}
.sub{{color:#6b7280;font-size:13px;margin:4px 0 16px;}}
.big{{font-size:28px;font-weight:800;}}
table{{border-collapse:collapse;width:100%;font-size:13px;margin:6px 0;}}
td,th{{border:1px solid #e0e7ff;padding:6px 10px;text-align:left;}} th{{background:#eff6ff;}}
.k{{color:#6b7280;width:34%;}} .v{{font-weight:600;}}
.verdict{{padding:14px 18px;border-radius:10px;font-size:20px;font-weight:800;margin:10px 0;}}
.up{{background:#f0fdf4;color:#15803d;border:2px solid #16a34a;}}
.dn{{background:#fff1f2;color:#b91c1c;border:2px solid #dc2626;}}
.wt{{background:#fffbeb;color:#b45309;border:2px solid #d97706;}}
.disc{{color:#9ca3af;font-size:11px;border-top:1px solid #e0e7ff;margin-top:24px;padding-top:10px;}}
@media print{{h2{{page-break-after:avoid;}}}}
</style></head><body>
<h1>💹 {name} <span style='font-size:16px;color:#6b7280;'>[{sym}]</span></h1>
<div class='sub'>Generated {now} · Timeframe {iv} · Period {per} · Market {mst_s.upper()}</div>
<div class='big'>₹{res['price']:,.2f}</div>
<div class='verdict {"up" if tr=="UPTREND" else "dn" if tr=="DOWNTREND" else "wt"}'>
Trend: {tr}  |  Signal: {res['sig']} ({res['conf']:.0f}% confidence)</div>
<h2>Summary</h2><table>
{row("Signal", res['sig'] + f" · Buy {res['bp']:.0f}% / Sell {res['sp']:.0f}%")}
{row("Trend", tr)}
{row("Higher timeframe", htf_txt)}
{row("Plan", plan['head'])}
{row("Timing", plan['timing'])}
{row("Circuit limits", circ)}
{row("Pre-open (IEP)", preopen)}
{row("52-week range", f"₹{res['lo52']:.2f} – ₹{res['hi52']:.2f} (pos {res['pos52']:.0f}%)")}
{row("ATR / Volume", f"₹{res['atr']:.2f} · {res['vr']:.1f}x avg")}
</table>
<h2>Trade Levels ({plan['when']})</h2><table>
{row("Action", plan['act'])}
{row("Buy at", f"₹{plan['buy_at']:.2f}")}
{row("Stop loss", f"₹{plan['sl']:.2f} ({plan['sl_pct']:.2f}% below)")}
{row("Target 1", f"₹{plan['t1']:.2f} (+{plan['t1p']:.2f}%)")}
{row("Target 2", f"₹{plan['t2']:.2f} (+{plan['t2p']:.2f}%)")}
{row("Target 3", f"₹{plan['t3']:.2f} (+{plan['t3p']:.2f}%)")}
{row("Risk : Reward", f"1 : {plan['rr']}")}
</table>
<h2>Support &amp; Resistance</h2>
<table><tr><th>Level</th><th>Type</th><th>Price</th><th>Distance</th></tr>{sr_rows}</table>
<h3>Pivots (Standard / Camarilla / Woodie)</h3>{piv_html}
<h3>Fibonacci</h3><p>{fib_html}</p>
<h2>All {res['n_sigs']} Indicators</h2>
<table><tr><th>Indicator</th><th>Value</th><th>Bias</th><th>Reading</th></tr>{ind_rows}</table>
<div class='disc'>⚠️ EDUCATIONAL PURPOSE ONLY · NOT FINANCIAL ADVICE · ALWAYS USE A STOP LOSS ·
Trade at your own risk · Generated by AI Trader Pro v13.6. To save as PDF: open this file and press Ctrl+P → "Save as PDF".</div>
</body></html>"""


def build_report_docx(sym, name, iv, per, res, plan, sr, nse, htf, mst_s):
    try:
        from docx import Document
        from docx.shared import Pt
        import io as _io
    except Exception:
        return None
    try:
        doc = Document()
        doc.add_heading(f"{name} [{sym}] — Analysis", level=0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')} · "
                          f"{iv} / {per} · Market {mst_s.upper()}")
        doc.add_paragraph(f"Price ₹{res['price']:.2f}").runs[0].bold = True
        p = doc.add_paragraph()
        r = p.add_run(f"Trend: {res['trend']} | Signal: {res['sig']} ({res['conf']:.0f}%)")
        r.bold = True; r.font.size = Pt(14)

        def tbl(title, rows):
            doc.add_heading(title, level=1)
            t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
            for k, v in rows:
                c = t.add_row().cells; c[0].text = str(k); c[1].text = str(v)

        htf_txt = f"{htf['label']}: {htf['trend']}" if htf and htf.get('trend') else "—"
        circ = (f"Upper ₹{nse['upper']:.2f} · Lower ₹{nse['lower']:.2f} · {nse.get('band','')}"
                if nse and nse.get('upper') else ("No band" if nse and nse.get('no_band') else "unavailable"))
        preopen = (f"₹{nse['iep']:.2f} ({nse.get('iep_pct',0):+.2f}%)" if nse and nse.get('iep') else "—")
        tbl("Summary", [("Signal", f"{res['sig']} · Buy {res['bp']:.0f}%/Sell {res['sp']:.0f}%"),
                        ("Trend", res['trend']), ("Higher timeframe", htf_txt),
                        ("Timing", plan['timing']), ("Circuit", circ), ("Pre-open", preopen),
                        ("52W range", f"₹{res['lo52']:.2f}-₹{res['hi52']:.2f}")])
        tbl(f"Trade Levels ({plan['when']})",
            [("Action", plan['act']), ("Buy at", f"₹{plan['buy_at']:.2f}"),
             ("Stop loss", f"₹{plan['sl']:.2f} ({plan['sl_pct']:.2f}%)"),
             ("Target 1", f"₹{plan['t1']:.2f} (+{plan['t1p']:.2f}%)"),
             ("Target 2", f"₹{plan['t2']:.2f} (+{plan['t2p']:.2f}%)"),
             ("Target 3", f"₹{plan['t3']:.2f} (+{plan['t3p']:.2f}%)"),
             ("Risk:Reward", f"1:{plan['rr']}")])
        if sr and sr.get('pivots'):
            doc.add_heading("Pivots", level=1)
            for sn, d in sr['pivots'].items():
                doc.add_paragraph(f"{sn}: " + " · ".join(f"{k} ₹{v:.2f}" for k, v in d.items()))
        doc.add_heading(f"All {res['n_sigs']} Indicators", level=1)
        it = doc.add_table(rows=1, cols=4); it.style = 'Light Grid Accent 1'
        hdr = it.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Indicator", "Value", "Bias", "Reading"
        for s in res['sigs']:
            c = it.add_row().cells
            c[0].text = s['n']; c[1].text = str(s['v'])
            c[2].text = "BUY" if s['b'] else "SELL"; c[3].text = s['t']
        doc.add_paragraph("⚠️ Educational only · Not financial advice · Always use a stop loss.")
        buf = _io.BytesIO(); doc.save(buf); return buf.getvalue()
    except Exception:
        return None


def live_candle_view(df, res):
    try:
        d = df.tail(14)
        if len(d) < 4:
            return None
        price = float(d['Close'].iloc[-1])
        closes = d['Close'].values
        greens = int((d['Close'] >= d['Open']).sum()); reds = len(d) - greens
        slope = closes[-1] - closes[-6] if len(closes) >= 6 else closes[-1] - closes[0]
        vw = res.get('vw', np.nan)
        above_vwap = (not np.isnan(vw)) and price > vw
        last_up = float(d['Close'].iloc[-1]) >= float(d['Open'].iloc[-1])
        streak = 1
        arr = (d['Close'] >= d['Open']).values
        for i in range(len(arr) - 2, -1, -1):
            if arr[i] == arr[-1]:
                streak += 1
            else:
                break
        up = slope > 0 and greens >= reds and (above_vwap or np.isnan(vw))
        dn = slope < 0 and reds > greens and (not above_vwap or np.isnan(vw))
        rng_hi = float(df['High'].tail(30).max()); rng_lo = float(df['Low'].tail(30).min())
        rpos = (price - rng_lo) / max(rng_hi - rng_lo, 1e-9)
        vd = ((price - vw) / vw * 100) if (not np.isnan(vw) and vw > 0) else 0
        extended = rpos > 0.85 or vd > 2.5
        if up and extended:
            verdict, vc, vico = "UP — BUT EXTENDED (pullback risk, don't chase)", "#ea580c", "🛑"
        elif up:
            verdict, vc, vico = "UPTREND FORMING", "#16a34a", "📈"
        elif dn:
            verdict, vc, vico = "DOWNTREND FORMING", "#dc2626", "📉"
        else:
            verdict, vc, vico = "CHOPPY — NO CLEAR DIRECTION", "#b45309", "➡️"
        blocks = ""
        for i, (_, r) in enumerate(d.iterrows()):
            g = r['Close'] >= r['Open']
            clr = "#16a34a" if g else "#dc2626"
            body = max(abs(r['Close'] - r['Open']), 0.01)
            hgt = 14 + min(int(body / max(price * 0.004, 0.01) * 6), 34)
            border = "border:2px solid #1a1f36;" if i == len(d) - 1 else ""
            blocks += f"<div title='O{r['Open']:.2f} H{r['High']:.2f} L{r['Low']:.2f} C{r['Close']:.2f}' style='width:16px;height:{hgt}px;background:{clr};border-radius:3px;{border}'></div>"
        last = d.iloc[-1]
        chg = (float(last['Close']) - float(last['Open']))
        chgp = chg / float(last['Open']) * 100 if last['Open'] else 0
        vwtxt = (f"{'above' if above_vwap else 'below'} VWAP" if not np.isnan(vw) else "")
        return {
            "verdict": verdict, "vc": vc, "vico": vico, "blocks": blocks,
            "streak": streak, "last_up": last_up, "greens": greens, "reds": reds,
            "chg": chg, "chgp": chgp, "vwtxt": vwtxt, "price": price,
            "oc": f"O ₹{last['Open']:.2f} · H ₹{last['High']:.2f} · L ₹{last['Low']:.2f} · C ₹{last['Close']:.2f}",
        }
    except Exception:
        return None


# ============================================================
# FULL ANALYSIS PAGE (single stock)
# ============================================================
def show_analysis(sym, name, iv, per):
    with st.spinner(f"⏳ Loading {name}..."):
        res = run_analysis(sym, iv, per)
    if res is None:
        st.markdown(f"""<div style='background:#fff1f2;border:2px solid #dc2626;border-radius:16px;padding:24px;'>
        <div style='font-size:18px;font-weight:800;color:#dc2626;'>❌ Cannot Load {sym}</div>
        <div style='color:#374151;font-size:13px;margin-top:8px;'>Try exact NSE symbol · check internet ·
        best combo: 15m timeframe + 1mo period</div></div>""", unsafe_allow_html=True)
        return
    mst_s, ml, mm = mkt_status()
    with st.spinner("🧠 Computing SR, Fibonacci, Pivots, Plan, Circuit, HTF..."):
        nse = get_nse_quote(sym)
        ck = f"nse_good_{sym}"
        if nse and (nse.get('upper') or nse.get('no_band')):
            st.session_state[ck] = nse
        elif st.session_state.get(ck):
            nse = st.session_state[ck]
        yahoo_price = res['price']
        price_src = "Yahoo (may be ~15 min delayed)"
        if nse and nse.get('ltp') and nse['ltp'] > 0 and abs(nse['ltp'] - yahoo_price) / yahoo_price < 0.25:
            res['price'] = round(nse['ltp'], 2)
            price_src = "NSE live"
        sr = get_sr_fib(sym, res['price'], res['atr'], nse)
        wk_hi = nse.get('wk_hi') if nse else None
        wk_lo = nse.get('wk_lo') if nse else None
        if not (wk_hi and wk_lo) and sr and sr.get('hi52'):
            wk_hi, wk_lo = sr['hi52'], sr['lo52']
        if wk_hi and wk_lo:
            res['hi52'] = round(wk_hi, 2); res['lo52'] = round(wk_lo, 2)
            rng = wk_hi - wk_lo
            res['pos52'] = round((res['price'] - wk_lo) / rng * 100, 1) if rng > 0 else 50
        htf = htf_trend(sym, iv)
        ema2 = ema200_filter(sym, res['price'])
        news = get_stock_news(sym)
        session = mst_s
        plan = make_plan(res['price'], res['atr'], sr, res['sig'], session, nse)
    price = res['price']; sc = res['sc']

    gap_text = ""
    if sr:
        gg = sr['gap_pct']; gclr = "#16a34a" if gg > 0 else "#dc2626" if gg < 0 else "#6b7280"
        gap_text = f"<span style='color:{gclr};font-weight:700;font-size:13px;'> {gg:+.1f}% gap</span>"
    mclr = "#16a34a" if mst_s == "open" else "#f59e0b" if mst_s == "pre" else "#dc2626"
    st.markdown(f"""<div style='background:white;border:1px solid #e0e7ff;border-radius:18px;padding:18px 24px;
    margin-bottom:18px;box-shadow:0 4px 20px rgba(0,0,0,0.06);display:flex;justify-content:space-between;
    align-items:center;flex-wrap:wrap;gap:10px;'><div>
    <div style='display:flex;align-items:center;gap:10px;flex-wrap:wrap;'>
    <span style='font-size:26px;font-weight:900;color:#1a1f36;'>{name}</span>
    <span style='background:#e0e7ff;color:#3730a3;font-size:12px;font-weight:700;padding:4px 12px;border-radius:20px;'>{sym}</span>
    <span style='background:{mclr}22;color:{mclr};font-size:12px;font-weight:700;padding:4px 12px;border-radius:20px;border:1px solid {mclr}44;'>{ml}</span>
    {gap_text}</div>
    <div style='color:#6b7280;font-size:12px;margin-top:6px;'>⏱️ {iv} · {per} · {datetime.now().strftime("%d %b %Y %H:%M")}
     ·  Trend: <b style='color:{sc};'>{res["trend"]}</b>  ·  52W pos: <b>{res["pos52"]:.1f}%</b>
     ·  Vol: <b>{res["vr"]:.1f}x</b></div></div>
    <div style='text-align:right;'><div style='font-size:40px;font-weight:900;color:#1d4ed8;'>₹{price:,.2f}</div>
    <div style='font-size:10px;font-weight:700;color:{"#16a34a" if price_src=="NSE live" else "#b45309"};'>{"🟢 NSE LIVE" if price_src=="NSE live" else "⚠️ "+price_src.upper()+" · CHECK BROKER FOR EXACT LTP"}</div>
    <div style='color:#6b7280;font-size:12px;'>52W H: ₹{res["hi52"]:.2f} | L: ₹{res["lo52"]:.2f}</div></div></div>""",
                unsafe_allow_html=True)

    st.markdown("<div style='background:#fffbeb;border:1px solid #fde68a;border-radius:10px;"
                "padding:8px 16px;margin-bottom:10px;color:#92400e;font-size:12px;'>"
                "<b>Reality check:</b> the levels (price, circuit, support/resistance) are exact — but the "
                "<b>direction call is a probability, not a prediction</b>. Even good setups fail ~40–50% of "
                "the time. The stop-loss is your protection; size small enough that a loss doesn't hurt. "
                "Track results in the 📓 Journal — trust your measured hit-rate, not any confidence number.</div>",
                unsafe_allow_html=True)

    if mst_s == 'open':
        stc = session_time_context()
        pcol = {"early": "#16a34a", "mid": "#2563eb", "late": "#ea580c",
                "closing": "#dc2626", "done": "#dc2626"}[stc['phase']]
        pemoji = {"early": "🟢", "mid": "🔵", "late": "🟠", "closing": "🔴", "done": "🔴"}[stc['phase']]
        st.markdown(f"""<div style='background:{pcol}12;border:2px solid {pcol};border-radius:14px;
        padding:14px 20px;margin-bottom:14px;display:flex;align-items:center;gap:14px;flex-wrap:wrap;'>
        <div style='font-size:28px;'>{pemoji}</div>
        <div style='flex:1;min-width:220px;'>
        <div style='font-size:13px;font-weight:900;color:{pcol};'>⏰ {stc['left_txt'].upper()} · {stc['phase'].upper()} SESSION</div>
        <div style='color:#374151;font-size:13px;margin-top:2px;'>{stc['note']}</div></div></div>""",
                    unsafe_allow_html=True)

    buy_c = sum(1 for s in res['sigs'] if s['b']); sell_c = len(res['sigs']) - buy_c
    lc = live_candle_view(res['df'], res)
    if lc:
        live_on = st.session_state.get('live_on', False)
        if mst_s == 'open':
            title = "🔴 LIVE CANDLE ANALYSIS"
            live_tag = ("<span style='background:#dc2626;color:white;border-radius:6px;padding:2px 8px;"
                        "font-size:11px;font-weight:700;'>● LIVE</span>" if live_on else
                        "<span style='background:#9ca3af;color:white;border-radius:6px;padding:2px 8px;"
                        "font-size:11px;font-weight:700;'>turn on 🔴 LIVE auto-refresh below to update live</span>")
        else:
            title = "📊 CANDLE ANALYSIS (last session)"
            live_tag = ("<span style='background:#6b7280;color:white;border-radius:6px;padding:2px 8px;"
                        "font-size:11px;font-weight:700;'>market closed — updates live once it opens</span>")
        st.markdown(f"""<div style='background:white;border:2px solid {lc['vc']};border-radius:16px;
        padding:16px 20px;margin-bottom:14px;'>
        <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;'>
        <div style='font-size:14px;font-weight:900;color:{lc['vc']};'>{lc['vico']} {title} · {lc['verdict']}</div>
        {live_tag}</div>
        <div style='display:flex;align-items:flex-end;gap:4px;height:52px;margin:12px 0 8px;'>{lc['blocks']}</div>
        <div style='color:#374151;font-size:12px;'>Latest candle: <b style='color:{"#16a34a" if lc['last_up'] else "#dc2626"};'>{'UP' if lc['last_up'] else 'DOWN'} {lc['chgp']:+.2f}%</b>
        · {lc['streak']} in a row · last 14: {lc['greens']}🟢 / {lc['reds']}🔴 · {lc['vwtxt']}</div>
        <div style='color:#9ca3af;font-size:11px;margin-top:2px;'>{lc['oc']}</div></div>""",
                    unsafe_allow_html=True)
    st.markdown(f"""<div class='{res["bg"]}'><div style='display:flex;justify-content:space-between;
    align-items:flex-start;flex-wrap:wrap;gap:16px;'><div style='flex:1;min-width:260px;'>
    <div style='font-size:11px;color:#6b7280;font-weight:700;letter-spacing:2px;text-transform:uppercase;'>AI Signal · {res["n_sigs"]} Indicators</div>
    <div style='font-size:46px;font-weight:900;color:{sc};line-height:1.1;margin-top:8px;'>{res["sig"]}</div>
    <div style='margin-top:14px;'><span style='background:{plan["ac"]};color:white;font-size:15px;font-weight:800;
    padding:10px 24px;border-radius:24px;'>{plan["act"]}</span></div>
    <div style='color:#374151;font-size:14px;margin-top:14px;background:rgba(255,255,255,0.5);border-radius:10px;padding:10px 14px;'>{plan["msg"]}</div></div>
    <div><div style='text-align:center;background:rgba(255,255,255,0.6);border-radius:16px;padding:16px 24px;'>
    <div style='font-size:60px;font-weight:900;color:{sc};line-height:1;'>{res["conf"]:.0f}<span style='font-size:24px;'>%</span></div>
    <div style='color:#6b7280;font-size:13px;font-weight:700;margin-top:4px;'>INDICATOR AGREEMENT</div>
    <div style='color:#b45309;font-size:10px;font-weight:700;margin-top:4px;max-width:170px;'>⚠️ NOT win probability.
Indicators agreeing ≠ price will move that way.</div>
    <div style='color:#6b7280;font-size:12px;margin-top:8px;'>🟢 {buy_c} buy · 🔴 {sell_c} sell</div></div></div></div>
    <div style='display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:18px;'>
    <div style='background:rgba(22,163,74,0.1);border:1px solid rgba(22,163,74,0.3);border-radius:12px;padding:12px;'>
    <div style='display:flex;justify-content:space-between;'><span style='color:#374151;font-size:13px;font-weight:600;'>BUY STRENGTH</span>
    <span style='color:#16a34a;font-weight:900;font-size:22px;'>{res["bp"]:.0f}%</span></div>
    <div class='pbar-container' style='margin-top:8px;'><div class='pbar-buy' style='width:{res["bp"]}%;'></div></div></div>
    <div style='background:rgba(220,38,38,0.08);border:1px solid rgba(220,38,38,0.2);border-radius:12px;padding:12px;'>
    <div style='display:flex;justify-content:space-between;'><span style='color:#374151;font-size:13px;font-weight:600;'>SELL STRENGTH</span>
    <span style='color:#dc2626;font-weight:900;font-size:22px;'>{res["sp"]:.0f}%</span></div>
    <div class='pbar-container' style='margin-top:8px;'><div class='pbar-sell' style='width:{res["sp"]}%;'></div></div></div></div></div>""",
                unsafe_allow_html=True)

    tr = res['trend']; is_buy = 'BUY' in res['sig']; is_sell = 'SELL' in res['sig']
    tim = entry_timing(res, res['df'], sr, ema2)
    if tr == 'UPTREND' and not is_sell:
        if tim['quality'] == 'GOOD':
            tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("UPTREND — FRESH ENTRY",
                "Trend is up and price is NOT overextended — a reasonable entry, ideally near support.",
                "#15803d", "linear-gradient(135deg,#f0fdf4,#dcfce7)", "📈")
        elif tim['quality'] == 'CAUTION':
            tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("UPTREND — BUT NOT FRESH (wait for dip)",
                "Trend is up, but this isn't an early entry. Prefer a pullback toward support, or size smaller.",
                "#b45309", "linear-gradient(135deg,#fffbeb,#fef3c7)", "⚠️")
        else:
            tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("UPTREND — LATE / EXTENDED · DON'T CHASE",
                "The up-move likely already happened — buying here risks catching the top. WAIT for a pullback to support.",
                "#ea580c", "linear-gradient(135deg,#fff7ed,#ffedd5)", "🛑")
    elif tr == 'DOWNTREND' or is_sell:
        tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("DOWNTREND — AVOID BUYING",
            "Lower highs & lower lows. Wrong side for longs — wait for a reversal.",
            "#b91c1c", "linear-gradient(135deg,#fff1f2,#ffe4e6)", "📉")
    elif tr == 'VOLATILE':
        tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("VOLATILE — RISKY",
            "Wide swings both ways. Trade small or wait for it to settle.",
            "#b45309", "linear-gradient(135deg,#fffbeb,#fef3c7)", "⚡")
    else:
        tv_txt, tv_sub, tv_clr, tv_bg, tv_ico = ("SIDEWAYS — WAIT",
            "No clear trend yet. Let a direction form before entering.",
            "#b45309", "linear-gradient(135deg,#fffbeb,#fef3c7)", "➡️")
    adx_txt = (f" · ADX {res['adx']:.0f} "
               + ("(strong)" if not np.isnan(res.get('adx', np.nan)) and res['adx'] > 25 else "(weak)")) \
              if not np.isnan(res.get('adx', np.nan)) else ""
    htf_chip = ""
    if htf and htf.get("trend"):
        ht = htf["trend"]
        agree = (ht == "UP" and tr == "UPTREND") or (ht == "DOWN" and tr == "DOWNTREND")
        hclr = "#16a34a" if ht == "UP" else "#dc2626" if ht == "DOWN" else "#b45309"
        mark = "✓ confirms" if agree else "✗ differs" if ht in ("UP", "DOWN") else "flat"
        htf_chip = (f"<span style='background:{hclr}18;color:{hclr};border:1px solid {hclr}44;"
                    f"font-size:12px;font-weight:700;padding:4px 12px;border-radius:20px;margin-left:8px;'>"
                    f"{htf['label']} trend: {ht.title()} · {mark}</span>")
    st.markdown(f"""<div style='background:{tv_bg};border:3px solid {tv_clr};border-radius:18px;
    padding:20px 26px;margin:16px 0;display:flex;align-items:center;gap:18px;flex-wrap:wrap;'>
    <div style='font-size:46px;'>{tv_ico}</div><div style='flex:1;min-width:220px;'>
    <div style='font-size:12px;color:#6b7280;font-weight:700;letter-spacing:1px;'>TREND VERDICT{adx_txt}</div>
    <div style='font-size:30px;font-weight:900;color:{tv_clr};line-height:1.1;'>{tv_txt}</div>
    <div style='color:#374151;font-size:13px;margin-top:6px;'>{tv_sub} {htf_chip}</div></div></div>""",
                unsafe_allow_html=True)

    stage_clr = {"GOOD": "#16a34a", "CAUTION": "#b45309", "POOR": "#ea580c",
                 "AVOID": "#dc2626", "WATCH": "#2563eb", "WAIT": "#b45309"}.get(tim['quality'], "#6b7280")
    reasons = ""
    if tim['warns']:
        reasons += "<div style='margin-top:6px;'><b style='color:#b45309;'>⚠️ Late/risk signals:</b> " + \
                   " · ".join(tim['warns']) + "</div>"
    if tim['goods']:
        reasons += "<div style='margin-top:4px;'><b style='color:#16a34a;'>✓ In favour:</b> " + \
                   " · ".join(tim['goods']) + "</div>"
    st.markdown(f"""<div style='background:white;border:1px solid #e0e7ff;border-left:5px solid {stage_clr};
    border-radius:0 12px 12px 0;padding:12px 18px;margin:-6px 0 14px;'>
    <span style='background:{stage_clr}18;color:{stage_clr};font-weight:800;font-size:12px;
    padding:3px 12px;border-radius:20px;'>ENTRY TIMING: {tim['stage']} · {tim['quality']}</span>
    <span style='color:#374151;font-size:13px;margin-left:8px;'>{tim['msg']}</span>{reasons}</div>""",
                unsafe_allow_html=True)

    if ema2:
        ico = "✅" if ema2['ok'] else "⛔"
        gc = "Golden cross (50>200)" if ema2['golden'] else "Death cross (50<200)"
        gclr = "#16a34a" if ema2['golden'] else "#dc2626"
        approx = " (approx — <200 daily bars)" if ema2.get('approx') else ""
        near_txt = (" · <b>price is AT the 200 EMA — major support/resistance, watch for bounce or rejection</b>"
                    if ema2['near'] else "")
        st.markdown(f"""<div style='background:{ema2['vc']}10;border:2px solid {ema2['vc']};
        border-radius:14px;padding:14px 20px;margin-bottom:14px;'>
        <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:10px;'>
        <div><span style='font-size:13px;font-weight:900;color:{ema2['vc']};'>{ico} 200 EMA METHOD · {ema2['verdict']}</span>
        <div style='color:#374151;font-size:12px;margin-top:3px;'>Daily 200 EMA ₹{ema2['ema200']:.2f}{approx} ·
        price is <b>{abs(ema2['dist']):.1f}% {'above' if ema2['above'] else 'below'}</b> it ·
        <span style='color:{gclr};font-weight:700;'>{gc}</span>{near_txt}</div></div></div>
        <div style='color:#6b7280;font-size:11px;margin-top:6px;'>Rule of the 200 EMA method: only take longs when price is
        <b>above</b> the daily 200 EMA. Below it = bear structure, where 'uptrend' bounces usually fail.</div></div>""",
                    unsafe_allow_html=True)

    if news:
        head = (f"<div style='font-size:13px;font-weight:900;color:{news['vc']};'>📰 NEWS CHECK · {news['verdict']}"
                f" <span style='color:#9ca3af;font-weight:600;'>({news['pos']}👍 / {news['neg']}👎)</span></div>")
        rows = ""
        for it in news['items']:
            sc_ = {"pos": "#16a34a", "neg": "#dc2626", "neu": "#6b7280"}[it['sent']]
            tag = {"pos": "👍 good", "neg": "👎 bad", "neu": "• neutral"}[it['sent']]
            link = f"<a href='{it['url']}' target='_blank' style='color:#1a1f36;text-decoration:none;'>{it['title']}</a>" if it['url'] else it['title']
            rows += (f"<div style='padding:7px 0;border-top:1px solid #f1f5f9;'>"
                     f"<span style='background:{sc_}18;color:{sc_};font-size:10px;font-weight:700;padding:1px 7px;border-radius:5px;'>{tag}</span> "
                     f"<span style='font-size:13px;color:#1a1f36;'>{link}</span> "
                     f"<span style='color:#9ca3af;font-size:11px;'>· {it['pub']}</span></div>")
        st.markdown(f"""<div style='background:white;border:2px solid {news['vc']};border-radius:14px;
        padding:14px 20px;margin-bottom:14px;'>{head}{rows}
        <div style='color:#9ca3af;font-size:11px;margin-top:8px;'>Headlines are a rough keyword read — open a link to verify.
        News can override the chart, so if something big broke, trust the news over the levels.</div></div>""",
                    unsafe_allow_html=True)

    if nse:
        cc = st.columns(3)
        with cc[0]:
            st.markdown(f"<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>PREV CLOSE</div><div style='font-size:26px;font-weight:900;color:#1a1f36;'>{'₹%.2f'%nse['prev'] if nse.get('prev') else '—'}</div><div style='color:#9ca3af;font-size:10px;'>NSE</div></div>", unsafe_allow_html=True)
        with cc[1]:
            if nse.get("upper"):
                st.markdown(f"<div class='mc-red'><div style='color:#6b7280;font-size:11px;font-weight:700;'>🔴 UPPER CIRCUIT</div><div style='font-size:26px;font-weight:900;color:#dc2626;'>₹{nse['upper']:.2f}</div><div style='color:#9ca3af;font-size:10px;'>max today · {nse.get('band','')}</div></div>", unsafe_allow_html=True)
            else:
                st.markdown("<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>UPPER CIRCUIT</div><div style='font-size:18px;font-weight:800;color:#6b7280;'>No band</div><div style='color:#9ca3af;font-size:10px;'>F&amp;O stock</div></div>", unsafe_allow_html=True)
        with cc[2]:
            if nse.get("lower"):
                st.markdown(f"<div class='mc-green'><div style='color:#6b7280;font-size:11px;font-weight:700;'>🟢 LOWER CIRCUIT</div><div style='font-size:26px;font-weight:900;color:#16a34a;'>₹{nse['lower']:.2f}</div><div style='color:#9ca3af;font-size:10px;'>min today · {nse.get('band','')}</div></div>", unsafe_allow_html=True)
            else:
                st.markdown("<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>LOWER CIRCUIT</div><div style='font-size:18px;font-weight:800;color:#6b7280;'>No band</div><div style='color:#9ca3af;font-size:10px;'>F&amp;O stock</div></div>", unsafe_allow_html=True)
        if nse.get("iep") and mst_s in ("pre", "closed"):
            ic = "#16a34a" if (nse.get("iep_chg") or 0) > 0 else "#dc2626" if (nse.get("iep_chg") or 0) < 0 else "#6b7280"
            st.markdown(f"<div style='background:#eff6ff;border:2px solid #3b82f6;border-radius:12px;padding:12px 18px;margin-top:8px;'><span style='color:#1d4ed8;font-weight:800;'>🌅 Pre-open (IEP): ₹{nse['iep']:.2f}</span> <span style='color:{ic};font-weight:700;'>({nse.get('iep_chg',0):+.2f} / {nse.get('iep_pct',0):+.2f}%)</span> <span style='color:#6b7280;font-size:12px;'>— NSE's likely opening price, set in the 9:00–9:15 auction.</span></div>", unsafe_allow_html=True)
    elif not sym.startswith("^"):
        st.markdown("<div style='background:#f9fafb;border:1px solid #e5e7eb;border-radius:10px;padding:8px 14px;margin-top:6px;color:#9ca3af;font-size:12px;'>Circuit limits & pre-open unavailable (NSE not reachable right now).</div>", unsafe_allow_html=True)

    st.markdown('<div class="sh">💾 SAVE THIS ANALYSIS</div>', unsafe_allow_html=True)
    try:
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        fnbase = f"{name.replace(' ', '_')[:20]}_{stamp}"
        html_report = build_report_html(sym, name, iv, per, res, plan, sr, nse, htf, mst_s)
        docx_bytes = build_report_docx(sym, name, iv, per, res, plan, sr, nse, htf, mst_s)
        dcols = st.columns(3 if docx_bytes else 2)
        with dcols[0]:
            st.download_button("📄 Download report (HTML)", data=html_report,
                               file_name=f"{fnbase}.html", mime="text/html",
                               **STRETCH)
        with dcols[1]:
            st.download_button("📝 Download as text", data=_report_text(sym, name, res, plan, sr, nse, htf),
                               file_name=f"{fnbase}.txt", mime="text/plain",
                               **STRETCH)
        if docx_bytes:
            with dcols[2]:
                st.download_button("📘 Download Word (.docx)", data=docx_bytes,
                                   file_name=f"{fnbase}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   **STRETCH)
        st.markdown("<div style='color:#6b7280;font-size:12px;margin-top:4px;'>"
                    "Saves every detail — trend, signal, circuit, levels, all pivots, Fibonacci &amp; "
                    "all indicators. <b>For a PDF:</b> open the HTML file and press <b>Ctrl+P → Save as PDF</b>."
                    + ("" if docx_bytes else " (Install <code>python-docx</code> to also get a Word file.)")
                    + "</div>", unsafe_allow_html=True)
    except Exception as e:
        st.caption(f"Report export unavailable: {type(e).__name__}")

    jc1, jc2 = st.columns([1, 2])
    with jc1:
        if st.button("📓 Save to Journal (verify tomorrow)", **STRETCH, key="jrn_save"):
            entry = {"saved": now_ist().strftime("%Y-%m-%d %H:%M"),
                     "sym": sym, "name": name, "price": res['price'],
                     "trend": res['trend'], "signal": res['sig'], "conf": res['conf'],
                     "stage": tim['stage'], "quality": tim['quality'],
                     "sit": plan['sit'], "buy_at": plan['buy_at'], "sl": plan['sl'],
                     "t1": plan['t1'], "t2": plan['t2'], "t3": plan['t3'], "rr": plan['rr'],
                     "news": news['verdict'] if news else "—",
                     "ema200_ok": bool(ema2 and ema2.get('ok'))}
            if journal_save(entry):
                st.success("✅ Saved. Open the 📓 Journal tab tomorrow to see if it was right.")
            else:
                st.error("Couldn't write journal file.")
    with jc2:
        st.markdown("<div style='color:#6b7280;font-size:12px;margin-top:6px;'>Saves today's plan to "
                    "<code>trade_journal_NAME.json</code> (NAME = your memory name). Next day the <b>📓 Journal</b> tab checks whether "
                    "price hit your targets or stop — building your real accuracy record.</div>",
                    unsafe_allow_html=True)

    st.markdown('<div class="sh">📊 KEY METRICS</div>', unsafe_allow_html=True)
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    rv = res.get('rsi', np.nan); adxv = res.get('adx', np.nan); cciv = res.get('cci', np.nan)
    mfiv = res.get('mfi', np.nan); bbwv = res.get('bbw', np.nan)
    with m1: st.metric("RSI (14)", f"{rv:.1f}" if not np.isnan(rv) else "N/A",
                       "Oversold" if not np.isnan(rv) and rv < 30 else "Overbought" if not np.isnan(rv) and rv > 70 else "Neutral")
    with m2: st.metric("ADX", f"{adxv:.1f}" if not np.isnan(adxv) else "N/A",
                       "Strong" if not np.isnan(adxv) and adxv > 25 else "Weak trend")
    with m3: st.metric("CCI", f"{cciv:.0f}" if not np.isnan(cciv) else "N/A",
                       "Buy" if not np.isnan(cciv) and cciv < -100 else "Sell" if not np.isnan(cciv) and cciv > 100 else "Neutral")
    with m4: st.metric("MFI", f"{mfiv:.1f}" if not np.isnan(mfiv) else "N/A",
                       "Inflow" if not np.isnan(mfiv) and mfiv < 40 else "Outflow" if not np.isnan(mfiv) and mfiv > 60 else "Balanced")
    with m5: st.metric("BB Width", f"{bbwv:.2f}%" if not np.isnan(bbwv) else "N/A",
                       "⚡ SQUEEZE!" if not np.isnan(bbwv) and bbwv < 1.5 else "Normal")
    with m6: st.metric("ATR", f"₹{res['atr']:.2f}", f"σ ₹{res['std20']:.2f}" if res['std20'] else "")

    st.markdown(f'<div class="sh">{plan["head"]} · ENTRY · STOP · TARGETS</div>', unsafe_allow_html=True)
    ep1, ep2, ep3 = st.columns(3)
    with ep1:
        el = "💰 BUY NOW AT" if plan['sit'] == 'BUY_NOW' else "⏳ BUY WHEN PRICE IS"
        st.markdown(f"<div style='background:{'#f0fdf4' if plan['sit']=='BUY_NOW' else '#fff7ed'};border:2px solid {plan['ac']};border-radius:16px;padding:22px;text-align:center;'><div style='color:#6b7280;font-size:11px;font-weight:700;'>{el}</div><div style='font-size:40px;font-weight:900;color:{plan['ac']};margin:10px 0;'>₹{plan['buy_at']:,.2f}</div><div style='color:{plan['ac']};font-size:12px;font-weight:600;'>{plan['timing']}</div></div>", unsafe_allow_html=True)
    with ep2:
        st.markdown(f"<div style='background:#fff1f2;border:2px solid #dc2626;border-radius:16px;padding:22px;text-align:center;'><div style='color:#6b7280;font-size:11px;font-weight:700;'>🛑 STOP LOSS</div><div style='font-size:40px;font-weight:900;color:#dc2626;margin:10px 0;'>₹{plan['sl']:,.2f}</div><div style='color:#dc2626;font-size:13px;font-weight:700;'>{plan['sl_pct']:.1f}% below entry</div></div>", unsafe_allow_html=True)
    with ep3:
        st.markdown(f"<div style='background:#eff6ff;border:2px solid #3b82f6;border-radius:16px;padding:22px;text-align:center;'><div style='color:#6b7280;font-size:11px;font-weight:700;'>⚖️ RISK : REWARD</div><div style='font-size:40px;font-weight:900;color:#1d4ed8;margin:10px 0;'>1 : {plan['rr']}</div><div style='color:#3b82f6;font-size:13px;font-weight:700;'>at Target 2</div></div>", unsafe_allow_html=True)

    tp1, tp2, tp3 = st.columns(3)
    for col, tprice, tpct, lbl, css, clr in [
        (tp1, plan['t1'], plan['t1p'], "🥉 TARGET 1 · book part", "mc-green", "#16a34a"),
        (tp2, plan['t2'], plan['t2p'], "🥈 TARGET 2 ⭐ main", "mc-green", "#16a34a"),
        (tp3, plan['t3'], plan['t3p'], "🥇 TARGET 3 · runner", "mc-purple", "#9333ea")]:
        with col:
            st.markdown(f"<div class='{css}'><div style='font-size:12px;font-weight:700;color:#6b7280;'>{lbl}</div><div style='font-size:34px;font-weight:900;color:{clr};margin:8px 0;line-height:1;'>₹{tprice:,.2f}</div><div style='font-size:13px;color:#6b7280;font-weight:600;'>+{tpct:.2f}% from entry</div></div>", unsafe_allow_html=True)
    st.markdown(f"<div style='background:#eff6ff;border:2px solid #3b82f6;border-radius:14px;padding:14px 20px;margin:14px 0;'><span style='color:#1d4ed8;font-weight:800;font-size:14px;'>📋 Plan:</span> <span style='color:#374151;font-size:14px;'>Buy near <b>₹{plan['buy_at']:.2f}</b> · stop <b>₹{plan['sl']:.2f}</b> · book at T1 <b>₹{plan['t1']:.2f}</b>, main exit T2 <b>₹{plan['t2']:.2f}</b>. You decide the quantity based on your own money.</span></div>", unsafe_allow_html=True)

    st.markdown('<div class="sh">📐 SUPPORT · RESISTANCE · FIBONACCI · CAMARILLA · WOODIE</div>', unsafe_allow_html=True)
    if sr:
        s1, s2 = st.columns(2)
        with s1:
            ns = sr.get('ns')
            if ns:
                st.markdown(f"<div style='background:#f0fdf4;border:2px solid #16a34a;border-radius:16px;padding:18px;'><div style='color:#16a34a;font-weight:800;font-size:14px;'>🟢 NEAREST SUPPORT — BUY ZONE</div><div style='font-size:34px;font-weight:900;color:#16a34a;margin:8px 0;'>₹{ns['price']:.2f}</div><div style='color:#374151;font-size:12px;'>{abs(ns['dist']):.1f}% below · SL ₹{round(ns['price']*0.993,2):.2f}</div></div>", unsafe_allow_html=True)
        with s2:
            nr = sr.get('nr')
            if nr:
                st.markdown(f"<div style='background:#fff1f2;border:2px solid #dc2626;border-radius:16px;padding:18px;'><div style='color:#dc2626;font-weight:800;font-size:14px;'>🔴 NEAREST RESISTANCE — BOOK PROFIT</div><div style='font-size:34px;font-weight:900;color:#dc2626;margin:8px 0;'>₹{nr['price']:.2f}</div><div style='color:#374151;font-size:12px;'>{abs(nr['dist']):.1f}% above · book profit here</div></div>", unsafe_allow_html=True)

        tab_std, tab_fib, tab_cam, tab_wood = st.tabs(
            ["📐 Standard Pivots", "🌀 Fibonacci", "🎯 Camarilla", "🔷 Woodie"])
        with tab_std:
            st.markdown("<div style='background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:12px 16px;margin-bottom:12px;'><b style='color:#1d4ed8;'>📐 Full level ladder</b> — pivots + swings + 52-week + circuits, typed by position vs current price (a broken support becomes resistance).</div>", unsafe_allow_html=True)
            for lvl in sr['levels']:
                lp = lvl['price']; dist = lvl['dist']; tp = lvl['type']
                is_cur = abs(dist) < 0.5
                if is_cur: css, lbl, clr, dt = "lv-cur", "📍 CURRENT PRICE", "#3b82f6", ""
                elif tp == 'R': css, lbl, clr, dt = "lv-r", f"🔴 {lvl['name']} — RESISTANCE", "#ef4444", f"+{abs(dist):.1f}%"
                elif tp == 'P': css, lbl, clr, dt = "lv-p", f"🟡 {lvl['name']} — PIVOT", "#f59e0b", f"{abs(dist):.1f}%"
                else: css, lbl, clr, dt = "lv-s", f"🟢 {lvl['name']} — SUPPORT", "#22c55e", f"{abs(dist):.1f}%"
                st.markdown(f"<div class='{css}'><div style='display:flex;justify-content:space-between;align-items:center;'><div><span style='color:{clr};font-weight:700;font-size:13px;'>{lbl}</span><span style='color:#9ca3af;font-size:11px;margin-left:8px;'>{dt}</span></div><div style='font-size:20px;font-weight:900;color:#1a1f36;'>₹{lp:,.2f}</div></div></div>", unsafe_allow_html=True)
        with tab_fib:
            st.markdown(f"<div style='background:#faf5ff;border:1px solid #d8b4fe;border-radius:12px;padding:12px 16px;margin-bottom:12px;'><b style='color:#7e22ce;'>🌀 Fibonacci Retracement</b> — 20-day swing (High ₹{sr['sw20_hi']:.2f} → Low ₹{sr['sw20_lo']:.2f})</div>", unsafe_allow_html=True)
            for ln, lp in sr['fib_levels'].items():
                dist = round((lp - price) / price * 100, 2) if price > 0 else 0
                is_cur = abs(dist) < 0.5; is_key = ln in ['38.2%', '50%', '61.8%']
                css = "lv-cur" if is_cur else "lv-fib"; clr = "#3b82f6" if is_cur else "#a855f7"
                st.markdown(f"<div class='{css}'><div style='display:flex;justify-content:space-between;align-items:center;'><div><span style='color:{clr};font-weight:{'800' if is_key else '700'};font-size:13px;'>🌀 Fib {ln}{' ⭐ KEY' if is_key else ''}</span><span style='color:#9ca3af;font-size:11px;margin-left:8px;'>{'← HERE' if is_cur else f'{dist:+.1f}%'}</span></div><div style='font-size:20px;font-weight:900;color:#1a1f36;'>₹{lp:,.2f}</div></div></div>", unsafe_allow_html=True)
            st.markdown("<div style='background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:12px 16px;margin:12px 0 8px;'><b style='color:#1d4ed8;'>📈 Fibonacci Extensions</b> — profit targets</div>", unsafe_allow_html=True)
            for en, ep_ in sr['fib_ext'].items():
                dist = round((ep_ - price) / price * 100, 2) if price > 0 else 0
                st.markdown(f"<div class='lv-r'><div style='display:flex;justify-content:space-between;align-items:center;'><span style='color:#2563eb;font-weight:700;'>📈 Extension {en}</span><div><span style='font-size:18px;font-weight:900;color:#1a1f36;'>₹{ep_:,.2f}</span><span style='color:#6b7280;font-size:12px;margin-left:8px;'>{dist:+.1f}%</span></div></div></div>", unsafe_allow_html=True)
        with tab_cam:
            st.markdown("<div style='background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:12px 16px;margin-bottom:12px;'><b style='color:#1d4ed8;'>🎯 Camarilla Pivots</b> — tight intraday levels. R3/S3 are strong reversal zones; a close beyond R4/S4 signals a breakout.</div>", unsafe_allow_html=True)
            st.markdown(pivot_table_html(sr['pivots']['Camarilla'], price), unsafe_allow_html=True)
        with tab_wood:
            st.markdown("<div style='background:#fffbeb;border:1px solid #fde68a;border-radius:12px;padding:12px 16px;margin-bottom:12px;'><b style='color:#b45309;'>🔷 Woodie Pivots</b> — weights yesterday's close more heavily, so it reacts faster to momentum than Standard pivots.</div>", unsafe_allow_html=True)
            st.markdown(pivot_table_html(sr['pivots']['Woodie'], price), unsafe_allow_html=True)

    # ── ADVANCED ANALYTICS ──
    daily_df = sr['daily'] if sr and sr.get('daily') is not None else res['df']

    reg = detect_regime(res)
    st.markdown('<div class="sh">📊 MARKET REGIME</div>', unsafe_allow_html=True)
    rg1, rg2 = st.columns(2)
    with rg1:
        st.markdown(f"<div style='background:{reg['rc']}12;border:2px solid {reg['rc']};border-radius:14px;padding:18px;text-align:center;'><div style='color:#6b7280;font-size:11px;font-weight:700;'>CURRENT REGIME</div><div style='font-size:24px;font-weight:900;color:{reg['rc']};margin:6px 0;'>{reg['regime']}</div><div style='color:#6b7280;font-size:12px;'>ADX {reg['adx']:.0f}</div></div>" if not np.isnan(reg['adx']) else f"<div style='background:{reg['rc']}12;border:2px solid {reg['rc']};border-radius:14px;padding:18px;text-align:center;'><div style='font-size:24px;font-weight:900;color:{reg['rc']};'>{reg['regime']}</div></div>", unsafe_allow_html=True)
    with rg2:
        st.markdown(f"<div style='background:{reg['vc']}12;border:2px solid {reg['vc']};border-radius:14px;padding:18px;text-align:center;'><div style='color:#6b7280;font-size:11px;font-weight:700;'>VOLATILITY</div><div style='font-size:24px;font-weight:900;color:{reg['vc']};margin:6px 0;'>{reg['vol']}</div><div style='color:#6b7280;font-size:12px;'>strategy: {'trend-follow' if 'STRONG' in reg['regime'] else 'range/mean-revert' if 'RANG' in reg['regime'] else 'wait'}</div></div>", unsafe_allow_html=True)

    st.markdown('<div class="sh">🤖 MACHINE LEARNING FORECAST</div>', unsafe_allow_html=True)
    with st.spinner("Training ML models (RF · GB · Linear)…"):
        ml = ml_predict(sym, price)
    if ml and "error" in ml:
        st.info("Install scikit-learn for ML forecasts:  `pip install scikit-learn scipy`")
    elif not ml:
        st.caption("Not enough daily history for ML on this stock.")
    else:
        st.markdown("<div style='background:#fffbeb;border:1px solid #fde68a;border-radius:10px;padding:10px 16px;margin-bottom:10px;color:#92400e;font-size:12px;'>⚠️ ML forecasts are statistical estimates from past patterns — markets are not truly predictable. Use the <b>backtested accuracy %</b> as a reality check: near 50% means little better than a coin flip. Never trade on this alone.</div>", unsafe_allow_html=True)
        mlc = st.columns(len(ml))
        for col, (h, v) in zip(mlc, sorted(ml.items())):
            dclr = "#16a34a" if v['dir'] == "UP" else "#dc2626"
            acc = v['acc']; aclr = "#16a34a" if acc >= 55 else "#b45309" if acc >= 50 else "#dc2626"
            with col:
                st.markdown(f"<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>{h}-DAY</div><div style='font-size:22px;font-weight:900;color:{dclr};margin:4px 0;'>₹{v['price']:.2f}</div><div style='color:{dclr};font-size:12px;font-weight:700;'>{v['ret']:+.2f}% {v['dir']}</div><div style='color:#9ca3af;font-size:10px;margin-top:4px;'>95% CI ₹{v['lo']:.0f}–{v['hi']:.0f}</div><div style='background:{aclr}18;color:{aclr};border-radius:6px;padding:2px 8px;font-size:11px;font-weight:700;margin-top:6px;display:inline-block;'>{acc:.0f}% accuracy</div></div>", unsafe_allow_html=True)

    prob = target_probability(price, plan, sym)
    if prob:
        st.markdown('<div class="sh">📈 TARGET PROBABILITY</div>', unsafe_allow_html=True)
        pcols = st.columns(4)
        for col, key in zip(pcols[:3], ["T1", "T2", "T3"]):
            pv = prob[key]
            with col:
                st.markdown(f"<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>{key} · ₹{pv['price']:.2f}</div><div style='font-size:28px;font-weight:900;color:{pv['color']};margin:4px 0;'>{pv['p']:.0f}%</div><div style='color:{pv['color']};font-size:12px;font-weight:700;'>{pv['rating']} chance</div></div>", unsafe_allow_html=True)
        with pcols[3]:
            st.markdown(f"<div class='mc-red'><div style='color:#6b7280;font-size:11px;font-weight:700;'>STOP HIT</div><div style='font-size:28px;font-weight:900;color:#dc2626;margin:4px 0;'>{prob['SL']['p']:.0f}%</div><div style='color:#dc2626;font-size:12px;font-weight:700;'>risk of SL</div></div>", unsafe_allow_html=True)
        st.caption("Rough estimate from ~3-day return volatility (normal model). Real markets have fat tails — treat as a guide, not a promise.")

    vp = volume_profile(daily_df)
    if vp:
        st.markdown('<div class="sh">📊 VOLUME PROFILE</div>', unsafe_allow_html=True)
        v1, v2, v3 = st.columns(3)
        with v1: st.markdown(f"<div class='mc-purple'><div style='color:#6b7280;font-size:11px;font-weight:700;'>POINT OF CONTROL</div><div style='font-size:24px;font-weight:900;color:#9333ea;'>₹{vp['poc']:.2f}</div><div style='color:#6b7280;font-size:11px;'>highest-volume price</div></div>", unsafe_allow_html=True)
        with v2: st.markdown(f"<div class='mc-blue'><div style='color:#6b7280;font-size:11px;font-weight:700;'>VALUE AREA (70%)</div><div style='font-size:20px;font-weight:900;color:#1d4ed8;'>₹{vp['va_lo']:.0f}–₹{vp['va_hi']:.0f}</div><div style='color:#6b7280;font-size:11px;'>fair-value zone</div></div>", unsafe_allow_html=True)
        with v3: st.markdown(f"<div class='mc-white'><div style='color:#6b7280;font-size:11px;font-weight:700;'>POSITION</div><div style='font-size:13px;font-weight:700;color:#1a1f36;margin-top:8px;'>{vp['pos']}</div></div>", unsafe_allow_html=True)

    pats = detect_patterns(daily_df)
    if pats:
        st.markdown('<div class="sh">🔍 CHART PATTERNS</div>', unsafe_allow_html=True)
        for nm, action, rel in pats:
            rc = {"High": "#16a34a", "Medium": "#b45309", "Low": "#6b7280"}[rel]
            st.markdown(f"<div style='background:white;border:1px solid #e0e7ff;border-left:5px solid {rc};border-radius:0 12px 12px 0;padding:12px 18px;margin:5px 0;'><b style='color:#1a1f36;'>{nm}</b> <span style='background:{rc}18;color:{rc};border-radius:6px;padding:2px 8px;font-size:11px;font-weight:700;margin-left:8px;'>{rel} reliability</span><div style='color:#6b7280;font-size:12px;margin-top:3px;'>{action}</div></div>", unsafe_allow_html=True)
        st.caption("Pattern detection is heuristic — confirm visually on the chart before acting.")

    rm = risk_metrics(sym)
    if rm:
        st.markdown('<div class="sh">⚖️ RISK METRICS (from ~2y daily history)</div>', unsafe_allow_html=True)
        k1, k2, k3, k4, k5, k6 = st.columns(6)
        with k1: st.metric("Sharpe", f"{rm['sharpe']}", "risk-adj return")
        with k2: st.metric("Max Drawdown", f"{rm['maxdd']}%", "worst drop")
        with k3: st.metric("Calmar", f"{rm['calmar']}", "ret/drawdown")
        with k4: st.metric("Win Rate", f"{rm['win']}%", "up days")
        with k5: st.metric("VaR 95%", f"{rm['var95']}%", "1-day worst")
        with k6: st.metric("Volatility", rm['vrating'], f"{rm['avol']}% ann.")

    if sr and sr.get('daily') is not None and len(sr['daily']) >= 2:
        try:
            dd = sr['daily']
            dm = demark_pivot(float(dd['High'].iloc[-1]), float(dd['Low'].iloc[-1]),
                              float(dd['Close'].iloc[-1]), float(dd['Open'].iloc[-1]))
            if dm:
                st.markdown(f"<div style='background:#faf5ff;border:1px solid #d8b4fe;border-radius:10px;padding:10px 16px;margin-top:8px;'><b style='color:#7e22ce;'>🎯 DeMark Pivots:</b> <span style='color:#374151;'>PP ₹{dm['PP']:.2f} · R1 ₹{dm['R1']:.2f} · S1 ₹{dm['S1']:.2f}</span></div>", unsafe_allow_html=True)
        except Exception:
            pass

    st.markdown('<div class="sh">📈 ADVANCED PRICE CHART — 5 Panels</div>', unsafe_allow_html=True)
    try:
        st.plotly_chart(make_chart(res['df'], name, plan, sr), **STRETCH)
    except Exception:
        st.info("Chart couldn't render for this data — the numbers above are still valid.")

    st.markdown('<div class="sh">🔬 ALL 12 INDICATOR SIGNALS</div>', unsafe_allow_html=True)
    cats = {}
    for s in res['sigs']:
        cats.setdefault(s.get('cat', 'Other'), []).append(s)
    for cat_name, cat_sigs in cats.items():
        st.markdown(f"**{cat_name} Indicators**")
        cols = st.columns(min(len(cat_sigs), 4))
        for col, s in zip(cols, cat_sigs):
            css = "ind-buy" if s['b'] else "ind-sell"; icon = "🟢" if s['b'] else "🔴"
            with col:
                st.markdown(f"<div class='{css}'><span style='color:{'#16a34a' if s['b'] else '#dc2626'};font-weight:700;'>{icon} {s['n']}</span> <code style='background:{'#dcfce7' if s['b'] else '#fee2e2'};color:{'#166534' if s['b'] else '#991b1b'};font-size:11px;padding:2px 8px;border-radius:4px;'>{s['v']}</code><div style='color:#4b5563;font-size:11px;margin-top:4px;'>{s['t']}</div></div>", unsafe_allow_html=True)

    with st.expander("📊 Raw OHLCV Data (Last 50 candles)"):
        st.dataframe(res['df'].tail(50), **STRETCH)
    st.markdown("<div style='text-align:center;color:#9ca3af;font-size:10px;padding:14px;border-top:1px solid #e0e7ff;margin-top:20px;'>⚠️ EDUCATIONAL PURPOSE ONLY · NOT FINANCIAL ADVICE · ALWAYS USE STOP LOSS · TRADE AT YOUR OWN RISK</div>", unsafe_allow_html=True)


# ============================================================
# ============================================================
# 🔴🔴🔴  LIVE 500-STOCK DASHBOARD ENGINE  🔴🔴🔴
#   bulk batched downloads · rotating refresh · flip alerts
#   · breadth · per-stock tabs · sortable table · CSV export
# ============================================================
# ============================================================
DASH_SRC = {
    "⭐ Famous 68 (most liquid)": FAMOUS,
    "🏅 Nifty 50": NIFTY50,
    "📈 Next 50": NEXT50,
    "💎 Midcap picks": MIDCAP,
    "🧩 Curated ~200 (all lists)": ALL_STOCKS,
    "🌐 Full NSE (auto-fill to your count)": None,
    "✍️ My custom list (paste symbols)": "custom",
}

DASH_PER = {"5m": "5d", "15m": "1mo", "30m": "3mo", "1h": "3mo"}
SECTORS_FOR_BREADTH = {"💻 IT": IT_S, "🏦 Bank": BANK_S, "⚡ Power": POWER_S,
                        "🛡️ Defence": DEF_S, "🚗 Auto": AUTO_S, "💊 Pharma": PHARMA_S}


def build_watchlist(src_key, n, custom_txt=""):
    """Watchlist of yahoo symbols (+ pretty names) for the dashboard."""
    d = DASH_SRC.get(src_key)
    out = []; nm = {}; seen = set()

    def add(sym, name):
        if sym and sym not in seen:
            seen.add(sym); out.append(sym); nm[sym] = name

    if d == "custom":
        for tok in re.split(r"[,\s]+", (custom_txt or "").upper()):
            t = tok.strip()
            if not t:
                continue
            if "." not in t and not t.startswith("^"):
                t += ".NS"
            add(t, t.replace(".NS", ""))
        return out[:500], nm
    if d is None:  # full NSE: curated first, then live-universe fill
        for k, v in ALL_STOCKS.items():
            add(v, k)
        try:
            uni = fetch_universe()["sym_map"]
            for s, yh in sorted(uni.items()):
                if len(out) >= n:
                    break
                add(yh, s)
        except Exception:
            pass
        return out[:n], nm
    items = list(d.items())[:n]
    for k, v in items:
        add(v, k)
    return out, nm


@st.cache_data(ttl=90, max_entries=120, show_spinner=False)
def fetch_chunk(syms, iv, per):
    """ONE batched yfinance download for up to ~50 symbols — this is what
    makes 500-stock live scanning possible (10 requests instead of 500)."""
    try:
        data = yf.download(list(syms), period=per, interval=iv, group_by='ticker',
                           threads=True, progress=False, auto_adjust=True)
    except Exception:
        return {}
    out = {}
    try:
        if data is None or data.empty:
            return {}
        single = (len(syms) == 1) or (not isinstance(data.columns, pd.MultiIndex))
        for s in syms:
            try:
                df = data if single else data[s]
                if df is None or len(df) == 0:
                    continue
                df = df.rename(columns=str.title)
                need = {'Open', 'High', 'Low', 'Close', 'Volume'}
                if not need.issubset(set(df.columns)):
                    continue
                df = df[list(need)].dropna()
                if len(df) >= 30:
                    out[s] = df
            except Exception:
                continue
    except Exception:
        pass
    return out


def day_chg_from_intraday(df):
    """% change vs PREVIOUS day's close, from an intraday series."""
    try:
        dates = sorted(set(df.index.date))
        if len(dates) >= 2:
            prev_day = df[df.index.date == dates[-2]]
            prev_close = float(prev_day['Close'].iloc[-1])
            last = float(df['Close'].iloc[-1])
            if prev_close > 0:
                return round((last / prev_close - 1) * 100, 2)
    except Exception:
        pass
    return 0.0


def daily_context(daily):
    """Daily-chart trend + 200EMA position for one stock (fast, no fetch)."""
    out = {"dtr": "—", "ema200": None, "above200": None}
    try:
        if daily is None or len(daily) < 30:
            return out
        c = daily['Close']; n = len(c)
        e200 = ta.trend.ema_indicator(c, 200) if n >= 200 else ta.trend.ema_indicator(c, min(n - 1, 150))
        e200v = safe(e200.iloc[-1])
        d = daily.tail(60); H = d['High'].values; L = d['Low'].values
        sh = []; slo = []
        for i in range(2, len(H) - 2):
            if H[i] > max(H[i-1], H[i-2], H[i+1], H[i+2]): sh.append(H[i])
            if L[i] < min(L[i-1], L[i-2], L[i+1], L[i+2]): slo.append(L[i])
        hh = len(sh) >= 2 and sh[-1] > sh[-2]; hl = len(slo) >= 2 and slo[-1] > slo[-2]
        lh = len(sh) >= 2 and sh[-1] < sh[-2]; ll = len(slo) >= 2 and slo[-1] < slo[-2]
        if hh and hl:   tr = "UPTREND"
        elif lh and ll: tr = "DOWNTREND"
        elif hh and ll: tr = "VOLATILE"
        else:           tr = "SIDEWAYS"
        out["dtr"] = tr
        if not np.isnan(e200v) and e200v > 0:
            out["ema200"] = round(e200v, 2)
            out["above200"] = bool(float(c.iloc[-1]) > e200v)
        return out
    except Exception:
        return out


def dash_row(name, sym, intra, daily):
    """One dashboard row: intraday signal + daily trend + quick plan + score."""
    try:
        res = compute_signals(intra.tail(400))
        if not res:
            return None
        dc = daily_context(daily)
        plan = make_plan(res['price'], res['atr'], None, res['sig'])
        chg = day_chg_from_intraday(intra)
        score = (res['bp'] * 0.4 + res['conf'] * 0.2
                 + (20 if dc['dtr'] == 'UPTREND' else 0)
                 + (10 if dc['above200'] else 0)
                 + (5 if res['trend'] == 'UPTREND' else 0)
                 + min(max(res['vr'], 0), 2.5) * 2
                 + max(min(chg, 5), -5))
        spark = [round(float(x), 2) for x in intra['Close'].tail(80).tolist()]
        return {"name": name, "sym": sym, "price": res['price'], "chg": chg,
                "sig": res['sig'], "sc": res['sc'], "conf": res['conf'],
                "bp": res['bp'], "sp": res['sp'], "dtr": dc['dtr'],
                "above200": dc['above200'], "ema200": dc['ema200'],
                "rsi": res['rsi'], "adx": res['adx'], "vr": res['vr'], "atr": res['atr'],
                "itrend": res['trend'], "score": round(score, 1),
                "buy_at": plan['buy_at'], "sl": plan['sl'], "t1": plan['t1'],
                "t2": plan['t2'], "rr": plan['rr'], "act": plan['act'], "sit": plan['sit'],
                "spark": spark, "ts": time.time()}
    except Exception:
        return None


def _ingest(syms, got, got_d, names, rows, prev, alerts):
    """Build dashboard rows from fetched frames + track flips & uptrend-since."""
    for sym in syms:
        intra = got.get(sym)
        if intra is None:
            continue
        r = dash_row(names.get(sym, sym.replace(".NS", "")), sym, intra, got_d.get(sym))
        if not r:
            continue
        old = prev.get(sym)
        if old and (old.get("sig") != r["sig"] or old.get("dtr") != r["dtr"]):
            alerts.insert(0, {"ts": now_ist().strftime("%H:%M:%S"),
                              "name": r["name"], "sym": sym,
                              "txt": f"{old.get('sig', '—')} → {r['sig']} · trend {old.get('dtr', '—')} → {r['dtr']}"})
        prev[sym] = {"sig": r["sig"], "dtr": r["dtr"]}
        old_row = rows.get(sym)
        if r["dtr"] == "UPTREND":
            r["up_since"] = (old_row.get("up_since")
                             if (old_row and old_row.get("dtr") == "UPTREND" and old_row.get("up_since"))
                             else time.time())
        else:
            r["up_since"] = None
        rows[sym] = r


def _sweep(syms, iv, per, with_daily=True, progress=False, label=""):
    """Batched download of a symbol list (chunks of 50)."""
    got, got_d = {}, {}
    CH = 50
    prog = st.progress(0.0) if progress else None
    for i in range(0, len(syms), CH):
        chunk = tuple(syms[i:i + CH])
        got.update(fetch_chunk(chunk, iv, per))
        if with_daily:
            got_d.update(fetch_chunk(chunk, '1d', '6mo'))
        if prog:
            try:
                prog.progress(min((i + CH) / max(len(syms), 1), 1.0),
                              text=f"🔴 {label} {min(i + CH, len(syms))}/{len(syms)} symbols…")
            except Exception:
                prog.progress(min((i + CH) / max(len(syms), 1), 1.0))
        try:
            time.sleep(0.12)
        except Exception:
            pass
    if prog:
        prog.empty()
    return got, got_d


def dash_refresh(ss, watch, iv, per, batch, progress=False, with_daily=True, ensure_min=0):
    """Refresh the stalest/missing batch — and GUARANTEE a minimum number of
    live rows by retrying missing symbols in extra passes."""
    state = ss.get("dash") or {}
    rows = state.get("rows", {}); prev = state.get("prev", {}); alerts = state.get("alerts", [])
    names = ss.get("dash_names") or {}
    missing = [s for s in watch if s not in rows]
    if missing:
        target = missing[:batch]
    else:
        target = sorted(watch, key=lambda s: rows.get(s, {}).get("ts", 0))[:batch]
    got, got_d = _sweep(target, iv, per, with_daily, progress, "Live sweep")
    _ingest(target, got, got_d, names, rows, prev, alerts)
    # ── minimum-count guarantee: retry whatever is still missing (max 2 passes)
    pass_n = 0
    while ensure_min and len(rows) < min(ensure_min, len(watch)) and pass_n < 2:
        pass_n += 1
        miss2 = [s for s in watch if s not in rows]
        if not miss2:
            break
        if progress:
            try:
                st.info(f"🔁 Guarantee pass {pass_n} — filling {len(miss2)} missing symbols (target: ≥{ensure_min} live)…")
            except Exception:
                pass
        g2, gd2 = _sweep(miss2, iv, per, with_daily, False)
        _ingest(miss2, g2, gd2, names, rows, prev, alerts)
    rows = {s: r for s, r in rows.items() if s in watch}
    alerts[:] = alerts[:40]
    ss["dash"] = {"rows": rows, "prev": prev, "alerts": alerts}
    return ss["dash"]


@st.cache_data(ttl=90, show_spinner=False)
def fetch_indices():
    """NIFTY 50 + BANK NIFTY pulse for the terminal header."""
    out = {"NIFTY 50": None, "BANK NIFTY": None}
    try:
        got = fetch_chunk(tuple(["^NSEI", "^NSEBANK"]), "15m", "5d")
        for lbl, s in out.items():
            sym = "^NSEI" if lbl == "NIFTY 50" else "^NSEBANK"
            df = got.get(sym)
            if df is not None and len(df):
                out[lbl] = {"price": float(df['Close'].iloc[-1]),
                            "chg": day_chg_from_intraday(df)}
    except Exception:
        pass
    return out


def svg_spark(vals, w=120, h=30):
    """Tiny inline SVG sparkline — crisp on any screen, no JS needed."""
    try:
        v = [x for x in (vals or []) if x == x]
        if len(v) < 2:
            return ""
        lo, hi = min(v), max(v); rng = (hi - lo) or 1e-9
        pts = [f"{i/(len(v)-1)*w:.1f},{h-2-(x-lo)/rng*(h-4):.1f}" for i, x in enumerate(v)]
        up = v[-1] >= v[0]
        stroke = "#22c55e" if up else "#ef4444"
        fill = "rgba(34,197,94,0.12)" if up else "rgba(239,68,68,0.12)"
        return (f"<svg width='{w}' height='{h}' viewBox='0 0 {w} {h}' style='display:block;'>"
                f"<polygon points='0,{h} {' '.join(pts)} {w},{h}' fill='{fill}'/>"
                f"<polyline points='{' '.join(pts)}' fill='none' stroke='{stroke}' stroke-width='1.8'/></svg>")
    except Exception:
        return ""


def _pill(txt, fg, bg):
    return f"<span style='background:{bg};color:{fg};border-radius:7px;padding:1px 8px;font-size:10px;font-weight:800;margin-right:4px;'>{txt}</span>"


def _kpi(label, value, sub, color):
    return (f"<div style='flex:1;min-width:128px;background:#0f172a;border:1px solid #1e293b;border-radius:14px;padding:13px 15px;'>"
            f"<div style='color:#64748b;font-size:10px;font-weight:800;letter-spacing:1px;'>{label}</div>"
            f"<div style='color:{color};font-size:29px;font-weight:900;font-family:ui-monospace,Menlo,Consolas,monospace;line-height:1.2;'>{value}</div>"
            f"<div style='color:#64748b;font-size:11px;margin-top:1px;'>{sub}</div></div>")


def _lb_row(i, r, now):
    chg = r.get("chg", 0); cc = "#22c55e" if chg >= 0 else "#ef4444"
    badges = ""
    if r.get("up_since") and now - r["up_since"] < 3600:
        badges += _pill("🆕 NEW UPTREND", "#4ade80", "rgba(34,197,94,0.16)")
    if r.get("above200"):
        badges += _pill("⬆ 200EMA", "#93c5fd", "rgba(59,130,246,0.16)")
    if (r.get("vr") or 0) >= 2:
        badges += _pill(f"🔊 VOL {r['vr']:.1f}x", "#fbbf24", "rgba(245,158,11,0.16)")
    spark = svg_spark(r.get("spark"))
    return (f"<div style='display:flex;align-items:center;gap:12px;background:#0f172a;border:1px solid #1e293b;"
            f"border-left:3px solid #22c55e;border-radius:12px;padding:8px 14px;margin:5px 0;flex-wrap:wrap;'>"
            f"<div style='color:#475569;font-weight:900;font-size:15px;width:32px;font-family:monospace;'>{i}</div>"
            f"<div style='min-width:148px;'><div style='color:#f1f5f9;font-weight:800;font-size:14px;'>{r['name'][:19]}</div>"
            f"<div style='color:#64748b;font-size:10px;'>{r['sym'].replace('.NS','')} · {r['sig'].title()}</div></div>"
            f"<div style='min-width:96px;color:#e2e8f0;font-weight:800;font-size:14px;font-family:monospace;'>₹{r['price']:,.2f}</div>"
            f"<div style='min-width:72px;'><span style='color:{cc};font-weight:900;font-size:13px;font-family:monospace;'>{chg:+.2f}%</span></div>"
            f"<div style='min-width:126px;'><div style='display:flex;align-items:center;gap:6px;'>"
            f"<div style='background:#1e293b;width:70px;height:6px;border-radius:3px;'><div style='background:#22c55e;width:{min(r['conf'],100):.0f}%;height:6px;border-radius:3px;'></div></div>"
            f"<span style='color:#94a3b8;font-size:10px;font-family:monospace;'>{r['conf']:.0f}%</span></div>"
            f"<div style='display:flex;align-items:center;gap:6px;margin-top:3px;'>"
            f"<div style='background:#1e293b;width:70px;height:6px;border-radius:3px;'><div style='background:#3b82f6;width:{min(r['score'],100):.0f}%;height:6px;border-radius:3px;'></div></div>"
            f"<span style='color:#94a3b8;font-size:10px;font-family:monospace;'>score {r['score']:.0f}</span></div></div>"
            f"<div style='min-width:126px;'>{spark}<div style='margin-top:2px;'>{badges}</div></div></div>")


def spark_fig(r):
    y = r.get("spark") or []
    up = (y[-1] >= y[0]) if len(y) >= 2 else True
    clr = "#16a34a" if up else "#dc2626"
    fig = go.Figure(go.Scatter(y=y, mode="lines", line=dict(color=clr, width=2.2),
                               fill="tozeroy",
                               fillcolor="rgba(22,163,74,0.08)" if up else "rgba(220,38,38,0.08)"))
    fig.update_layout(height=230, margin=dict(l=8, r=8, t=6, b=6), showlegend=False,
                      paper_bgcolor="white", plot_bgcolor="white",
                      xaxis=dict(showticklabels=False, showgrid=False),
                      yaxis=dict(showticklabels=True, showgrid=True, gridcolor="#e0e7ff", tickfont=dict(size=9)))
    return fig


def _fmt_age(sec):
    sec = int(max(sec, 0))
    return f"{sec//60}m {sec%60}s" if sec >= 60 else f"{sec}s"


# ============================================================
# ⚡ LIVE MOVERS — follow the MONEY, not the math.
#   Pure price action: who is REALLY going up right now?
#   (steady green candles · last-hour slope · near day-high ·
#    volume confirming). No indicator predictions here.
# ============================================================
def compute_movers(got):
    """Score every stock's LIVE climb from today's 5-minute candles."""
    out = []
    today = now_ist().date()
    # effective session: today if candles exist, else the latest traded day
    # (post-midnight / pre-open → full review of the LAST session)
    all_dates = set()
    for _s, _df in got.items():
        try:
            if _df is not None and len(_df):
                for _ix in _df.index:
                    all_dates.add(_dist(_ix))
        except Exception:
            pass
    sess = max((d for d in all_dates if d <= today), default=today)
    for sym, df in got.items():
        try:
            if df is None or len(df) < 12:
                continue
            dates = [_dist(ix) for ix in df.index]
            td = [i for i, d in enumerate(dates) if d == sess]
            if len(td) < 3:      # 9:30 AM IST = 3 candles — already scorable
                continue
            d0 = td[0]
            prev_close = float(df["Close"].iloc[d0 - 1]) if d0 > 0 else float(df["Open"].iloc[d0])
            t = df.iloc[d0:td[-1] + 1]
            o = float(t["Open"].iloc[0]); last = float(t["Close"].iloc[-1])
            hi = float(t["High"].max()); lo = float(t["Low"].min())
            cl = t["Close"].values; op = t["Open"].values; vol = t["Volume"].values
            n = len(cl)
            chg_day = (last - prev_close) / prev_close * 100 if prev_close else 0.0
            g, oo = cl[-24:], op[-24:]
            green = (sum(1 for i in range(len(g)) if g[i] > oo[i]) / len(g)) if len(g) else 0.5
            base = cl[-13] if n >= 13 else cl[0]
            slope1h = (last - base) / base * 100 if base else 0.0
            pos = (last - lo) / (hi - lo) if hi > lo else 0.5
            vrec = vol[-6:]
            vold = vol[:-6] if len(vol) > 6 else vol
            vr = (vrec.mean() / vold.mean()) if (len(vold) and vold.mean() > 0) else 1.0

            def _c(x):
                return max(-1.0, min(1.0, x))
            s = (35 * _c(chg_day / 3.0) + 25 * _c((green - 0.5) / 0.35) + 20 * _c(slope1h / 1.0)
                 + 10 * _c((pos - 0.5) / 0.5) + 10 * _c((vr - 1.0) / 1.5))
            climb = round(50 + s / 2, 1)
            steady = (0.2 <= chg_day <= 4.5) and green >= 0.55
            state = ("CLIMBING" if (climb >= 68 and chg_day > 0 and slope1h > 0)
                     else "WATCH" if climb >= 58 else "—")
            out.append({"sym": sym, "last": round(last, 2), "chg_day": round(chg_day, 2),
                        "green": int(round(green * 100)), "slope1h": round(float(slope1h), 2),
                        "pos": int(round(pos * 100)), "vr": round(float(vr), 2),
                        "climb": climb, "steady": steady, "state": state,
                        "hi": round(hi, 2), "lo": round(lo, 2)})
        except Exception:
            continue
    out.sort(key=lambda m: -m["climb"])
    return out


def _mv_row(i, m, name):
    climbing = m["state"] == "CLIMBING"
    border = "#22c55e" if climbing else ("#3b82f6" if m["state"] == "WATCH" else "#334155")
    cc = "#22c55e" if m["chg_day"] >= 0 else "#ef4444"
    badge = ("<span style='background:rgba(34,197,94,.16);color:#4ade80;font-size:9.5px;"
             "font-weight:900;padding:2px 8px;border-radius:8px;'>⚡ CLIMBING</span>" if climbing else
             ("<span style='background:rgba(59,130,246,.16);color:#93c5fd;font-size:9.5px;"
               "font-weight:900;padding:2px 8px;border-radius:8px;'>👀 WATCH</span>" if m["state"] == "WATCH" else ""))
    if m["steady"]:
        badge += (" <span style='background:rgba(245,158,11,.16);color:#fbbf24;font-size:9.5px;"
                  "font-weight:900;padding:2px 8px;border-radius:8px;'>🐢 SLOW-STEADY</span>")
    sc = m["climb"]; scw = max(0, min(100, sc)); scc = "#22c55e" if sc >= 68 else ("#3b82f6" if sc >= 58 else "#64748b")
    return (f"<div style='display:flex;align-items:center;gap:12px;background:#0f172a;border:1px solid #1e293b;"
            f"border-left:3px solid {border};border-radius:12px;padding:8px 14px;margin:5px 0;flex-wrap:wrap;'>"
            f"<div style='color:#475569;font-weight:900;font-size:15px;width:26px;font-family:monospace;'>{i}</div>"
            f"<div style='min-width:150px;'><div style='color:#f1f5f9;font-weight:800;font-size:14px;'>{name[:20]}</div>"
            f"<div style='color:#64748b;font-size:10px;'>{m['sym'].replace('.NS','')} · hi ₹{m['hi']:,.2f} · lo ₹{m['lo']:,.2f}</div></div>"
            f"<div style='min-width:92px;color:#e2e8f0;font-weight:800;font-size:14px;font-family:monospace;'>₹{m['last']:,.2f}</div>"
            f"<div style='min-width:70px;color:{cc};font-weight:900;font-size:13px;font-family:monospace;'>{m['chg_day']:+.2f}%</div>"
            f"<div style='min-width:104px;'><div style='color:#64748b;font-size:9px;'>GREEN CANDLES</div>"
            f"<div style='background:#1e293b;width:70px;height:6px;border-radius:3px;'><div style='background:#22c55e;width:{m['green']}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>{m['green']}% · 1h {m['slope1h']:+.2f}%</div></div>"
            f"<div style='min-width:88px;'><div style='color:#64748b;font-size:9px;'>VOLUME · DAY POS</div>"
            f"<div style='color:{'#fbbf24' if m['vr'] >= 1.3 else '#94a3b8'};font-size:11px;font-family:monospace;font-weight:800;'>{m['vr']:.2f}×</div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>{m['pos']}% of range</div></div>"
            f"<div style='min-width:104px;'><div style='color:#64748b;font-size:9px;'>CLIMB SCORE</div>"
            f"<div style='background:#1e293b;width:74px;height:6px;border-radius:3px;'><div style='background:{scc};width:{scw}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:{scc};font-size:10px;font-family:monospace;font-weight:800;'>{sc:.0f}/100</div></div>"
            f"<div style='min-width:130px;'>{badge}</div></div>")


def live_movers_tab(ss, mst_s):
    # ♾️ AUTO-RESUME — the radar keeps running across page refreshes
    if not ss.get("mv_on"):
        _rt = rt_load().get("mv") or {}
        if _rt.get("on") and _rt.get("watch"):
            ss["mv_on"] = True
            ss["mv_watch"] = _rt["watch"]; ss["mv_names"] = _rt.get("names") or {}
            ss["mv"] = _rt.get("movers") or None
            ss["mv_last"] = _rt.get("last_scan", 0)
            ss["mv_ts_str"] = _rt.get("ts_str") or "—"
            ss["mv_prev"] = _rt.get("prev") or []
            ss["mv_alerts"] = _rt.get("alerts") or []
            if _rt.get("src"):
                ss["mv_src"] = _rt["src"]
            if _rt.get("n"):
                ss["mv_n"] = _rt["n"]
            ss["_mv_resumed"] = True

    st.markdown("""<div style='background:linear-gradient(135deg,#052e16,#14532d);border-radius:18px;
    padding:18px 22px;margin-bottom:14px;'>
    <div style='color:white;font-size:20px;font-weight:900;'>⚡ LIVE MOVERS — follow the money, not the math</div>
    <div style='color:#bbf7d0;font-size:12.5px;margin-top:6px;line-height:1.7;'>This tab makes <b>NO predictions</b>.
    It watches live 5-minute candles of the whole board and shows who is <b>actually going up NOW</b>: steady green
    candles · rising in the last hour · holding near the day's high · volume confirming. When a stock starts
    climbing you get a 🔔 alert below. 🐢 SLOW-STEADY = the "slowly going up" ones you asked for (up 0.2–4.5%
    without spiking) — usually the safest to ride.</div></div>""", unsafe_allow_html=True)

    # ── settings ──
    with st.expander("⚙️ UNIVERSE · REFRESH", expanded=not ss.get("mv_watch")):
        k1, k2 = st.columns(2)
        with k1:
            _keys = list(DASH_SRC.keys())
            _def = _keys.index(ss["dash_src"]) if ss.get("dash_src") in _keys else (
                _keys.index("🌐 Full NSE (auto-fill to your count)")
                if "🌐 Full NSE (auto-fill to your count)" in _keys else 0)
            mv_src = st.selectbox("Universe", _keys, index=_def, key="mv_src")
            mv_n = st.slider("How many stocks", 100, 500, ss.get("dash_n", 500), 50, key="mv_n")
        with k2:
            st.selectbox("Auto-refresh every", ["1 min", "2 min", "3 min", "5 min"], index=1, key="mv_int")
            st.caption("Each scan = one live sweep of the whole board (5-minute candles).")
        s1, s2, s3 = st.columns(3)
        with s1:
            start_mv = st.button("⚡ START LIVE MOVERS", type="primary", **STRETCH, key="mv_start")
        with s2:
            rescan = st.button("🔄 Scan now", **STRETCH, key="mv_rescan")
        with s3:
            stop_mv = st.button("⏹ Stop", **STRETCH, key="mv_stop")

    if stop_mv:
        ss["mv_on"] = False
        rt_clear("mv")
    if start_mv:
        watch, names = build_watchlist(mv_src, mv_n)
        ss["mv_watch"] = watch; ss["mv_names"] = names
        ss["mv_on"] = True; ss["mv"] = None; ss["mv_last"] = 0
        rt_save("mv", on=True, src=mv_src, n=mv_n, watch=watch, names=names)

    if not ss.get("mv_on"):
        st.markdown("<div style='background:#0b1220;border-radius:20px;padding:48px;text-align:center;'>"
                    "<div style='font-size:48px;'>⚡</div>"
                    "<div style='font-size:21px;font-weight:900;color:#f1f5f9;margin-top:10px;'>LIVE MOVERS RADAR</div>"
                    "<div style='color:#64748b;font-size:13px;margin-top:8px;'>Up to 500 stocks watched live · "
                    "who is climbing right now · slow-steady riders · 🔔 alerts when a climb starts</div>"
                    "<div style='color:#94a3b8;font-size:12px;margin-top:12px;'>↑ Press "
                    "<b style='color:#22c55e;'>⚡ START LIVE MOVERS</b></div></div>", unsafe_allow_html=True)
        return

    if ss.get("_mv_resumed"):
        ss["_mv_resumed"] = False
        st.info("♾️ Live Movers resumed automatically — a page refresh does NOT stop it. "
                "Press ⏹ Stop to end the session.")

    # auto-refresh
    try:
        from streamlit_autorefresh import st_autorefresh
        _sec = int(ss.get("mv_int", "2 min").split()[0]) * 60
        st_autorefresh(interval=_sec * 1000, key="mv_tick")
    except Exception:
        pass

    watch = ss.get("mv_watch") or []
    names = ss.get("mv_names") or {}
    _sec = int(ss.get("mv_int", "2 min").split()[0]) * 60
    due = time.time() - ss.get("mv_last", 0) > (_sec - 10)
    if rescan or due or not ss.get("mv"):
        with st.spinner("⚡ Scanning the board live (5-minute candles)…"):
            got, _ = _sweep(watch, "5m", "2d", with_daily=False, progress=True)
        movers = compute_movers(got)
        ss["mv"] = movers
        ss["mv_last"] = time.time()
        ss["mv_ts_str"] = now_ist().strftime("%d %b %Y · %H:%M")
        # 🔔 alerts: stocks that JUST entered CLIMBING
        prev_climb = set(ss.get("mv_prev") or [])
        now_climb = {m["sym"] for m in movers if m["state"] == "CLIMBING"}
        alerts = ss.get("mv_alerts") or []
        for m in [x for x in movers if x["sym"] in (now_climb - prev_climb)]:
            alerts.insert(0, {"ts": now_ist().strftime("%H:%M:%S"), "sym": m["sym"],
                              "name": names.get(m["sym"], m["sym"].replace(".NS", "")),
                              "txt": (f"started climbing · {m['chg_day']:+.2f}% today · {m['green']}% green candles · "
                                      f"1h {m['slope1h']:+.2f}% · vol {m['vr']:.1f}× · ₹{m['last']:,.2f}"
                                      + (" · 🐢 slow-steady" if m["steady"] else ""))})
        ss["mv_alerts"] = alerts[:40]
        ss["mv_prev"] = sorted(now_climb)
        for a in ss["mv_alerts"][:3]:
            try:
                st.toast(f"🔔 {a['name']} — {a['txt'][:70]}")
            except Exception:
                pass
        rt_save("mv", on=True, watch=watch, names=names, movers=movers,
                last_scan=ss["mv_last"], ts_str=ss.get("mv_ts_str"), prev=ss.get("mv_prev") or [],
                alerts=ss.get("mv_alerts") or [], src=ss.get("mv_src"), n=ss.get("mv_n"))

    movers = ss.get("mv") or []
    if not movers:
        if mst_s == "open":
            st.markdown("<div style='background:#0b2447;border:1px solid #3b82f6;border-radius:12px;padding:14px 18px;"
                        "color:#bfdbfe;font-size:13px;line-height:1.8;'>🕘 <b>Market is open but the session just "
                        "started.</b> Five-minute candles are still building (first ones arrive after 9:15).<br>"
                        "✔️ Press <b>🔄 Scan now</b> again — from <b>~9:30 AM</b> you'll get early scores, and they "
                        "become reliable from <b>~9:45 AM</b> (6+ candles).<br>"
                        "⏰ Keep auto-refresh ON — the radar fills up by itself.</div>", unsafe_allow_html=True)
        elif mst_s == "pre":
            st.info("🌅 Market opens at 9:15 AM IST. Press ⚡ START before the open — the radar will begin scoring "
                    "as soon as the first candles form (~9:20–9:30).")
        else:
            st.info("🔴 Market is closed — no candles for today. On the next trading day: START before 9:15, "
                    "first scores from ~9:30, reliable from ~9:45 AM IST.")
        return
    if mst_s == "open" and now_ist().time() < dtime(9, 45):
        st.markdown("<div style='background:#3f2d04;border:1px solid #f59e0b;border-radius:10px;padding:8px 14px;"
                    "color:#fde68a;font-size:12px;margin-bottom:10px;'>⚠️ Early session (before 9:45) — scores are "
                    "based on few candles. Opening spikes can fake a climb; wait for 9:45–10:00 for confirmations."
                    "</div>", unsafe_allow_html=True)

    _mv_age = int((time.time() - ss.get("mv_last", 0)) / 60)
    _mv_ts = ss.get("mv_ts_str") or "—"
    _mvc = "#4ade80" if _mv_age <= 5 else ("#fbbf24" if _mv_age <= 12 else "#f87171")
    st.markdown(f"<div style='background:#0b1220;border:1px solid #1e293b;border-radius:10px;padding:8px 14px;"
                f"color:#94a3b8;font-size:12px;margin-bottom:10px;'>🕒 Last scan: <b style='color:#e2e8f0;'>{_mv_ts} IST</b>"
                f" · <b style='color:{_mvc};'>{_mv_age} min ago</b>"
                + (" — <b style='color:#f87171;'>data is old, climbs may have ended. Press 🔄 Scan now.</b>"
                   if _mv_age > 12 and mst_s == "open" else "") + "</div>", unsafe_allow_html=True)

    climbing = [m for m in movers if m["state"] == "CLIMBING"]
    steady = [m for m in movers if m["steady"] and m["chg_day"] > 0]
    watching = [m for m in movers if m["state"] == "WATCH"]
    if mst_s == "closed":
        st.markdown("<div style='background:#3f2d04;border:1px solid #f59e0b;border-radius:10px;padding:8px 14px;"
                    "color:#fde68a;font-size:12px;margin-bottom:10px;'>🔴 Market closed — this is TODAY'S full-session "
                    "climb review: what live money actually did today. Alerts resume next trading day.</div>",
                    unsafe_allow_html=True)
    a1, a2, a3, a4 = st.columns(4)
    with a1: st.metric("⚡ Climbing now", len(climbing))
    with a2: st.metric("🐢 Slow-steady riders", len(steady))
    with a3: st.metric("👀 Watch list", len(watching))
    with a4: st.metric("Scanned", len(movers),
                       f"avg move {sum(m['chg_day'] for m in movers)/max(len(movers),1):+.2f}%")

    alerts = ss.get("mv_alerts") or []
    if alerts:
        with st.expander(f"🔔 CLIMB ALERTS — this session ({len(alerts)})", expanded=True):
            for a in alerts[:12]:
                st.markdown(f"<div style='background:rgba(34,197,94,.08);border:1px solid rgba(34,197,94,.35);"
                            f"border-radius:10px;padding:8px 14px;margin:4px 0;color:#d1fae5;font-size:12px;'>"
                            f"<b style='color:#4ade80;'>🔔 {a['name']}</b> · <span style='color:#64748b;'>{a['ts']}</span>"
                            f" · {a['txt']}</div>", unsafe_allow_html=True)

    st.markdown("<div style='color:#94a3b8;font-size:12px;font-weight:900;margin:10px 0 4px;'>"
                "⚡ TOP CLIMBERS — ranked by live climb score</div>", unsafe_allow_html=True)
    rows_html = "".join(_mv_row(i + 1, m, names.get(m["sym"], m["sym"].replace(".NS", "")))
                        for i, m in enumerate(movers[:30]))
    st.markdown(rows_html, unsafe_allow_html=True)

    with st.expander("📋 Full board — all scanned stocks (sortable)"):
        disp = pd.DataFrame([{"Stock": names.get(m["sym"], m["sym"].replace(".NS", "")),
                              "Symbol": m["sym"], "Price": m["last"], "Day%": m["chg_day"],
                              "GreenCandles%": m["green"], "Slope1h%": m["slope1h"],
                              "DayPosition%": m["pos"], "Vol×": m["vr"], "ClimbScore": m["climb"],
                              "State": m["state"], "SlowSteady": "🐢" if m["steady"] else "",
                              "ScanTime": _mv_ts}
                             for m in movers])
        try:
            st.dataframe(disp, **STRETCH, height=420, hide_index=True)
        except Exception:
            st.dataframe(disp, **STRETCH)
        try:
            st.download_button("⬇️ Download movers (CSV)", data=disp.to_csv(index=False).encode(),
                               file_name=f"live_movers_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv", **STRETCH, key="mv_csv")
        except Exception:
            pass


# ============================================================
# 🎯 COMBO — the accuracy booster: LIVE climb ∩ CALCULATION.
#   A stock qualifies only if REAL money is moving it up now
#   (⚡ climb score) AND the calculation agrees (uptrend / BUY /
#   above 200EMA / confidence). Both green = highest-probability.
# ============================================================
def combo_scan(watch, names):
    """One sweep → both engines: live movers + dashboard calculation."""
    got, gotd = _sweep(watch, "5m", "2d", with_daily=True, progress=True)
    movers = {m["sym"]: m for m in compute_movers(got)}
    rows = {}
    for sym in watch:
        intra = got.get(sym)
        if intra is None:
            continue
        try:
            r = dash_row(names.get(sym, sym.replace(".NS", "")), sym, intra, gotd.get(sym))
        except Exception:
            r = None
        if r:
            rows[sym] = r
    out = []
    for sym, r in rows.items():
        mv = movers.get(sym)
        if not mv:
            continue
        sup = res = sup_dist = res_dist = None
        try:
            d = gotd.get(sym)
            if d is not None and len(d) >= 30:
                lo10 = float(d["Low"].tail(10).min()); lo20 = float(d["Low"].tail(20).min())
                support = max(lo20 * 0.995, lo10)
                hi20 = float(d["High"].tail(20).max())
                sup, res = round(support, 2), round(hi20, 2)
                sup_dist = round((mv["last"] - support) / support * 100, 2)
                res_dist = round((hi20 - mv["last"]) / mv["last"] * 100, 2)
        except Exception:
            pass
        sig = (r.get("sig") or "").upper()
        live_pts = (2 if mv["climb"] >= 68 else 1 if mv["climb"] >= 58 else 0) \
                   + (0.5 if mv["steady"] else 0) + (0.5 if mv["vr"] >= 1.3 else 0)
        calc_pts = (1 if "BUY" in sig else 0) + (1 if r.get("dtr") == "UPTREND" else 0) \
                   + (0.5 if r.get("above200") else 0) + (0.5 if (r.get("conf") or 0) >= 65 else 0)
        combo = round(min(live_pts, 3.0) / 3.0 * 60 + min(calc_pts, 3.0) / 3.0 * 40, 1)
        if mv["climb"] >= 68 and "BUY" in sig and r.get("dtr") == "UPTREND":
            verdict = "🎯 PERFECT"
        elif mv["climb"] >= 58 and calc_pts >= 1:
            verdict = "✅ MATCH"
        elif mv["climb"] >= 58:
            verdict = "⚠️ LIVE ONLY"
        elif calc_pts >= 2:
            verdict = "🧮 CALC ONLY"
        else:
            verdict = "—"
        out.append({"name": r.get("name", sym.replace(".NS", "")), "sym": sym,
                    "price": mv["last"], "chg_day": mv["chg_day"], "climb": mv["climb"],
                    "green": mv["green"], "slope1h": mv["slope1h"], "vr": mv["vr"],
                    "steady": mv["steady"], "sig": r.get("sig"), "dtr": r.get("dtr"),
                    "conf": r.get("conf"), "above200": bool(r.get("above200")),
                    "score": r.get("score"), "combo": combo, "verdict": verdict,
                    "buy_at": r.get("buy_at"), "sl": r.get("sl"), "t1": r.get("t1"), "t2": r.get("t2"),
                    "support": sup, "resistance": res, "sup_dist": sup_dist, "res_dist": res_dist,
                    "from_support": sup_dist is not None and sup_dist <= 1.5,
                    "at_resistance": res_dist is not None and res_dist <= 1.0})
    order = {"🎯 PERFECT": 0, "✅ MATCH": 1, "⚠️ LIVE ONLY": 2, "🧮 CALC ONLY": 3, "—": 4}
    out.sort(key=lambda x: (order.get(x["verdict"], 9), -x["combo"]))
    return out


def _trade_html(c):
    """💰 BUY / SL / SELL levels + resistance warning on combo cards."""
    b, sl = c.get("buy_at"), c.get("sl")
    t1, t2 = c.get("t1"), c.get("t2")
    res, rd = c.get("resistance"), c.get("res_dist")
    if not b or not sl:
        return ""
    res_line = ""
    if res:
        wr = (" · ⚠️ AT RESISTANCE — may reverse!" if (rd or 99) <= 1.0 else "")
        res_line = (f"<div style='color:#94a3b8;font-size:9px;margin-top:2px;'>RES ₹{res:,.0f}"
                    f" (+{rd:.1f}% above){wr}</div>")
    return (f"<div style='min-width:150px;'><div style='color:#64748b;font-size:9px;'>💰 TRADE · BUY & SELL</div>"
            f"<div style='font-size:10.5px;font-family:monospace;line-height:1.5;'>"
            f"<span style='color:#4ade80;font-weight:900;'>BUY ₹{b:,.2f}</span> · "
            f"<span style='color:#f87171;'>SL ₹{sl:,.2f}</span><br>"
            f"<span style='color:#22c55e;font-weight:900;'>SELL T1 ₹{(t1 or b):,.2f}</span> · "
            f"<span style='color:#16a34a;font-weight:900;'>T2 ₹{(t2 or b):,.2f}</span></div>{res_line}</div>")


def _alive_html(c, now_map):
    """🕒 NOW vs SCAN — is the climb still alive, fading, or already ENDED?"""
    nowp = (now_map or {}).get(c["sym"])
    if not nowp or not c.get("price"):
        return ""
    d = (nowp - c["price"]) / c["price"] * 100
    if d >= -0.12:
        b, col = "✅ STILL ON", "#22c55e"
    elif d >= -0.35:
        b, col = "⚠️ FADING", "#fbbf24"
    else:
        b, col = "❌ ENDED", "#ef4444"
    return (f"<div style='min-width:108px;'><div style='color:#64748b;font-size:9px;'>🕒 NOW vs SCAN</div>"
            f"<div style='color:{col};font-size:11px;font-family:monospace;font-weight:800;'>₹{nowp:,.2f}</div>"
            f"<div style='color:{col};font-size:9.5px;font-weight:900;'>{d:+.2f}% · {b}</div></div>")


def _combo_row(i, c, scan_ts="", now_map=None):
    now_map = now_map or {}
    vc = {"🎯 PERFECT": "#22c55e", "✅ MATCH": "#4ade80",
          "⚠️ LIVE ONLY": "#fbbf24", "🧮 CALC ONLY": "#93c5fd"}.get(c["verdict"], "#64748b")
    cc = "#22c55e" if c["chg_day"] >= 0 else "#ef4444"
    cb = max(0, min(100, c["combo"]))
    ema = "200EMA ✓" if c["above200"] else "below 200EMA"
    return (f"<div style='display:flex;align-items:center;gap:12px;background:#0f172a;border:1px solid #1e293b;"
            f"border-left:3px solid {vc};border-radius:12px;padding:8px 14px;margin:5px 0;flex-wrap:wrap;'>"
            f"<div style='color:#475569;font-weight:900;font-size:15px;width:26px;font-family:monospace;'>{i}</div>"
            f"<div style='min-width:148px;'><div style='color:#f1f5f9;font-weight:800;font-size:14px;'>{c['name'][:19]}</div>"
            f"<div style='color:#64748b;font-size:10px;'>{c['sym'].replace('.NS','')} · {c['sig'].title()} · {ema}"
            f"{' · 🕒 ' + scan_ts if scan_ts else ''}</div></div>"
            f"<div style='min-width:92px;color:#e2e8f0;font-weight:800;font-size:14px;font-family:monospace;'>₹{c['price']:,.2f}</div>"
            f"<div style='min-width:68px;color:{cc};font-weight:900;font-size:13px;font-family:monospace;'>{c['chg_day']:+.2f}%</div>"
            f"<div style='min-width:98px;'><div style='color:#64748b;font-size:9px;'>⚡ LIVE CLIMB</div>"
            f"<div style='background:#1e293b;width:66px;height:6px;border-radius:3px;'><div style='background:#16a34a;width:{c['climb']:.0f}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>{c['climb']:.0f}/100 · {c['green']}% green</div></div>"
            f"<div style='min-width:98px;'><div style='color:#64748b;font-size:9px;'>🧮 CALCULATION</div>"
            f"<div style='background:#1e293b;width:66px;height:6px;border-radius:3px;'><div style='background:#3b82f6;width:{min(c['conf'] or 0,100):.0f}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>conf {c['conf']:.0f}% · {str(c.get('dtr') or '—')[:8].title()}</div></div>"
            f"{_alive_html(c, now_map)}"
            f"{_trade_html(c)}"
            f"<div style='min-width:98px;'><div style='color:#64748b;font-size:9px;'>🎯 COMBO SCORE</div>"
            f"<div style='background:#1e293b;width:66px;height:6px;border-radius:3px;'><div style='background:{vc};width:{cb}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:{vc};font-size:10px;font-family:monospace;font-weight:800;'>{c['combo']:.0f}/100</div></div>"
            f"<div style='min-width:120px;'><span style='background:{'rgba(34,197,94,.16)' if 'PERFECT' in c['verdict'] or 'MATCH' in c['verdict'] else 'rgba(59,130,246,.16)'};"
            f"color:{vc};font-size:10.5px;font-weight:900;padding:3px 10px;border-radius:8px;'>{c['verdict']}</span>"
            + (" <span style='background:rgba(245,158,11,.16);color:#fbbf24;font-size:9.5px;font-weight:900;padding:2px 7px;border-radius:8px;'>🐢</span>" if c["steady"] else "")
            + (" <span style='background:rgba(59,130,246,.16);color:#93c5fd;font-size:9.5px;font-weight:900;padding:2px 7px;border-radius:8px;'>🛡️ FROM SUPPORT</span>" if c.get("from_support") else "")
            + (" <span style='background:rgba(239,68,68,.14);color:#f87171;font-size:9.5px;font-weight:900;padding:2px 7px;border-radius:8px;'>⚠️ AT RESISTANCE</span>" if c.get("at_resistance") else "")
            + "</div></div>")


# ============================================================
# 🚀 UPTREND STARTING — the support-bounce radar.
#   As per our calculation: which stock has REACHED its support
#   point and is now slowly turning up? Buy at the bounce,
#   sell into strength — levels shown on every card.
# ============================================================
def compute_bounces(got, gotd):
    out = []
    all_dates = set()
    for _s, _df in got.items():
        try:
            if _df is not None and len(_df):
                for _ix in _df.index:
                    all_dates.add(_dist(_ix))
        except Exception:
            pass
    today = now_ist().date()
    sess = max((d for d in all_dates if d <= today), default=today)
    for sym, df in got.items():
        try:
            if df is None or len(df) < 12:
                continue
            dates = [_dist(ix) for ix in df.index]
            td = [i for i, d in enumerate(dates) if d == sess]
            if len(td) < 3:
                continue
            d0 = td[0]
            prev_close = float(df["Close"].iloc[d0 - 1]) if d0 > 0 else float(df["Open"].iloc[d0])
            t = df.iloc[d0:td[-1] + 1]
            last = float(t["Close"].iloc[-1]); hi = float(t["High"].max()); lo = float(t["Low"].min())
            cl = t["Close"].values; op = t["Open"].values; vol = t["Volume"].values
            chg_day = (last - prev_close) / prev_close * 100 if prev_close else 0.0
            g, oo = cl[-24:], op[-24:]
            green = (sum(1 for i in range(len(g)) if g[i] > oo[i]) / len(g)) if len(g) else 0.5
            base = cl[-13] if len(cl) >= 13 else cl[0]
            slope1h = (last - base) / base * 100 if base else 0.0
            rec = (last - lo) / lo * 100 if lo else 0.0
            vrec = vol[-6:]; vold = vol[:-6] if len(vol) > 6 else vol
            vr = (vrec.mean() / vold.mean()) if (len(vold) and vold.mean() > 0) else 1.0
            d = gotd.get(sym)
            if d is None or len(d) < 30:
                continue
            lo10 = float(d["Low"].tail(10).min()); lo20 = float(d["Low"].tail(20).min())
            support = max(lo20 * 0.995, lo10)
            hi20 = float(d["High"].tail(20).max())
            sup_dist = (last - support) / support * 100
            res_dist = (hi20 - last) / last * 100
            if not (-0.5 <= sup_dist <= 1.5):     # only stocks AT/ON support
                continue
            turning = (rec >= 0.25 and slope1h > -0.05) or slope1h > 0.15 or green >= 0.65
            starting = turning and chg_day <= 4.0
            score = (40 * max(0.0, (1.5 - sup_dist)) / 1.5 + 30 * min(rec, 1.5) / 1.5
                     + 15.0 * green + 15.0 * min(vr, 2.0) / 2.0)
            sl = round(support * 0.985, 2)
            r_ = max(last - sl, last * 0.004)
            buy = round(min(last, support * 1.01), 2)
            out.append({"sym": sym, "last": round(last, 2), "chg_day": round(chg_day, 2),
                        "support": round(support, 2), "resistance": round(hi20, 2),
                        "sup_dist": round(sup_dist, 2), "res_dist": round(res_dist, 2),
                        "rec": round(rec, 2), "green": int(round(green * 100)),
                        "slope1h": round(float(slope1h), 2), "vr": round(float(vr), 2),
                        "score": round(score, 1),
                        "state": "🚀 UPTREND STARTING" if starting else "🛡️ AT SUPPORT",
                        "starting": starting, "buy": buy, "sl": sl,
                        "t1": round(last + 1.5 * r_, 2), "t2": round(last + 2.5 * r_, 2)})
        except Exception:
            continue
    out.sort(key=lambda b: (not b["starting"], -b["score"]))
    return out


def _bc_row(i, b, name):
    st_ = "#22c55e" if b["starting"] else "#3b82f6"
    cc = "#22c55e" if b["chg_day"] >= 0 else "#ef4444"
    badge = (f"<span style='background:rgba(34,197,94,.16);color:#4ade80;font-size:10px;font-weight:900;"
             f"padding:3px 10px;border-radius:8px;'>🚀 UPTREND STARTING</span>" if b["starting"] else
             f"<span style='background:rgba(59,130,246,.16);color:#93c5fd;font-size:10px;font-weight:900;"
             f"padding:3px 10px;border-radius:8px;'>🛡️ AT SUPPORT · waiting for turn</span>")
    return (f"<div style='display:flex;align-items:center;gap:12px;background:#0f172a;border:1px solid #1e293b;"
            f"border-left:3px solid {st_};border-radius:12px;padding:8px 14px;margin:5px 0;flex-wrap:wrap;'>"
            f"<div style='color:#475569;font-weight:900;font-size:15px;width:26px;font-family:monospace;'>{i}</div>"
            f"<div style='min-width:140px;'><div style='color:#f1f5f9;font-weight:800;font-size:14px;'>{name[:19]}</div>"
            f"<div style='color:#64748b;font-size:10px;'>{b['sym'].replace('.NS','')} · +{b['rec']:.2f}% off day's low</div></div>"
            f"<div style='min-width:88px;color:#e2e8f0;font-weight:800;font-size:14px;font-family:monospace;'>₹{b['last']:,.2f}</div>"
            f"<div style='min-width:64px;color:{cc};font-weight:900;font-size:13px;font-family:monospace;'>{b['chg_day']:+.2f}%</div>"
            f"<div style='min-width:118px;'><div style='color:#64748b;font-size:9px;'>🛡️ SUPPORT (BUY ZONE)</div>"
            f"<div style='color:#4ade80;font-size:11.5px;font-family:monospace;font-weight:800;'>₹{b['support']:,.2f}</div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>{b['sup_dist']:+.2f}% away</div></div>"
            f"<div style='min-width:118px;'><div style='color:#64748b;font-size:9px;'>💰 TRADE LEVELS</div>"
            f"<div style='font-size:10.5px;font-family:monospace;line-height:1.5;'>"
            f"<span style='color:#4ade80;font-weight:900;'>BUY ₹{b['buy']:,.2f}</span> · "
            f"<span style='color:#f87171;'>SL ₹{b['sl']:,.2f}</span><br>"
            f"<span style='color:#22c55e;font-weight:900;'>SELL T1 ₹{b['t1']:,.2f}</span> · "
            f"<span style='color:#16a34a;font-weight:900;'>T2 ₹{b['t2']:,.2f}</span></div></div>"
            f"<div style='min-width:104px;'><div style='color:#64748b;font-size:9px;'>🎯 RESISTANCE (SELL ZONE)</div>"
            f"<div style='color:#fbbf24;font-size:11.5px;font-family:monospace;font-weight:800;'>₹{b['resistance']:,.2f}</div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>+{b['res_dist']:.2f}% above"
            + (" · ⚠️ NEAR" if b["res_dist"] <= 1.0 else "") + "</div></div>"
            f"<div style='min-width:98px;'><div style='color:#64748b;font-size:9px;'>BOUNCE SCORE</div>"
            f"<div style='background:#1e293b;width:70px;height:6px;border-radius:3px;'><div style='background:{st_};width:{min(b['score'],100):.0f}%;height:6px;border-radius:3px;'></div></div>"
            f"<div style='color:#94a3b8;font-size:9.5px;font-family:monospace;'>{b['score']:.0f}/100 · {b['green']}% green</div></div>"
            f"<div style='min-width:150px;'>{badge}</div></div>")


def bounce_tab(ss, mst_s):
    # ♾️ AUTO-RESUME
    if not ss.get("bc_on"):
        _rt = rt_load().get("bc") or {}
        if _rt.get("on") and _rt.get("watch"):
            ss["bc_on"] = True
            ss["bc_watch"] = _rt["watch"]; ss["bc_names"] = _rt.get("names") or {}
            ss["bc"] = _rt.get("bounces") or None
            ss["bc_last"] = _rt.get("last_scan", 0)
            ss["bc_ts_str"] = _rt.get("ts_str") or "—"
            ss["bc_prev"] = _rt.get("prev") or []
            ss["bc_alerts"] = _rt.get("alerts") or []
            ss["_bc_resumed"] = True

    st.markdown("""<div style='background:linear-gradient(135deg,#0f2b1e,#1e3a8a);border-radius:18px;
    padding:18px 22px;margin-bottom:12px;'>
    <div style='color:white;font-size:20px;font-weight:900;'>🚀 UPTREND STARTING — support-bounce radar</div>
    <div style='color:#bbf7d0;font-size:12.5px;margin-top:6px;line-height:1.7;'>As per our calculation: which stocks
    have <b>REACHED their SUPPORT point</b> and are <b>slowly starting to turn up</b>? 🚀 UPTREND STARTING = at
    support + bouncing now — the safest buy zone (buy at support, stop just below, sell into T1/T2).
    🛡️ AT SUPPORT = sitting on support, turn not confirmed yet — watchlist it. Every card shows the exact
    <b>BUY, STOP-LOSS and SELL</b> values, plus the resistance where the rally may stall.</div></div>""",
                unsafe_allow_html=True)

    with st.expander("⚙️ UNIVERSE · REFRESH", expanded=not ss.get("bc_watch")):
        k1, k2 = st.columns(2)
        with k1:
            _keys = list(DASH_SRC.keys())
            _def = _keys.index("🌐 Full NSE (auto-fill to your count)") if "🌐 Full NSE (auto-fill to your count)" in _keys else 0
            st.selectbox("Universe", _keys, index=_def, key="bc_src")
            st.slider("How many stocks", 100, 500, 500, 50, key="bc_n")
        with k2:
            st.selectbox("Auto-refresh every", ["1 min", "2 min", "3 min", "5 min"], index=2, key="bc_int")
            st.caption("One scan = live 5-min candles + daily history (support/resistance) for the whole board.")
        s1, s2, s3 = st.columns(3)
        with s1:
            start_bc = st.button("🚀 START BOUNCE RADAR", type="primary", use_container_width=True, key="bc_start")
        with s2:
            rescan_bc = st.button("🔄 Scan now", use_container_width=True, key="bc_rescan")
        with s3:
            stop_bc = st.button("⏹ Stop", use_container_width=True, key="bc_stop")

    if stop_bc:
        ss["bc_on"] = False
        rt_clear("bc")
    if start_bc:
        watch, names = build_watchlist(ss.get("bc_src"), ss.get("bc_n", 500))
        ss["bc_watch"] = watch; ss["bc_names"] = names
        ss["bc_on"] = True; ss["bc"] = None; ss["bc_last"] = 0; ss["bc_alerts"] = []; ss["bc_prev"] = []
        rt_save("bc", on=True, src=ss.get("bc_src"), n=ss.get("bc_n", 500), watch=watch, names=names)

    if not ss.get("bc_on"):
        st.markdown("<div style='background:#0b1220;border-radius:20px;padding:44px;text-align:center;'>"
                    "<div style='font-size:44px;'>🚀</div>"
                    "<div style='font-size:20px;font-weight:900;color:#f1f5f9;margin-top:10px;'>SUPPORT-BOUNCE RADAR</div>"
                    "<div style='color:#64748b;font-size:13px;margin-top:8px;'>All-India board · who just reached support "
                    "and started turning up · exact BUY / SL / SELL levels · 🔔 alerts</div>"
                    "<div style='color:#94a3b8;font-size:12px;margin-top:12px;'>↑ Press "
                    "<b style='color:#22c55e;'>🚀 START BOUNCE RADAR</b></div></div>", unsafe_allow_html=True)
        return

    if ss.get("_bc_resumed"):
        ss["_bc_resumed"] = False
        st.info("♾️ Bounce radar resumed automatically — a page refresh does NOT stop it. Press ⏹ Stop to end it.")

    try:
        from streamlit_autorefresh import st_autorefresh
        _sec = int(ss.get("bc_int", "3 min").split()[0]) * 60
        st_autorefresh(interval=_sec * 1000, key="bc_tick")
    except Exception:
        pass

    watch = ss.get("bc_watch") or []
    names = ss.get("bc_names") or {}
    _sec = int(ss.get("bc_int", "3 min").split()[0]) * 60
    due = time.time() - ss.get("bc_last", 0) > (_sec - 10)
    if rescan_bc or due or not ss.get("bc"):
        with st.spinner("🚀 Scanning for support bounces (live candles + support levels)…"):
            got, gotd = _sweep(watch, "5m", "2d", with_daily=True, progress=True)
            bounces = compute_bounces(got, gotd)
            ss["bc"] = bounces
            ss["bc_last"] = time.time()
            ss["bc_ts_str"] = now_ist().strftime("%d %b %Y · %H:%M")
            prev_start = set(ss.get("bc_prev") or [])
            now_start = {b["sym"] for b in bounces if b["starting"]}
            alerts = ss.get("bc_alerts") or []
            for b in [x for x in bounces if x["sym"] in (now_start - prev_start)]:
                alerts.insert(0, {"ts": now_ist().strftime("%H:%M:%S"), "sym": b["sym"],
                                  "name": names.get(b["sym"], b["sym"].replace(".NS", "")),
                                  "txt": (f"reached support ₹{b['support']:,.2f} and TURNING UP · now ₹{b['last']:,.2f} "
                                          f"(+{b['rec']:.2f}% off low) · BUY ₹{b['buy']:,.2f} · SL ₹{b['sl']:,.2f} "
                                          f"· SELL T1 ₹{b['t1']:,.2f}")})
            ss["bc_alerts"] = alerts[:40]
            ss["bc_prev"] = sorted(now_start)
            for a in ss["bc_alerts"][:3]:
                try:
                    st.toast(f"🚀 {a['name']} — {a['txt'][:70]}")
                except Exception:
                    pass
            rt_save("bc", on=True, watch=watch, names=names, bounces=bounces,
                    last_scan=ss["bc_last"], ts_str=ss["bc_ts_str"],
                    prev=ss["bc_prev"], alerts=ss["bc_alerts"],
                    src=ss.get("bc_src"), n=ss.get("bc_n"))

    bounces = ss.get("bc") or []
    if not bounces:
        st.info("No stocks are sitting at support right now — that's fine (cash is a position). "
                "Re-scan later; 🔔 alerts fire the moment one turns up.")
        return

    _age = int((time.time() - ss.get("bc_last", 0)) / 60)
    _ts = ss.get("bc_ts_str") or "—"
    _agc = "#4ade80" if _age <= 7 else ("#fbbf24" if _age <= 15 else "#f87171")
    st.markdown(f"<div style='background:#0b1220;border:1px solid #1e293b;border-radius:10px;padding:8px 14px;"
                f"color:#94a3b8;font-size:12px;margin-bottom:10px;'>🕒 Last scan: <b style='color:#e2e8f0;'>{_ts} IST</b>"
                f" · <b style='color:{_agc};'>{_age} min ago</b>"
                + (" — <b style='color:#f87171;'>old scan. Press 🔄 Scan now.</b>" if _age > 15 and mst_s == "open" else "")
                + "</div>", unsafe_allow_html=True)

    starting = [b for b in bounces if b["starting"]]
    waiting = [b for b in bounces if not b["starting"]]
    a1, a2, a3 = st.columns(3)
    with a1: st.metric("🚀 Uptrend starting", len(starting), "at support + turning up")
    with a2: st.metric("🛡️ At support", len(waiting), "waiting for confirmation")
    with a3: st.metric("Scanned", ss.get("bc_n", 500))

    alerts = ss.get("bc_alerts") or []
    if alerts:
        with st.expander(f"🔔 BOUNCE ALERTS — this session ({len(alerts)})", expanded=True):
            for a in alerts[:12]:
                st.markdown(f"<div style='background:rgba(34,197,94,.08);border:1px solid rgba(34,197,94,.35);"
                            f"border-radius:10px;padding:8px 14px;margin:4px 0;color:#d1fae5;font-size:12px;'>"
                            f"<b style='color:#4ade80;'>🚀 {a['name']}</b> · <span style='color:#64748b;'>{a['ts']}</span>"
                            f" · {a['txt']}</div>", unsafe_allow_html=True)

    if starting:
        st.markdown(f"<div style='color:#94a3b8;font-size:12px;font-weight:900;margin:10px 0 4px;'>"
                    f"🚀 UPTREND STARTING — reached support & turning up ({len(starting)})</div>", unsafe_allow_html=True)
        st.markdown("".join(_bc_row(i + 1, b, names.get(b["sym"], b["sym"].replace(".NS", "")))
                            for i, b in enumerate(starting[:25])), unsafe_allow_html=True)
    if waiting:
        with st.expander(f"🛡️ AT SUPPORT — waiting for the turn ({len(waiting)})"):
            st.markdown("".join(_bc_row(i + 1, b, names.get(b["sym"], b["sym"].replace(".NS", "")))
                                for i, b in enumerate(waiting[:25])), unsafe_allow_html=True)
    with st.expander("📋 Full bounce board + download"):
        disp = pd.DataFrame([{"Stock": names.get(b["sym"], b["sym"].replace(".NS", "")), "Symbol": b["sym"],
                              "Price": b["last"], "Day%": b["chg_day"], "Support": b["support"],
                              "SupDist%": b["sup_dist"], "BUY@": b["buy"], "SL": b["sl"],
                              "SellT1": b["t1"], "SellT2": b["t2"], "Resistance": b["resistance"],
                              "ResDist%": b["res_dist"], "OffLow%": b["rec"], "GreenCandles%": b["green"],
                              "Slope1h%": b["slope1h"], "Vol×": b["vr"], "BounceScore": b["score"],
                              "State": b["state"], "ScanTime": _ts} for b in bounces])
        try:
            st.dataframe(disp, use_container_width=True, height=420, hide_index=True)
        except Exception:
            st.dataframe(disp, use_container_width=True)
        try:
            st.download_button("⬇️ Download bounce board (CSV)", data=disp.to_csv(index=False).encode(),
                               file_name=f"support_bounces_{now_ist().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv", use_container_width=True, key="bc_csv")
        except Exception:
            pass


def combo_tab(ss, mst_s):
    # ♾️ AUTO-RESUME — the combo scan keeps running across page refreshes
    if not ss.get("cb_on"):
        _rt = rt_load().get("cb") or {}
        if _rt.get("on") and _rt.get("watch"):
            ss["cb_on"] = True
            ss["cb_watch"] = _rt["watch"]; ss["cb_names"] = _rt.get("names") or {}
            ss["cb"] = _rt.get("combos") or None
            ss["cb_last"] = _rt.get("last_scan", 0)
            ss["cb_ts_str"] = _rt.get("ts_str") or "—"
            if _rt.get("src"):
                ss["cb_src"] = _rt["src"]
            if _rt.get("n"):
                ss["cb_n"] = _rt["n"]
            ss["_cb_resumed"] = True

    st.markdown("""<div style='background:linear-gradient(135deg,#1e1b4b,#0f3d2e);border-radius:18px;
    padding:18px 22px;margin-bottom:12px;'>
    <div style='color:white;font-size:20px;font-weight:900;'>🎯 COMBO — live money ∩ calculation</div>
    <div style='color:#c7d2fe;font-size:12.5px;margin-top:6px;line-height:1.7;'>The accuracy booster you asked for:
    a stock appears here only when <b> BOTH </b> agree — ⚡ it is <b>actually climbing right now</b> (live candles)
    <b>AND</b> 🧮 the <b>calculation</b> says uptrend/BUY/above-200EMA. 🎯 PERFECT = strongest possible agreement.
    ⚠️ LIVE ONLY = money moving but math neutral (risky momentum). 🧮 CALC ONLY = math likes it, money not yet
    (breakout watchlist).</div></div>""", unsafe_allow_html=True)
    st.markdown("""<div style='background:#0b1220;border:1px dashed #1e293b;border-radius:12px;padding:10px 16px;
    color:#94a3b8;font-size:12px;margin-bottom:10px;line-height:1.8;'>⏰ <b style='color:#e2e8f0;">Best time to run this
    (IST):</b> 🥇 <b style='color:#4ade80;'>9:45 – 11:00 AM</b> (opening noise settles, real trend confirms) ·
    🥈 1:30 – 2:30 PM (afternoon build) · 🥉 2:45 – 3:10 PM (closing strength). <b style='color:#f87171;">Avoid
    9:15–9:35</b> — opening fake spikes. One clean scan at ~9:45 + one re-scan at ~2:45 is the professional routine.
    </div>""", unsafe_allow_html=True)
    if mst_s == "open" and now_ist().time() < dtime(9, 45):
        st.markdown("<div style='background:#3f2d04;border:1px solid #f59e0b;border-radius:10px;padding:8px 14px;"
                    "color:#fde68a;font-size:12px;margin-bottom:10px;'>⚠️ It's before 9:45 AM — you CAN scan now, but "
                    "scores settle after 9:45–10:00. A re-scan at 9:45+ is worth it.</div>", unsafe_allow_html=True)

    with st.expander("⚙️ UNIVERSE · REFRESH", expanded=not ss.get("cb_watch")):
        k1, k2 = st.columns(2)
        with k1:
            _keys = list(DASH_SRC.keys())
            _def = _keys.index("🌐 Full NSE (auto-fill to your count)") if "🌐 Full NSE (auto-fill to your count)" in _keys else 0
            st.selectbox("Universe", _keys, index=_def, key="cb_src")
            st.slider("How many stocks", 100, 500, 500, 50, key="cb_n")
        with k2:
            st.selectbox("Auto-refresh every", ["1 min", "2 min", "3 min", "5 min"], index=1, key="cb_int")
            st.caption("One scan = live 5-minute candles + daily history for the whole board (~1–2 min).")
        s1, s2, s3 = st.columns(3)
        with s1:
            start_cb = st.button("🎯 START COMBO SCAN", type="primary", **STRETCH, key="cb_start")
        with s2:
            rescan_cb = st.button("🔄 Scan now", **STRETCH, key="cb_rescan")
        with s3:
            stop_cb = st.button("⏹ Stop", **STRETCH, key="cb_stop")

    if stop_cb:
        ss["cb_on"] = False
        rt_clear("cb")
    if start_cb:
        watch, names = build_watchlist(ss.get("cb_src"), ss.get("cb_n", 500))
        ss["cb_watch"] = watch; ss["cb_names"] = names
        ss["cb_on"] = True; ss["cb"] = None; ss["cb_last"] = 0
        rt_save("cb", on=True, src=ss.get("cb_src"), n=ss.get("cb_n", 500), watch=watch, names=names)

    if not ss.get("cb_on"):
        st.markdown("<div style='background:#0b1220;border-radius:20px;padding:44px;text-align:center;'>"
                    "<div style='font-size:44px;'>🎯</div>"
                    "<div style='font-size:20px;font-weight:900;color:#f1f5f9;margin-top:10px;'>COMBO SCANNER</div>"
                    "<div style='color:#64748b;font-size:13px;margin-top:8px;'>All-India board · live climb + "
                    "calculation agreement · 🎯 PERFECT picks · 🐢 slow-steady riders</div>"
                    "<div style='color:#94a3b8;font-size:12px;margin-top:12px;'>↑ Best at <b>9:45–11:00 AM IST</b> · "
                    "press <b style='color:#22c55e;'>🎯 START COMBO SCAN</b></div></div>", unsafe_allow_html=True)
        return

    if ss.get("_cb_resumed"):
        ss["_cb_resumed"] = False
        st.info("♾️ Combo scan resumed automatically — a page refresh does NOT stop it. "
                "Press ⏹ Stop to end the session.")

    try:
        from streamlit_autorefresh import st_autorefresh
        _sec = int(ss.get("cb_int", "2 min").split()[0]) * 60
        st_autorefresh(interval=_sec * 1000, key="cb_tick")
    except Exception:
        pass

    watch = ss.get("cb_watch") or []
    names = ss.get("cb_names") or {}
    _sec = int(ss.get("cb_int", "2 min").split()[0]) * 60
    due = time.time() - ss.get("cb_last", 0) > (_sec - 10)
    if rescan_cb or due or not ss.get("cb"):
        with st.spinner("🎯 Combo scan — live candles + calculation for the whole board…"):
            ss["cb"] = combo_scan(watch, names)
            ss["cb_last"] = time.time()
            ss["cb_ts_str"] = now_ist().strftime("%d %b %Y · %H:%M")
            ss.pop("cb_recheck", None)
        rt_save("cb", on=True, watch=watch, names=names, combos=ss["cb"],
                last_scan=ss["cb_last"], ts_str=ss.get("cb_ts_str"),
                src=ss.get("cb_src"), n=ss.get("cb_n", 500))

    combos = ss.get("cb") or []
    if not combos:
        st.info("No scorable stocks yet — see the timing note above (first reliable scores from ~9:45 AM IST; "
                "after close you get the full-day review).")
        return
    if mst_s == "closed":
        st.markdown("<div style='background:#3f2d04;border:1px solid #f59e0b;border-radius:10px;padding:8px 14px;"
                    "color:#fde68a;font-size:12px;margin-bottom:10px;'>🔴 Market closed — this is today's FULL-SESSION "
                    "combo review: who had both live money AND calculation agreement today.</div>", unsafe_allow_html=True)

    # 🕒 FRESHNESS — when was this scan made? (an old scan's uptrend may have ENDED)
    _age = int((time.time() - ss.get("cb_last", 0)) / 60)
    _ts = ss.get("cb_ts_str") or "—"
    _agc = "#4ade80" if _age <= 6 else ("#fbbf24" if _age <= 15 else "#f87171")
    st.markdown(f"<div style='background:#0b1220;border:1px solid #1e293b;border-radius:10px;padding:8px 14px;"
                f"color:#94a3b8;font-size:12px;margin-bottom:10px;'>🕒 Last scan: <b style='color:#e2e8f0;'>{_ts} IST</b>"
                f" · <b style='color:{_agc};'>{_age} min ago</b>"
                + (" — <b style='color:#f87171;'>old scan: an uptrend may have ENDED. Press 🔄 Scan now.</b>"
                   if _age > 15 and mst_s == "open" else "")
                + " · every card below shows 🕒 <b>NOW vs SCAN</b> — still on / fading / ended.</div>",
                unsafe_allow_html=True)

    # 🔄 live re-check of the top picks (current price vs scan price)
    _rc = ss.get("cb_recheck") or {}
    if time.time() - _rc.get("ts", 0) > 90:
        try:
            _syms = [c["sym"] for c in combos[:30]]
            _got = fetch_chunk(tuple(_syms), "5m", "1d")
            ss["cb_recheck"] = {"ts": time.time(),
                                "prices": {s: round(float(_got[s]["Close"].iloc[-1]), 2)
                                           for s in _syms if s in _got and len(_got[s])}}
        except Exception:
            ss["cb_recheck"] = {"ts": time.time(), "prices": {}}
        _rc = ss["cb_recheck"]
    _now_map = _rc.get("prices") or {}   # useful live AND after close (review)
    _scan_short = (_ts.split("·")[-1]).strip() if _ts != "—" else ""

    perfect = [c for c in combos if c["verdict"] == "🎯 PERFECT"]
    match = [c for c in combos if c["verdict"] == "✅ MATCH"]
    liveonly = [c for c in combos if c["verdict"] == "⚠️ LIVE ONLY"]
    calconly = [c for c in combos if c["verdict"] == "🧮 CALC ONLY"]
    a1, a2, a3, a4, a5 = st.columns(5)
    with a1: st.metric("🎯 Perfect", len(perfect), "live + calc agree")
    with a2: st.metric("✅ Match", len(match), "good agreement")
    with a3: st.metric("⚠️ Live only", len(liveonly), "momentum, no math")
    with a4: st.metric("🧮 Calc only", len(calconly), "breakout watchlist")
    with a5: st.metric("Scanned", len(combos))

    top = perfect + match
    if top:
        st.markdown(f"<div style='color:#94a3b8;font-size:12px;font-weight:900;margin:10px 0 4px;'>"
                    f"🎯 THE PICKS — live climb + calculation agree ({len(top)})</div>", unsafe_allow_html=True)
        st.markdown("".join(_combo_row(i + 1, c, _scan_short, _now_map) for i, c in enumerate(top[:30])),
                    unsafe_allow_html=True)
    else:
        st.info("No PERFECT/MATCH picks right now — the two engines don't agree on anything this moment. "
                "That's the system protecting you (no trade is better than a bad trade). Re-scan later.")
    if liveonly:
        with st.expander(f"⚠️ LIVE ONLY — climbing but calculation neutral ({len(liveonly)}) · higher risk"):
            st.markdown("".join(_combo_row(i + 1, c, _scan_short, _now_map) for i, c in enumerate(liveonly[:20])),
                        unsafe_allow_html=True)
    if calconly:
        with st.expander(f"🧮 CALC ONLY — math likes them, money not moving yet ({len(calconly)}) · breakout watchlist"):
            st.markdown("".join(_combo_row(i + 1, c, _scan_short, _now_map) for i, c in enumerate(calconly[:20])),
                        unsafe_allow_html=True)
    with st.expander("📋 Full combo board (sortable) + download"):
        disp = pd.DataFrame([{"Stock": c["name"], "Symbol": c["sym"], "Price": c["price"],
                              "Day%": c["chg_day"], "ClimbScore": c["climb"], "GreenCandles%": c["green"],
                              "Slope1h%": c["slope1h"], "Vol×": c["vr"], "Signal": c["sig"],
                              "Trend": c["dtr"], "Conf%": c["conf"], "200EMA": "Above" if c["above200"] else "Below",
                              "ComboScore": c["combo"], "Verdict": c["verdict"],
                              "ScanTime": _ts, "Buy@": c.get("buy_at"), "SL": c.get("sl"),
                              "SellT1": c.get("t1"), "SellT2": c.get("t2"),
                              "Support": c.get("support"), "Resist": c.get("resistance"),
                              "ResDist%": c.get("res_dist"),
                              "SlowSteady": "🐢" if c["steady"] else ""} for c in combos])
        try:
            st.dataframe(disp, **STRETCH, height=420, hide_index=True)
        except Exception:
            st.dataframe(disp, **STRETCH)
        try:
            st.download_button("⬇️ Download combo board (CSV)", data=disp.to_csv(index=False).encode(),
                               file_name=f"combo_picks_{now_ist().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv", **STRETCH, key="cb_csv")
        except Exception:
            pass


def dashboard_tab(ss, mst_s, ml, mm):
    MONO = "ui-monospace,Menlo,Consolas,monospace"

    # ♾️ AUTO-RESUME — the terminal keeps running across page refreshes
    if not ss.get("dash_watch"):
        _rt = rt_load().get("dash") or {}
        if _rt.get("watch"):
            ss["dash_watch"] = _rt["watch"]
            ss["dash_names"] = _rt.get("names") or {}
            _st = _rt.get("state") or {}
            if _st.get("rows"):
                for r in _st["rows"].values():
                    r.pop("spark", None)
                ss["dash"] = _st
            ss["dash_per"] = _rt.get("per") or "5d"
            if _rt.get("iv"):
                ss["dash_iv"] = _rt["iv"]
            if _rt.get("auto") is not None:
                ss["dash_auto"] = _rt["auto"]
            ss["_dash_resumed"] = True

    # ---------- settings ----------
    with st.expander("⚙️ WATCHLIST · SPEED · REFRESH", expanded=not ss.get("dash_watch")):
        c1, c2 = st.columns(2)
        with c1:
            _keys = list(DASH_SRC.keys())
            _full = _keys.index("🌐 Full NSE (auto-fill to your count)") if "🌐 Full NSE (auto-fill to your count)" in _keys else 0
            src = st.selectbox("Universe", _keys, index=_full, key="dash_src")
            custom_txt = ""
            if DASH_SRC.get(src) == "custom":
                custom_txt = st.text_area("Symbols (comma/space separated · max 500)",
                                          placeholder="RELIANCE, TCS, SUZLON, ZENSARTECH …",
                                          height=80, key="dash_custom")
            n_sel = st.slider("How many stocks (minimum 200 recommended)", 100, 500, 500, 50, key="dash_n",
                              help="Full-NSE fills up to this many. The board GUARANTEES at least "
                                   "min(200, list size) live stocks — missing symbols are auto-retried.")
        with c2:
            dash_iv = st.selectbox("Timeframe", ["5m", "15m", "30m", "1h"], index=1, key="dash_iv")
            st.checkbox("⚡ Light mode (5d history — fastest, recommended for 300–500 stocks)",
                        value=True, key="dash_light")
            st.slider("Stocks refreshed per cycle (stalest first)", 50, 500, 200, 50, key="dash_batch",
                      help="Each cycle re-downloads only this many — the whole board rotates, "
                           "so it stays live without hammering Yahoo.")
        b1, b2, b3 = st.columns(3)
        with b1:
            start = st.button("🚀 START TERMINAL", type="primary", **STRETCH, key="dash_start")
        with b2:
            refresh_now = st.button("🔄 Refresh a batch now", **STRETCH, key="dash_refresh_btn")
        with b3:
            wipe = st.button("🧹 Clear board", **STRETCH, key="dash_wipe")

    if wipe:
        ss["dash_watch"] = None; ss["dash"] = None
        rt_clear("dash")

    if start:
        watch, names = build_watchlist(src, n_sel, custom_txt)
        ss["dash_watch"] = watch
        ss["dash_names"] = names
        ss["dash"] = None
        ss["dash_per"] = "5d" if ss.get("dash_light", True) else DASH_PER.get(ss.get("dash_iv", "15m"), "1mo")

    watch = ss.get("dash_watch")
    if not watch:
        st.markdown("<div style='background:#0b1220;border-radius:20px;padding:56px;text-align:center;'>"
                    "<div style='font-size:52px;'>🔴</div>"
                    "<div style='font-size:22px;font-weight:900;color:#f1f5f9;margin-top:10px;'>"
                    "LIVE MARKET TERMINAL</div>"
                    "<div style='color:#64748b;font-size:13px;margin-top:8px;'>"
                    "Up to 500 NSE stocks · guaranteed 200+ live · uptrend leaderboard · flip alerts · "
                    "who JUST turned uptrend · NIFTY pulse</div>"
                    "<div style='color:#94a3b8;font-size:12px;margin-top:14px;'>↑ Choose a universe and press "
                    "<b style='color:#22c55e;'>🚀 START TERMINAL</b></div></div>", unsafe_allow_html=True)
        return

    if ss.get("_dash_resumed"):
        ss["_dash_resumed"] = False
        st.info("♾️ Terminal resumed automatically — a page refresh does NOT stop it. "
                "Press 🧹 Clear board to end the session.")

    iv = ss.get("dash_iv", "15m")
    per = ss.get("dash_per") or "5d"
    batch = ss.get("dash_batch", 200)

    if start or not ss.get("dash"):
        with st.spinner("🔴 First sweep — building the live board (guaranteeing 200+)…"):
            dash_refresh(ss, watch, iv, per, batch=len(watch), progress=True,
                         ensure_min=min(200, len(watch)))
    if refresh_now or ss.get("dash_auto"):
        dash_refresh(ss, watch, iv, per, batch=batch,
                     ensure_min=min(200, len(watch)) if len(ss.get("dash", {}).get("rows", {})) < 200 else 0)

    # ♾️ persist the running terminal — survives page refresh, resumes anywhere
    try:
        _st8 = ss.get("dash") or {}
        _rows8 = {k: {kk: vv for kk, vv in v.items() if kk != "spark"}
                  for k, v in (_st8.get("rows") or {}).items()}
        rt_save("dash", watch=watch, names=ss.get("dash_names") or {}, per=per, iv=iv,
                auto=bool(ss.get("dash_auto")), state={**_st8, "rows": _rows8})
    except Exception:
        pass

    state = ss.get("dash") or {"rows": {}, "alerts": []}
    rows = state.get("rows", {})
    data = list(rows.values())
    holes = len(watch) - len(rows)
    now = time.time()
    ages = [now - r["ts"] for r in data]
    avg_age = (sum(ages) / len(ages)) if ages else 0

    # ================= PANEL 1 — TERMINAL HEADER =================
    idx = fetch_indices()
    def _idx_cell(lbl, d):
        if not d:
            return f"<div style='text-align:center;opacity:.45;'><div style='color:#64748b;font-size:9px;font-weight:800;letter-spacing:1px;'>{lbl}</div><div style='color:#64748b;font-size:12px;'>—</div></div>"
        c = "#22c55e" if d["chg"] >= 0 else "#ef4444"
        return (f"<div style='text-align:center;'><div style='color:#64748b;font-size:9px;font-weight:800;letter-spacing:1px;'>{lbl}</div>"
                f"<div style='color:#f1f5f9;font-weight:900;font-size:15px;font-family:{MONO};'>{d['price']:,.1f}</div>"
                f"<div style='color:{c};font-size:11px;font-weight:800;font-family:{MONO};'>{d['chg']:+.2f}%</div></div>")
    live_c = "#22c55e" if len(rows) >= min(200, len(watch)) else "#f59e0b"
    st.markdown(f"""
    <div style='background:linear-gradient(135deg,#0b1220,#0f172a);border:1px solid #1e293b;border-radius:18px;
    padding:16px 22px;margin-bottom:12px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:14px;'>
      <div><div style='display:flex;align-items:center;gap:9px;'>
        <span style='width:10px;height:10px;border-radius:50%;background:#22c55e;display:inline-block;
        box-shadow:0 0 10px #22c55e;'></span>
        <span style='color:#f1f5f9;font-size:19px;font-weight:900;letter-spacing:.5px;'>LIVE MARKET TERMINAL</span></div>
        <div style='color:#64748b;font-size:11px;margin-top:3px;font-family:{MONO};'>
        {len(rows):,}/{len(watch):,} stocks live · avg age {_fmt_age(avg_age)} · tf {iv} ·
        {'light' if per=='5d' else 'full'} mode · {datetime.now().strftime('%d %b %H:%M:%S')}</div></div>
      <div style='display:flex;gap:18px;align-items:center;flex-wrap:wrap;'>
        {_idx_cell('NIFTY 50', idx.get('NIFTY 50'))}
        {_idx_cell('BANK NIFTY', idx.get('BANK NIFTY'))}
        <div style='text-align:center;'><div style='color:#64748b;font-size:9px;font-weight:800;letter-spacing:1px;'>BOARD</div>
        <div style='color:{live_c};font-size:13px;font-weight:900;font-family:{MONO};'>
        {'✓ 200+ GUARANTEED' if len(rows)>=200 else f'{len(rows)} live' + (' · filling…' if holes>0 else '')}</div>
        <div style='color:#64748b;font-size:10px;'>{ml}</div></div>
      </div></div>""", unsafe_allow_html=True)

    if not data:
        st.info("No data yet — press 🔄 Refresh a batch now.")
        return

    # live controls row
    q1, q2 = st.columns([1, 1])
    with q1:
        st.checkbox("🔴 LIVE auto-refresh (keeps scanning — late uptrends get caught)", value=False, key="dash_auto")
    with q2:
        st.selectbox("Every", ["1 min", "2 min", "3 min", "5 min"], index=1, key="dash_int")

    # ── 🧠 DAY MEMORY — ONE snapshot/day (today's board), verified next day ──
    try:
        _st = ss.get("dash") or {}
        if rows and time.time() - _st.get("last_snap", 0) > 600:
            if snaps_save(rows):
                ss["dash"]["last_snap"] = time.time()
    except Exception:
        pass
    mc1, mc2 = st.columns([3, 1])
    with mc1:
        _ls = (ss.get("dash") or {}).get("last_snap", 0)
        st.markdown("<div style='background:#0f172a;border:1px dashed #1e293b;border-radius:10px;padding:7px 12px;"
                    "color:#64748b;font-size:11px;'>🧠 <b style='color:#94a3b8;'>Day memory (1/day):</b> "
                    f"today's board last saved {datetime.fromtimestamp(_ls, TZ_IST).strftime('%d %b %H:%M') if _ls else '— (saves automatically once the board is live)'}"
                    f" · 1 snapshot per day · memory: {_ukey()} · verify next day in the 🌙 EOD Review tab</div>", unsafe_allow_html=True)
    with mc2:
        if st.button("📸 Snapshot now", key="dash_snap", **STRETCH):
            if snaps_save(rows):
                ss["dash"]["last_snap"] = time.time()
                st.success("✅ Today's calculation saved — verify it tomorrow in 🌙 EOD Review.")
            else:
                st.error("Couldn't write snapshot file.")

    # ================= PANEL 2 — KPI STRIP =================
    n_up = sum(1 for r in data if r["dtr"] == "UPTREND")
    n_dn = sum(1 for r in data if r["dtr"] == "DOWNTREND")
    n_sd = len(data) - n_up - n_dn
    sbuy = sum(1 for r in data if r["sig"] in ("STRONG BUY", "BUY"))
    ssell = sum(1 for r in data if r["sig"] in ("STRONG SELL", "SELL"))
    adv = sum(1 for r in data if r["chg"] > 0); dec = sum(1 for r in data if r["chg"] < 0)
    tot = len(data)
    pu = adv / tot * 100; pd_ = dec / tot * 100; pf = 100 - pu - pd_
    kpis = (_kpi("📈 UPTREND", n_up, f"{n_up/tot*100:.0f}% of board", "#22c55e")
            + _kpi("📉 DOWNTREND", n_dn, f"{n_dn/tot*100:.0f}% of board", "#ef4444")
            + _kpi("➡️ SIDEWAYS", n_sd, f"{n_sd/tot*100:.0f}% of board", "#94a3b8")
            + _kpi("🟢 BUY SIGNALS", sbuy, f"strong: {sum(1 for r in data if r['sig']=='STRONG BUY')}", "#4ade80")
            + _kpi("🔴 SELL SIGNALS", ssell, f"strong: {sum(1 for r in data if r['sig']=='STRONG SELL')}", "#f87171"))
    st.markdown(f"<div style='display:flex;gap:10px;flex-wrap:wrap;margin-bottom:10px;'>{kpis}</div>", unsafe_allow_html=True)
    st.markdown(f"""<div style='background:#0f172a;border:1px solid #1e293b;border-radius:14px;padding:12px 16px;margin-bottom:12px;'>
    <div style='display:flex;height:14px;border-radius:7px;overflow:hidden;'>
    <div style='background:#16a34a;width:{pu:.1f}%;'></div><div style='background:#334155;width:{pf:.1f}%;'></div>
    <div style='background:#dc2626;width:{pd_:.1f}%;'></div></div>
    <div style='display:flex;justify-content:space-between;color:#64748b;font-size:11px;font-weight:700;margin-top:6px;font-family:{MONO};'>
    <span style='color:#22c55e;'>▲ {adv} advancing ({pu:.0f}%)</span>
    <span>MARKET BREADTH — TODAY</span>
    <span style='color:#ef4444;'>▼ {dec} declining ({pd_:.0f}%)</span></div></div>""", unsafe_allow_html=True)

    # ================= PANEL 3 — FLIP ALERTS =================
    alerts = state.get("alerts", [])
    if alerts:
        chips = "".join(f"<span style='background:rgba(245,158,11,.12);border:1px solid rgba(245,158,11,.35);color:#fbbf24;"
                        f"font-size:11px;font-weight:700;padding:4px 10px;border-radius:8px;margin:3px 4px 3px 0;display:inline-block;'>"
                        f"⏱ {a['ts']} · <b>{a['name']}</b> · {a['txt']}</span>" for a in alerts[:14])
        st.markdown(f"<div style='background:#0f172a;border:1px solid #1e293b;border-radius:14px;padding:12px 16px;margin-bottom:12px;'>"
                    f"<div style='color:#f59e0b;font-size:11px;font-weight:900;letter-spacing:1px;margin-bottom:6px;'>🔔 LIVE FLIP ALERTS — signals & trends changing now</div>"
                    f"{chips}</div>", unsafe_allow_html=True)

    # ================= PANEL 4 — UPTREND LEADERBOARD =================
    ups = sorted([r for r in data if r["dtr"] == "UPTREND"], key=lambda x: -x["score"])
    st.markdown(f"""<div style='background:#0b1220;border:1px solid #14532d;border-radius:18px;padding:16px 18px;margin-bottom:12px;'>
    <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;'>
    <div style='color:#4ade80;font-size:15px;font-weight:900;letter-spacing:.5px;'>🚀 UPTREND LEADERBOARD
    <span style='color:#64748b;font-size:12px;font-weight:600;'>· {len(ups)} stocks rising now · ranked by score</span></div>
    <div style='color:#64748b;font-size:10px;font-family:{MONO};'>CONF bar · SCORE bar · 15m sparkline</div></div></div>""",
                unsafe_allow_html=True)
    if not ups:
        st.markdown("<div style='background:#0f172a;border:1px solid #1e293b;border-radius:14px;padding:18px;color:#94a3b8;"
                    "font-size:13px;'>No clean daily uptrend on this board right now — that's information too. "
                    "Keep 🔴 LIVE on; the 🆕 panel below will flag the moment one turns.</div>", unsafe_allow_html=True)
    else:
        top_show = ups[:40]
        rows_html = "".join(_lb_row(i + 1, r, now) for i, r in enumerate(top_show))
        st.markdown(rows_html, unsafe_allow_html=True)
        if len(ups) > 40:
            chips = "".join(f"<span style='background:rgba(34,197,94,.10);border:1px solid rgba(34,197,94,.3);color:#86efac;"
                            f"font-size:11px;font-weight:700;padding:3px 10px;border-radius:14px;margin:2px 3px 2px 0;display:inline-block;'>"
                            f"{r['name']} {r['chg']:+.1f}%</span>" for r in ups[40:120])
            st.markdown(f"<div style='margin-top:6px;'>{chips}</div>", unsafe_allow_html=True)
        # analyze buttons for top 12
        st.markdown("<div style='color:#64748b;font-size:11px;font-weight:800;margin:8px 0 6px;'>🔍 OPEN FULL ANALYSIS — top setups:</div>",
                    unsafe_allow_html=True)
        show_btn = ups[:12]
        for rr in range(0, len(show_btn), 4):
            cols = st.columns(4)
            for cc_, r in zip(cols, show_btn[rr:rr + 4]):
                with cc_:
                    if st.button(f"📊 {r['name'][:14]}", key=f"dg_{r['sym']}", **STRETCH):
                        ss["sym"] = r["sym"]; ss["stock_name"] = r["name"]; ss["analyzed"] = True
                        st.success(f"✅ {r['name']} loaded — open the '📊 Analyze Stock' tab.")

    # ================= PANEL 5 — JUST TURNED UPTREND =================
    fresh = sorted([r for r in data if r.get("up_since") and now - r["up_since"] < 3600],
                   key=lambda x: -x["score"])
    if fresh:
        chips = "".join(f"<span style='background:rgba(34,197,94,.14);border:1px solid #22c55e;color:#4ade80;"
                        f"font-size:12px;font-weight:800;padding:5px 12px;border-radius:9px;margin:3px 4px 3px 0;display:inline-block;'>"
                        f"🆕 {r['name']} · since {datetime.fromtimestamp(r['up_since']).strftime('%H:%M')} · {r['chg']:+.1f}% · score {r['score']:.0f}</span>"
                        for r in fresh[:25])
        st.markdown(f"<div style='background:#0f172a;border:1px solid #14532d;border-radius:14px;padding:12px 16px;margin-bottom:12px;'>"
                    f"<div style='color:#4ade80;font-size:11px;font-weight:900;letter-spacing:1px;margin-bottom:6px;'>"
                    f"🆕 JUST TURNED UPTREND (last 60 min) — the late movers you asked for</div>{chips}</div>",
                    unsafe_allow_html=True)

    # ================= PANEL 6 — DOWNTREND =================
    dns = sorted([r for r in data if r["dtr"] == "DOWNTREND"], key=lambda x: x["sp"], reverse=True)
    if dns:
        chips = "".join(f"<span style='background:rgba(239,68,68,.10);border:1px solid rgba(239,68,68,.35);color:#f87171;"
                        f"font-size:11px;font-weight:700;padding:3px 10px;border-radius:14px;margin:2px 3px 2px 0;display:inline-block;'>"
                        f"{r['name']} {r['chg']:+.1f}%</span>" for r in dns[:40])
        st.markdown(f"""<div style='background:#0b1220;border:1px solid #7f1d1d;border-radius:16px;padding:14px 16px;margin-bottom:12px;'>
        <div style='color:#f87171;font-size:13px;font-weight:900;margin-bottom:6px;'>📉 DOWNTREND — {len(dns)} stocks
        <span style='color:#64748b;font-size:11px;font-weight:600;'>· avoid longs · top 40 shown</span></div>{chips}</div>""",
                    unsafe_allow_html=True)

    # ================= PANEL 7 — SECTOR BREADTH =================
    sec_chips = ""
    for sn, sdict in SECTORS_FOR_BREADTH.items():
        mem = [rows[s] for s in sdict.values() if s in rows]
        if not mem:
            continue
        upn = sum(1 for r in mem if r["dtr"] == "UPTREND")
        pctu = upn / len(mem) * 100
        clr = "#4ade80" if pctu >= 60 else "#fbbf24" if pctu >= 40 else "#f87171"
        sec_chips += (f"<span style='background:#0f172a;border:1px solid #1e293b;color:{clr};font-size:11px;"
                      f"font-weight:800;padding:4px 12px;border-radius:9px;margin:3px 4px 3px 0;display:inline-block;'>"
                      f"{sn} · {upn}/{len(mem)} up ({pctu:.0f}%)</span>")
    if sec_chips:
        st.markdown(f"<div style='background:#0b1220;border:1px solid #1e293b;border-radius:14px;padding:12px 16px;margin-bottom:12px;'>"
                    f"<div style='color:#94a3b8;font-size:11px;font-weight:900;letter-spacing:1px;margin-bottom:6px;'>🏭 SECTOR BREADTH — who is in an uptrend</div>"
                    f"{sec_chips}</div>", unsafe_allow_html=True)

    # ================= DEEP-DIVE TABS (top 6) =================
    top6 = ups[:6] if ups else []
    if top6:
        st.markdown('<div class="sh">🔍 DEEP DIVE — TOP 6 UPTREND SETUPS</div>', unsafe_allow_html=True)
        tbs = st.tabs([f"{r['name'][:11]} · {r['conf']:.0f}%" for r in top6])
        for tb, r in zip(tbs, top6):
            with tb:
                k1, k2, k3, k4 = st.columns(4)
                with k1: st.metric("Price", f"₹{r['price']:,.2f}", f"{r['chg']:+.2f}% today")
                with k2: st.metric("Signal", r["sig"], f"{r['conf']:.0f}% conf")
                with k3: st.metric("Daily trend", "🟢 UPTREND", "score " + f"{r['score']:.0f}")
                with k4: st.metric("200 EMA", "Above ⬆" if r["above200"] else "Below ⬇",
                                   (f"₹{r['ema200']:,.2f}" if r.get("ema200") else "—"))
                try:
                    st.plotly_chart(spark_fig(r), **STRETCH)
                except Exception:
                    pass
                st.markdown(f"<div class='sc-r'><b style='color:{r['sc']};'>{r['act']}</b><div style='color:#374151;"
                            f"font-size:12px;margin-top:6px;'>Entry ₹{r['buy_at']:.2f} · Stop ₹{r['sl']:.2f} · "
                            f"T1 ₹{r['t1']:.2f} · T2 ₹{r['t2']:.2f} · R:R 1:{r['rr']} · intraday: {r['itrend'].title()} · "
                            f"RSI {r['rsi']:.0f} · vol {r['vr']:.1f}x avg</div>"
                            f"<div style='color:#9ca3af;font-size:11px;margin-top:4px;'>Quick plan from the live board — the full "
                            f"Analyze tab adds pivots, Fibonacci, circuits, news & ML.</div></div>", unsafe_allow_html=True)
                if st.button(f"📊 Open FULL analysis → {r['name'][:14]}", key=f"dt_{r['sym']}",
                             **STRETCH, type="primary"):
                    ss["sym"] = r["sym"]; ss["stock_name"] = r["name"]; ss["analyzed"] = True
                    st.success("Loaded — open the '📊 Analyze Stock' tab.")

    # ================= DATA TABLE + CSV =================
    with st.expander("📋 Full data table · CSV export · load any stock"):
        def _n(v, nd=1):
            try:
                return None if (v is None or (isinstance(v, float) and np.isnan(v))) else round(float(v), nd)
            except Exception:
                return None
        disp_rows = [{
            "Stock": r["name"], "Symbol": r["sym"], "Price": round(r["price"], 2),
            "% Today": r["chg"], "Signal": r["sig"], "Conf": r["conf"],
            "Buy%": r["bp"], "Sell%": r["sp"], "D-Trend": r["dtr"],
            "200EMA": ("Above" if r["above200"] else "Below") if r["above200"] is not None else "—",
            "I-Momentum": r["itrend"].title(), "RSI": _n(r["rsi"]), "ADX": _n(r["adx"]),
            "Vol×": r["vr"], "Score": r["score"], "Buy@": r["buy_at"], "SL": r["sl"],
            "T2": r["t2"], "R:R": r["rr"], "AgeS": _n(now - r["ts"], 0),
        } for r in sorted(data, key=lambda x: -x["score"])]
        disp = pd.DataFrame(disp_rows)
        try:
            st.dataframe(disp, **STRETCH, height=440, hide_index=True)
        except Exception:
            st.dataframe(disp, **STRETCH)
        lc1, lc2 = st.columns([2, 1])
        with lc1:
            pick = st.text_input("Load any board stock into the Analyze tab — type its exact name (e.g. HAL)",
                                 placeholder="Exact name from the table…", key="dash_pick")
        with lc2:
            st.markdown("<div style='height:10px;'></div>", unsafe_allow_html=True)
            if st.button("📊 Load into Analyze tab", key="dash_load", **STRETCH):
                hit = [r for r in data if r["name"].upper() == (pick or "").strip().upper()]
                if hit:
                    ss["sym"] = hit[0]["sym"]; ss["stock_name"] = hit[0]["name"]; ss["analyzed"] = True
                    st.success(f"✅ {hit[0]['name']} loaded — open the '📊 Analyze Stock' tab.")
                else:
                    st.warning("Type the exact stock name as shown in the table (e.g. HAL).")
        try:
            st.download_button("⬇️ Download board as CSV", data=disp.to_csv(index=False).encode(),
                               file_name=f"live_board_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv", **STRETCH)
        except Exception:
            pass

    st.markdown("<div style='color:#64748b;font-size:11px;margin-top:8px;'>⚠️ Live board = live data, but signals are "
                "probabilities, not predictions. Open the full analysis (and always use a stop) before trading. "
                "Educational use only.</div>", unsafe_allow_html=True)

    # ---- auto-refresh hook ----
    if ss.get("dash_auto"):
        secs = {"1 min": 60, "2 min": 120, "3 min": 180, "5 min": 300}[ss.get("dash_int", "2 min")]
        _smooth = False
        try:
            from streamlit_autorefresh import st_autorefresh
            st_autorefresh(interval=secs * 1000, key="dash_tick")
            _smooth = True
        except Exception:
            _smooth = False
        if not _smooth:
            st.info(f"🔴 LIVE (basic mode) — refreshing the stalest batch every {secs}s. "
                    "For perfectly smooth auto-refresh:  pip install streamlit-autorefresh")
            try:
                time.sleep(secs)
            except Exception:
                pass
            try:
                st.rerun()
            except RuntimeError:
                pass



# ============================================================
# 🌙 EOD REVIEW TAB — before vs after (did the board get it right?)
# ============================================================
def eod_review_tab(ss):
    st.markdown("""<div style='background:linear-gradient(135deg,#1e1b4b,#312e81);border-radius:18px;
    padding:20px 24px;margin-bottom:16px;'>
    <div style='color:white;font-size:20px;font-weight:900;'>🌙 EOD REVIEW — before vs after</div>
    <div style='color:#c7d2fe;font-size:13px;margin-top:6px;'>The 🔴 dashboard saves a <b>memory snapshot</b>
    automatically — ONE snapshot per day (that day's calculation) — exactly what it showed you. Pick a day below and press
    VERIFY: the app fetches what prices <b>ACTUALLY did afterwards</b> and scores every call — how many stocks went
    as per the calculation, how many hit T1/T2, how many hit the stop. This is your real accuracy record.</div></div>""",
                unsafe_allow_html=True)
    _uk = _ukey()
    _shared = _uk == "main"
    uc1, uc2, uc3 = st.columns([2.3, 1.4, 1])
    with uc1:
        st.markdown("<div style='background:" + ("#3f2d04;border:1px solid #f59e0b" if _shared else "#04260f;border:1px solid #16a34a") +
                    ";border-radius:10px;padding:8px 12px;color:#d1d5db;font-size:11.5px;'>🧑‍🤝‍🧑 Your memory name: <b style=\"" +
                    ("#fbbf24" if _shared else "#4ade80") + "\">" + _uk + "</b>" +
                    (" — <b>shared default.</b> If a friend also uses this app, each of you must set your own name here "
                     "(any word). Then memories, scorecards and journals can never mix." if _shared else
                     " — your own separate memory ✓ (snapshots, scorecard & journal are yours alone)") + "</div>",
                    unsafe_allow_html=True)
    with uc2:
        _nm = st.text_input("memory name", value="", key="mu_name", placeholder="type a name… e.g. ravi",
                            label_visibility="collapsed")
    with uc3:
        st.markdown("<div style='height:2px;'></div>", unsafe_allow_html=True)
        if st.button("🔗 Use this name", key="mu_go", **STRETCH):
            _clean = "".join(ch for ch in _nm.strip().lower() if ch.isalnum() or ch in "_-")[:16]
            if _clean:
                st.query_params["u"] = _clean
                st.rerun()
            else:
                st.error("Type a name first (letters/numbers).")

    with st.expander("💾 Where is my data saved? · Backup & Restore (Streamlit Cloud)"):
        st.markdown("""<div style='background:#0f172a;border:1px solid #1e293b;border-radius:12px;padding:12px 16px;
        color:#94a3b8;font-size:12px;line-height:1.8;'>
        ● <b style='color:#e2e8f0;'>On your own PC:</b> memory saves automatically in the same folder as this app —
        one set of files <b>per person</b> (<code>board_snapshots_NAME.json</code>, <code>eod_results_NAME.json</code>,
        <code>trade_journal_NAME.json</code>; NAME is your memory name, "main" if not set) — permanent, nothing to upload.<br>
        ● <b style='color:#e2e8f0;'>On Streamlit Cloud:</b> the files live inside the app's private cloud container —
        they survive while the app runs/sleeps, but are <b style='color:#f87171;'>wiped whenever the app redeploys</b>
        (each GitHub update). VERIFY always works (it fetches live market prices by itself) — only past days' memory
        needs restoring: press ⬇️ Backup before you update the app, and ⬆️ Restore after.</div>""",
                    unsafe_allow_html=True)
        _bk = {"snapshots": snaps_load(), "eod": eod_load(),
               "exported": datetime.now().strftime("%Y-%m-%d %H:%M")}
        st.download_button("⬇️ Backup memory (download .json)", data=_json.dumps(_bk).encode(),
                           file_name=f"ai_trader_memory_{_ukey()}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                           mime="application/json", **STRETCH, key="eod_backup_dl")
        _up = st.file_uploader("⬆️ Restore memory (.json backup)", type=["json"], key="eod_restore")
        if _up is not None:
            try:
                _d = _json.loads(_up.read().decode("utf-8"))
                _news = _newe = 0
                _olds = snaps_load(); _ids = {s.get("id") for s in _olds}
                _srcs = _d.get("snapshots") if isinstance(_d, dict) else (_d if isinstance(_d, list) else [])
                _add = [s for s in _srcs if isinstance(s, dict) and s.get("id") and s.get("id") not in _ids and s.get("rows")]
                if _add:
                    _json.dump(sorted(_olds + _add, key=lambda x: x.get("saved", ""))[-40:],
                               open(snap_file(), "w", encoding="utf-8"))
                    _news = len(_add)
                if isinstance(_d, dict) and _d.get("eod"):
                    _olde = eod_load(); _eids = {x.get("id") for x in _olde}
                    _adde = [x for x in _d["eod"] if isinstance(x, dict) and x.get("id") not in _eids]
                    if _adde:
                        _json.dump(sorted(_olde + _adde, key=lambda x: x.get("date", ""))[-100:],
                                   open(eod_file(), "w", encoding="utf-8"), indent=2)
                        _newe = len(_adde)
                if _news or _newe:
                    st.success(f"✅ Restored {_news} day-snapshots + {_newe} scorecard entries — pick a day below.")
                else:
                    st.info("Nothing new in that backup (I already have it all).")
            except Exception as _e:
                st.error(f"Couldn't read that backup: {_e}")

    snaps = snaps_load()
    if not snaps:
        st.info("🧠 No board memory yet. Open the 🔴 Live Dashboard tab → START TERMINAL → let it run "
                "(a snapshot saves automatically within seconds of the board going live) → come back here.")
        return
    snaps = sorted(snaps, key=lambda s: s.get("saved", ""), reverse=True)
    _tdy = datetime.now().strftime("%Y-%m-%d")
    labels = [f"{s['saved']}  ·  {len(s['rows'])} stocks  ·  "
              + ("✅ full results ready" if s['saved'][:10] < _tdy else "⏳ today · partial so far")
              for s in snaps]
    _defi = next((i for i, s in enumerate(snaps) if s['saved'][:10] < _tdy), 0)
    c1, c2 = st.columns([3, 1])
    with c1:
        si = st.selectbox("WHICH DAY'S calculation (the ‘before’) should I verify?", labels, index=_defi, key="eod_snap")
    with c2:
        st.markdown("<div style='height:10px;'></div>", unsafe_allow_html=True)
        go = st.button("▶ VERIFY NOW (the ‘after’)", type="primary", **STRETCH, key="eod_go")
    snap = snaps[labels.index(si)]
    st.caption("How it works: pick YESTERDAY'S calculation → VERIFY → the app itself fetches today's real market "
               "prices (Yahoo) and scores every call — direction accuracy, T1/T2 hits, stops. No manual file work: "
               "comparing with the market is always automatic.")
    if go:
        with st.spinner(f"🔍 Fetching what actually happened after {snap['saved']} …"):
            try:
                results, summary = verify_snapshot(snap)
                ss["eod_cache"] = {"id": snap["id"], "results": results, "summary": summary}
                if summary.get("n_valid", 0) >= max(10, int(summary["checked"] * 0.2)):
                    eod_save({"id": snap["id"], "date": snap["saved"][:10],
                              **{k: summary[k] for k in ("checked", "n_up", "up_ok", "up_acc",
                                                          "n_buy", "n_t1", "n_sl", "buy_wr")}})
            except Exception as e:
                st.error(f"Verification failed: {type(e).__name__}: {e}")
    cached = ss.get("eod_cache")
    if cached and cached.get("id") == snap.get("id"):
        s = cached["summary"]; res = cached["results"]
        st.markdown(f"<div style='background:#0f172a;border:1px solid #1e293b;border-radius:14px;padding:12px 16px;"
                    f"margin-bottom:12px;color:#94a3b8;font-size:12px;'>🔎 Verified snapshot <b>{s['snap']}</b> "
                    f"· {s['checked']} stocks checked against real prices after that time</div>", unsafe_allow_html=True)
        if s.get("n_valid", 0) == 0:
            st.markdown(f"""<div style='background:#3f2d04;border:1px solid #f59e0b;border-radius:12px;
            padding:14px 18px;margin-bottom:12px;color:#fde68a;font-size:13px;line-height:1.8;'>
            <b style='color:#fbbf24;font-size:15px;'>⚠️ Nothing could be verified yet — this is NOT 0% accuracy.</b><br>
            Every stock came back <b>NO DATA</b>: the app could not download any prices after this snapshot
            right now. Usual reasons:<br>
            1. You verified <b>before the next trading session finished</b> — verify after <b>~3:30 PM IST</b>
            on a trading day.<br>
            2. The snapshot was saved <b>after market close or on a weekend/holiday</b> — its "next day"
            hasn't traded yet.<br>
            3. Yahoo didn't answer from the server just then (it throttles cloud servers) — simply
            <b>press VERIFY again</b>.<br>
            Latest market candle the app can see right now:
            <b>{s.get('last_candle') or 'unknown — Yahoo not answering, press VERIFY again'}</b>.
            Nothing was saved to the scorecard for this attempt.</div>""", unsafe_allow_html=True)
        elif s.get("n_valid", 0) < s["checked"] * 0.5:
            st.warning(f"⚠️ Only {s['n_valid']}/{s['checked']} stocks could be verified — Yahoo throttling. "
                       "Press VERIFY again to fill in the rest before trusting these numbers.")
        _acc_txt = f"{s['up_acc']:.0f}%" if s.get("up_acc") is not None else "—"
        _wr_txt = f"{s['buy_wr']:.0f}%" if s.get("buy_wr") is not None else "—"
        m1, m2, m3, m4, m5, m6 = st.columns(6)
        with m1: st.metric("Stocks checked", s["checked"], f"{s.get('n_valid', 0)} with real data")
        with m2: st.metric("📈 Uptrend calls", s["n_up"], f"avg move {s['avg_up']:+.2f}%")
        with m3: st.metric("Uptrend accuracy", _acc_txt, f"{s['up_ok']}/{s['n_up']} closed higher")
        with m4: st.metric("🔴 Downtrend avg", f"{s['avg_dn']:+.2f}%", "avoid-long list")
        with m5: st.metric("BUY signals", s["n_buy"], f"{s['n_t1']} hit T1/T2 · {s['n_sl']} hit SL")
        with m6: st.metric("Buy win-rate", _wr_txt, "of resolved (T1 vs SL)")
        acc = max(0.0, min(100.0, s["up_acc"])) if s.get("up_acc") is not None else None
        if acc is not None:
            st.markdown(f"""<div style='background:#0f172a;border:1px solid #1e293b;border-radius:12px;padding:12px 16px;margin-bottom:12px;'>
            <div style='display:flex;height:14px;border-radius:7px;overflow:hidden;'>
            <div style='background:#16a34a;width:{acc:.1f}%;'></div>
            <div style='background:#334155;width:{100-acc:.1f}%;'></div></div>
            <div style='display:flex;justify-content:space-between;color:#64748b;font-size:11px;font-weight:700;margin-top:6px;'>
            <span style='color:#22c55e;'>▲ {s['up_ok']} calls right</span>
            <span>BOARD DIRECTION ACCURACY</span>
            <span style='color:#ef4444;'>▼ {s['n_up'] - s['up_ok']} calls wrong</span></div></div>""",
                        unsafe_allow_html=True)
        okr = sorted([x for x in res if isinstance(x.get("moved"), (int, float))], key=lambda x: -x["moved"])
        if okr:
            best = "".join(f"<span style='background:rgba(34,197,94,.14);border:1px solid #22c55e;color:#4ade80;"
                           f"font-size:11px;font-weight:800;padding:4px 10px;border-radius:9px;margin:3px 4px 3px 0;display:inline-block;'>"
                           f"👍 {x['name']} {x['moved']:+.1f}%</span>" for x in okr[:10])
            worst = "".join(f"<span style='background:rgba(239,68,68,.12);border:1px solid rgba(239,68,68,.4);color:#f87171;"
                            f"font-size:11px;font-weight:800;padding:4px 10px;border-radius:9px;margin:3px 4px 3px 0;display:inline-block;'>"
                            f"👎 {x['name']} {x['moved']:+.1f}%</span>" for x in okr[-10:][::-1])
            st.markdown(f"<div style='background:#0b1220;border:1px solid #1e293b;border-radius:14px;padding:12px 16px;"
                        f"margin-bottom:12px;'><div style='color:#94a3b8;font-size:11px;font-weight:900;margin-bottom:6px;'>"
                        f"🎯 BEST calls after this snapshot</div>{best}"
                        f"<div style='color:#94a3b8;font-size:11px;font-weight:900;margin:10px 0 6px;'>💀 WORST calls</div>{worst}</div>",
                        unsafe_allow_html=True)
        disp = pd.DataFrame([{
            "Stock": x.get("name"), "Symbol": x.get("sym"),
            "Signal@snap": x.get("sig"), "Trend@snap": x.get("dtr"),
            "Price@snap": x.get("price"), "After_Last": x.get("last"),
            "Moved%": x.get("moved"), "High": x.get("hi"), "Low": x.get("lo"),
            "T1": x.get("t1"), "T2": x.get("t2"), "SL": x.get("sl"),
            "Outcome": x.get("outcome"), "Score@snap": x.get("score"),
        } for x in res])
        try:
            st.dataframe(disp, **STRETCH, height=420, hide_index=True)
        except Exception:
            st.dataframe(disp, **STRETCH)
        try:
            st.download_button("⬇️ Download EOD review (CSV)", data=disp.to_csv(index=False).encode(),
                               file_name=f"eod_review_{snap['id']}.csv", mime="text/csv",
                               **STRETCH)
        except Exception:
            pass
    hist = eod_load()
    if hist:
        with st.expander(f"🏦 Scorecard history — {len(hist)} past verifications"):
            h = pd.DataFrame([{"Date": d.get("date"), "Stocks": d.get("checked"),
                               "UptrendCalls": d.get("n_up"), "UptrendAcc%": d.get("up_acc"),
                               "Buys": d.get("n_buy"), "T1/T2Wins": d.get("n_t1"),
                               "SLosses": d.get("n_sl"), "BuyWinRate%": d.get("buy_wr")}
                              for d in sorted(hist, key=lambda x: x.get("date", ""), reverse=True)])
            try:
                st.dataframe(h, **STRETCH, hide_index=True)
            except Exception:
                st.dataframe(h, **STRETCH)
            st.caption("Memory files live next to the app, one set per person (memory name) — "
                       "on your PC they persist forever; on Streamlit Cloud they survive until the app redeploys.")


# ============================================================
# SCANNER
# ============================================================
def resolve_category(cat):
    d = CAT_MAP.get(cat)
    if d is None:
        uni = fetch_universe()
        return dict(uni["name_map"])
    return d


def scan_one(args):
    name, sym, iv, per = args
    try:
        res = run_analysis(sym, iv, per)
        if not res:
            return None
        plan = make_plan(res['price'], res['atr'], None, res['sig'])
        return {
            'name': name, 'sym': sym, 'price': res['price'], 'sig': res['sig'], 'sc': res['sc'],
            'conf': res['conf'], 'bp': res['bp'], 'sp': res['sp'], 'atr': res['atr'],
            'tr': res['trend'], 'vr': res['vr'], 'gap': plan['gap_pct'], 'act': plan['act'],
            'sit': plan['sit'], 'ac': plan['ac'], 'buy_at': plan['buy_at'], 'sl': plan['sl'],
            't1': plan['t1'], 't2': plan['t2'], 't3': plan['t3'], 't2p': plan['t2p'],
            't1p': plan['t1p'], 't3p': plan['t3p'], 'sl_pct': plan['sl_pct'], 'rr': plan['rr'],
        }
    except Exception:
        return None


def run_scan(stocks, iv, per, min_conf, stype, workers=10, cap_n=None, stats_out=None):
    items = list(stocks.items())
    if cap_n:
        items = items[:cap_n]
    args = [(n, s, iv, per) for n, s in items]
    results = []; total = len(args); cnt = [0]

    def _pass(todo, label):
        got = 0
        bar = st.progress(0); stat = st.empty()
        with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
            futs = {ex.submit(scan_one, a): a for a in todo}
            for f in concurrent.futures.as_completed(futs):
                cnt[0] += 1
                bar.progress(min(cnt[0] / total, 1.0))
                stat.markdown(f"<div style='background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;"
                              f"padding:8px 16px;display:inline-block;color:#1d4ed8;font-size:13px;font-weight:600;'>"
                              f"🔍 {label} {min(cnt[0], total)}/{total} · ✅ data OK {len(results)}</div>",
                              unsafe_allow_html=True)
                try:
                    r = f.result(timeout=15)
                    if r:
                        results.append(r); got += 1
                except Exception:
                    pass
        bar.empty(); stat.empty()
        return got

    ok = _pass(args, "Scanning")
    # ── anti-throttle: if Yahoo didn't answer for many stocks, retry them ONCE
    done_syms = {r["sym"] for r in results}
    todo2 = [a for a in args if a[1] not in done_syms]
    retried = 0
    if todo2 and len(todo2) > max(5, int(total * 0.2)):
        cnt[0] = total - len(todo2)
        retried = _pass(todo2[:400], "Retrying failed (Yahoo throttle)")
    if stats_out is not None:
        stats_out["total"] = total
        stats_out["ok"] = len({r["sym"] for r in results})
        stats_out["fails"] = max(total - stats_out["ok"], 0)
    if stype == "BUY":
        results = [r for r in results if r['conf'] >= min_conf and r['bp'] > r['sp'] and r['sit'] not in ('SELL', 'GAP_DN')]
        results.sort(key=lambda x: (x['tr'] == 'UPTREND', x['bp'], x['conf']), reverse=True)
    elif stype == "SELL":
        results = [r for r in results if r['conf'] >= min_conf and r['sp'] > r['bp']]
        results.sort(key=lambda x: x['sp'], reverse=True)
    else:
        results = [r for r in results if r['conf'] >= min_conf]
        results.sort(key=lambda x: x['conf'], reverse=True)
    return results


# ============================================================
# MAIN APP
# ============================================================
def main():
    ss = st.session_state
    defaults = {'analyzed': False, 'sym': 'RELIANCE.NS', 'stock_name': 'RELIANCE',
                'capital': 10000, 'target': 500, 'iv': '15m', 'per': '1mo',
                'scan_results': None, 'search_results': None}
    for k, v in defaults.items():
        if k not in ss:
            ss[k] = v
    mst_s, ml, mm = mkt_status()
    mclr = "#22c55e" if mst_s == "open" else "#f59e0b" if mst_s == "pre" else "#ef4444"

    # live breadth badge for the navbar once the dashboard has data
    dash_state = ss.get("dash") or {}
    drows = dash_state.get("rows", {})
    breadth_badge = ""
    if drows:
        up = sum(1 for r in drows.values() if r["dtr"] == "UPTREND")
        dn = sum(1 for r in drows.values() if r["dtr"] == "DOWNTREND")
        breadth_badge = (f"<div style='background:rgba(255,255,255,0.12);border-radius:10px;padding:8px 16px;'>"
                         f"<div style='color:#fbbf24;font-weight:700;font-size:12px;'>🔴 Live board</div>"
                         f"<div style='color:white;font-weight:900;font-size:16px;'>📈{up} · 📉{dn} · {len(drows)} watched</div></div>")

    st.markdown(f"""<div class='navbar'><div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:10px;'>
    <div><span style='font-size:28px;font-weight:900;color:white;'>💹 AI Trader Pro</span>
    <span style='font-size:14px;color:#93c5fd;margin-left:12px;'>v13.6 · PRO TERMINAL · 500 LIVE</span></div>
    <div style='display:flex;gap:12px;align-items:center;flex-wrap:wrap;'>
    <div style='background:rgba(255,255,255,0.15);border-radius:10px;padding:8px 16px;text-align:center;'>
    <div style='color:{mclr};font-weight:700;font-size:13px;'>{ml}</div><div style='color:#93c5fd;font-size:10px;'>{mm}</div></div>
    {breadth_badge}
    <div style='background:rgba(255,255,255,0.12);border-radius:10px;padding:8px 16px;'>
    <div style='color:#fbbf24;font-weight:700;font-size:12px;'>Focus</div><div style='color:white;font-weight:900;font-size:16px;'>Uptrend + Levels</div></div>
    </div></div></div>""", unsafe_allow_html=True)

    tab_dash, tab_mv, tab_bnc, tab_cb, tab_analyze, tab_scan, tab_search, tab_journal, tab_eod, tab_guide = st.tabs(
        ["🔴 Live Dashboard (500)", "⚡ Live Movers (Now)", "🚀 Uptrend Starting", "🎯 Combo Picks",
         "📊 Analyze Stock", "🔍 Scanner", "🔎 Search Any Stock", "📓 Journal", "🌙 EOD Review", "📚 Trading Guide"])

    # ── TAB 0: LIVE DASHBOARD (the common board) ──
    with tab_dash:
        try:
            dashboard_tab(ss, mst_s, ml, mm)
        except Exception as e:
            st.markdown(f"<div style='background:#fff1f2;border:2px solid #dc2626;border-radius:14px;"
                        f"padding:16px 20px;'><b style='color:#dc2626;'>⚠️ Dashboard hit a snag</b>"
                        f"<div style='color:#374151;font-size:13px;margin-top:6px;'>Usually a temporary "
                        f"data issue — press 🔄 Refresh, reduce the watchlist size, or use ⚡ Light mode. "
                        f"Details: <code>{type(e).__name__}: {e}</code></div></div>", unsafe_allow_html=True)

    # ── TAB: LIVE MOVERS (pure price action — follow the money) ──
    with tab_mv:
        try:
            live_movers_tab(ss, mst_s)
        except Exception as e:
            st.warning(f"⚠️ Live Movers problem: {type(e).__name__}: {e}")

    # ── TAB: UPTREND STARTING (support bounces) ──
    with tab_bnc:
        try:
            bounce_tab(ss, mst_s)
        except Exception as e:
            st.warning(f"⚠️ Bounce radar problem: {type(e).__name__}: {e}")

    # ── TAB: COMBO (live climb + calculation agreement) ──
    with tab_cb:
        try:
            combo_tab(ss, mst_s)
        except Exception as e:
            st.warning(f"⚠️ Combo problem: {type(e).__name__}: {e}")

    # ── TAB 1: ANALYZE ──
    with tab_analyze:
        st.markdown('<div class="input-row">', unsafe_allow_html=True)
        st.markdown("<div style='font-size:15px;font-weight:800;color:#1d4ed8;margin-bottom:14px;'>⚙️ SETTINGS — pick timeframe, then choose a stock &amp; ANALYZE</div>", unsafe_allow_html=True)
        r3, r4 = st.columns(2)
        with r3:
            iv = st.selectbox("⏱️ Timeframe", ["1m", "5m", "15m", "30m", "1h", "1d"], index=2, key="iv_in"); ss.iv = iv
        with r4:
            best_per = {"1m": "5d", "5m": "1mo", "15m": "1mo", "30m": "3mo", "1h": "3mo", "1d": "1y"}
            per_opts = ["5d", "1mo", "3mo", "6mo", "1y"]
            per = st.selectbox("📅 Period", per_opts, index=per_opts.index(best_per.get(iv, "1mo")), key="per_in"); ss.per = per
        st.markdown("---")
        r5, r6, r7, r8 = st.columns([2, 2, 3, 2])
        with r5:
            cat = st.selectbox("📂 Category", list(CAT_MAP.keys()), key="cat_in")
        with r6:
            flt = st.text_input("🔍 Filter", "", placeholder="Type to filter...", key="flt_in").upper()
        with r7:
            cat_stocks = resolve_category(cat)
            filtered = {k: v for k, v in cat_stocks.items() if flt in k.upper()} if flt else cat_stocks
            keys = sorted(filtered.keys())[:800]
            if keys:
                sel = st.selectbox(f"📈 Stock ({len(filtered)})", keys, key="stk_in")
            else:
                st.warning("No match"); sel = None
        with r8:
            st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)
            analyze_btn = st.button("🚀 ANALYZE NOW", **STRETCH, type="primary", key="analyze_btn")

        lr1, lr2, lr3 = st.columns([1, 1, 2])
        with lr1:
            live_on = st.checkbox("🔴 LIVE auto-refresh", value=False, key="live_on")
        with lr2:
            live_secs = {"30 sec": 30, "1 min": 60, "2 min": 120, "5 min": 300}[
                st.selectbox("Every", ["30 sec", "1 min", "2 min", "5 min"], index=1, key="live_secs",
                             disabled=not live_on)]
        with lr3:
            if live_on:
                st.markdown(f"<div style='background:#fff1f2;border:1px solid #fecaca;border-radius:10px;"
                            f"padding:8px 14px;margin-top:2px;color:#dc2626;font-size:12px;font-weight:600;'>"
                            f"🔴 LIVE — refreshing every {live_secs}s. Best during market hours (9:15–3:30).</div>",
                            unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if analyze_btn and sel and filtered:
            ss.sym = filtered[sel]
            ss.stock_name = sel.split(" (")[0].split(" [")[0]
            ss.analyzed = True

        if ss.analyzed and ss.sym:
            if live_on:
                st.markdown(f"<div style='background:#fff1f2;border:1px solid #fecaca;border-radius:10px;padding:8px 16px;margin-bottom:10px;color:#dc2626;font-weight:700;'>🔴 LIVE · {ss.stock_name} · updates every {live_secs}s · {datetime.now().strftime('%H:%M:%S')}</div>", unsafe_allow_html=True)
            try:
                show_analysis(ss.sym, ss.stock_name, iv, per)
            except Exception as e:
                st.markdown(f"<div style='background:#fff1f2;border:2px solid #dc2626;border-radius:14px;"
                            f"padding:18px 22px;'><b style='color:#dc2626;font-size:16px;'>⚠️ Couldn't complete "
                            f"this analysis</b><div style='color:#374151;font-size:13px;margin-top:6px;'>"
                            f"Often a temporary data issue — try again, switch timeframe to 15m + period 1mo, "
                            f"or pick another stock. Details: <code>{type(e).__name__}: {e}</code></div></div>",
                            unsafe_allow_html=True)
            if live_on:
                _did = False
                try:
                    from streamlit_autorefresh import st_autorefresh
                    st_autorefresh(interval=live_secs * 1000, key="live_tick")
                    _did = True
                except Exception:
                    _did = False
                if not _did:
                    try:
                        time.sleep(live_secs)
                    except Exception:
                        pass
                    try:
                        st.rerun()
                    except RuntimeError:
                        pass
        else:
            st.markdown("<div style='background:white;border:2px dashed #bfdbfe;border-radius:20px;padding:60px;text-align:center;margin-top:10px;'><div style='font-size:52px;margin-bottom:16px;'>📊</div><div style='font-size:22px;font-weight:800;color:#1d4ed8;'>Select a stock above and click ANALYZE NOW</div><div style='color:#6b7280;font-size:14px;margin-top:8px;'>Signal · Fibonacci · Camarilla · Woodie · Chart · All 12 indicators · LIVE mode</div></div>", unsafe_allow_html=True)

    # ── TAB 2: SCANNER ──
    with tab_scan:
        st.markdown('<div class="input-row">', unsafe_allow_html=True)
        st.markdown("<div style='font-size:15px;font-weight:800;color:#1d4ed8;margin-bottom:14px;'>⚙️ SCANNER SETTINGS — tip: the 🔴 Live Dashboard scans up to 500 stocks far faster (batched)</div>", unsafe_allow_html=True)
        sc3, sc4, sc5 = st.columns(3)
        with sc3: s_cat = st.selectbox("📂 Category", list(CAT_MAP.keys()), key="sc_cat")
        with sc4: s_sig = st.radio("Signal", ["📈 BUY", "📉 SELL", "🔍 ALL"], key="sc_sig")
        with sc5: s_mc = st.slider("Min Conf %", 55, 90, 65, 5, key="sc_mc")
        s_iv = st.select_slider("Timeframe", ["5m", "15m", "30m", "1h"], value="15m", key="sc_iv")
        st.markdown('</div>', unsafe_allow_html=True)

        scan_stocks = resolve_category(s_cat)
        is_all = CAT_MAP.get(s_cat) is None
        total_n = len(scan_stocks)
        if is_all:
            cap_n = st.slider("How many stocks to scan (from the full NSE list)",
                              100, total_n, min(total_n, 1000), 100, key="sc_capn")
            if st.checkbox(f"🔥 Scan ALL {total_n:,} stocks (slowest — ~{max(1, total_n//180)}–{max(2, total_n//120)} min)",
                           key="sc_all"):
                cap_n = total_n
        else:
            cap_n = None
        b1, b2 = st.columns([1, 3])
        with b1:
            scan_btn = st.button("🚀 START SCAN", **STRETCH, type="primary", key="scan_btn")
        with b2:
            n_scan = cap_n if cap_n else total_n
            eta = max(1, int(n_scan / 160))
            note = f"Will scan <b>{n_scan:,}</b> stocks (~{eta}–{eta*2} min)" if is_all else f"Will scan <b>{total_n}</b> stocks"
            st.markdown(f"<div style='background:#eff6ff;border-radius:10px;padding:10px 16px;margin-top:4px;'><span style='color:#1d4ed8;font-size:13px;'>{note} · uptrend + confirmed setups first · keep this tab open while it runs</span></div>", unsafe_allow_html=True)

        if scan_btn:
            sm = {"📈 BUY": "BUY", "📉 SELL": "SELL", "🔍 ALL": "ALL"}[s_sig]
            _st8 = {}
            res = run_scan(scan_stocks, s_iv, "1mo", s_mc, sm, workers=18, cap_n=cap_n, stats_out=_st8)
            ss.scan_results = res
            _ok, _tot = _st8.get("ok", 0), _st8.get("total", 0)
            _fail = _st8.get("fails", 0)
            if res:
                st.success(f"✅ Found {len(res)} stocks! · data OK for {_ok}/{_tot} scanned · "
                           f"uptrend + strongest setups on top.")
            elif _fail > _ok:
                st.warning(f"⚠️ Yahoo didn't answer for {_fail} of {_tot} stocks (server throttling — "
                           f"NOT a scan problem). Press 🚀 START SCAN once more, or use the 🔴 Live Dashboard / "
                           f"🎯 Combo tab (batched downloads — they throttle far less).")
            else:
                st.warning(f"No stocks matched your filters (scanned {_ok} OK). Lower Min Conf % "
                           f"(try 55–60), switch Signal to 🔍 ALL, or pick a bigger category.")

        if ss.scan_results:
            total_found = len(ss.scan_results)
            rc1, rc2 = st.columns(2)
            with rc1:
                _sho = st.selectbox("📋 Results to show", ["200", "100", "50", "300", "500", "ALL"],
                                    index=0, key="sc_shown")
            show_n = total_found if _sho == "ALL" else min(int(_sho), total_found)
            with rc2:
                detail_n = st.slider("🖼️ Detail cards (top N)", 10, 200, 25, 5, key="sc_detail",
                                     help="Expandable detail cards for the top N. The FULL table + CSV "
                                          "download below always covers everything you selected.")
            results = ss.scan_results[:show_n]
            st.markdown(f"<div style='background:#f0fdf4;border:1px solid #bbf7d0;border-radius:10px;"
                        f"padding:8px 14px;color:#166534;font-size:12.5px;'>📋 Showing <b>{show_n}</b> of "
                        f"<b>{total_found}</b> matched stocks · full sortable table + ⬇️ CSV download below</div>",
                        unsafe_allow_html=True)
            st.markdown("### 🏆 Top 3")
            top3 = st.columns(min(3, len(results))); medals = ["🥇", "🥈", "🥉"]
            for i, (col, r) in enumerate(zip(top3, results[:3])):
                with col:
                    up = r['tr'] == 'UPTREND'
                    st.markdown(f"<div style='background:white;border:2px solid {r['sc']};border-radius:18px;padding:20px;text-align:center;box-shadow:0 6px 24px rgba(0,0,0,0.08);'><div style='font-size:28px;'>{medals[i]}</div><div style='font-size:15px;font-weight:800;color:#1a1f36;margin:6px 0;'>{r['name'][:22]}</div><div style='font-size:24px;font-weight:900;color:#1d4ed8;'>₹{r['price']:.2f}</div><div style='background:{r['sc']};color:white;border-radius:20px;padding:5px 16px;font-size:12px;font-weight:800;margin:10px auto;display:inline-block;'>{r['sig']}</div><div style='font-size:15px;font-weight:800;color:{'#16a34a' if up else '#6b7280'};margin-top:6px;'>{'📈 UPTREND' if up else r['tr'].title()}</div><div style='color:#6b7280;font-size:12px;margin-top:4px;'>Conf {r['conf']:.0f}% · Buy ₹{r['buy_at']:.2f}</div></div>", unsafe_allow_html=True)
            st.markdown("---"); st.markdown(f"### 📋 Top {min(detail_n, len(results))} — detail cards")
            for i, r in enumerate(results[:detail_n]):
                up = r['tr'] == 'UPTREND'
                with st.expander(f"{'📈' if up else '•'} #{i+1} · {r['name'][:26]} · ₹{r['price']:.2f} · {r['sig']} · {r['conf']:.0f}% · {r['tr'].title()} · Gap {r['gap']:+.1f}%", expanded=i < 3):
                    st.markdown(f"<div class='sc-r'><div style='font-size:15px;font-weight:800;color:{r['ac']};margin-bottom:12px;'>{r['act']}</div><div style='display:grid;grid-template-columns:repeat(5,1fr);gap:10px;text-align:center;'><div style='background:#eff6ff;border-radius:10px;padding:10px;'><div style='color:#6b7280;font-size:10px;'>TREND</div><div style='color:{'#16a34a' if up else '#6b7280'};font-weight:900;font-size:14px;'>{r['tr'].title()}</div></div><div style='background:white;border:1px solid #e0e7ff;border-radius:10px;padding:10px;'><div style='color:#6b7280;font-size:10px;'>BUY AT</div><div style='color:{r['ac']};font-weight:800;font-size:15px;'>₹{r['buy_at']:.2f}</div></div><div style='background:#fff1f2;border-radius:10px;padding:10px;'><div style='color:#6b7280;font-size:10px;'>STOP</div><div style='color:#dc2626;font-weight:800;font-size:15px;'>₹{r['sl']:.2f}</div></div><div style='background:#f0fdf4;border-radius:10px;padding:10px;'><div style='color:#6b7280;font-size:10px;'>TARGET 2</div><div style='color:#16a34a;font-weight:800;font-size:15px;'>₹{r['t2']:.2f}</div><div style='color:#6b7280;font-size:9px;'>+{r['t2p']:.1f}%</div></div><div style='background:#eff6ff;border-radius:10px;padding:10px;'><div style='color:#6b7280;font-size:10px;'>RISK:REWARD</div><div style='color:#1d4ed8;font-weight:900;font-size:16px;'>1:{r['rr']}</div></div></div><div style='color:#6b7280;font-size:11px;margin-top:10px;'>Conf {r['conf']:.0f}% · Vol {r['vr']:.1f}x · Buy strength {r['bp']:.0f}%</div></div>", unsafe_allow_html=True)
                    if st.button(f"📊 Full Analysis → {r['name'][:20]}", key=f"scf_{i}_{r['sym']}", **STRETCH, type="primary"):
                        ss.sym = r['sym']; ss.stock_name = r['name']; ss.analyzed = True
                        st.info("✅ Open the '📊 Analyze Stock' tab — it's loaded with this stock.")

            # ── 📋 FULL results table (all shown rows, sortable) + CSV download ──
            st.markdown("---")
            st.markdown(f"### 📋 Full Results Table ({len(results)} stocks)")
            disp = pd.DataFrame([{
                "#": i + 1, "Stock": r["name"], "Symbol": r["sym"], "Price": r["price"],
                "Signal": r["sig"], "Conf%": round(r["conf"], 1), "Trend": str(r["tr"]).title(),
                "Buy@": r["buy_at"], "SL": r["sl"], "T2": r["t2"], "T2%": r["t2p"],
                "R:R": f"1:{r['rr']}", "Gap%": r["gap"], "Vol×": round(r["vr"], 2),
                "BuyStr%": round(r["bp"], 1),
            } for i, r in enumerate(results)])
            try:
                st.dataframe(disp, **STRETCH, height=430, hide_index=True)
            except Exception:
                st.dataframe(disp, **STRETCH)
            try:
                st.download_button("⬇️ Download ALL results (CSV)", data=disp.to_csv(index=False).encode(),
                                   file_name=f"scan_results_{now_ist().strftime('%Y%m%d_%H%M')}.csv",
                                   mime="text/csv", **STRETCH, key="sc_csv")
            except Exception:
                pass

    # ── TAB 3: SEARCH ──
    with tab_search:
        st.markdown('<div class="input-row">', unsafe_allow_html=True)
        st.markdown("<div style='font-size:15px;font-weight:800;color:#1d4ed8;margin-bottom:12px;'>🔎 SEARCH ANY NSE/BSE STOCK (full live universe)</div>", unsafe_allow_html=True)
        q1, q2 = st.columns([4, 1])
        with q1:
            sq = st.text_input("Search", placeholder="Zensar Technologies, Oil India, Vedanta, Suzlon, SBIN...", key="sq_in", label_visibility="collapsed")
        with q2:
            sb = st.button("🔍 SEARCH", **STRETCH, key="sq_btn")
        st.markdown('</div>', unsafe_allow_html=True)
        if sb and sq:
            with st.spinner(f"Searching '{sq}'..."):
                ss.search_results = do_search(sq)
        if ss.search_results is not None:
            if ss.search_results:
                st.success(f"✅ Found {len(ss.search_results)} result(s)")
                cols = st.columns(min(3, len(ss.search_results)))
                for i, (label, data) in enumerate(ss.search_results.items()):
                    with cols[i % 3]:
                        p = data['price']
                        st.markdown(f"<div style='background:white;border:1px solid #e0e7ff;border-radius:16px;padding:20px;text-align:center;box-shadow:0 4px 16px rgba(0,0,0,0.06);margin:4px 0;'><div style='font-weight:800;color:#1a1f36;font-size:14px;'>{data['name']}</div><div style='font-size:30px;font-weight:900;color:#1d4ed8;margin:10px 0;'>₹{p:,.2f}</div><div style='color:#9ca3af;font-size:11px;'>{data['sym']}</div></div>", unsafe_allow_html=True)
                        if st.button("📊 Analyze", key=f"sr_{i}", **STRETCH):
                            ss.sym = data['sym']; ss.stock_name = data['name']; ss.analyzed = True
                            st.success("✅ Open the '📊 Analyze Stock' tab.")
            elif sq:
                st.error(f"❌ '{sq}' not found. Try exact symbol like {sq}.NS")
        st.markdown("### 🔥 Popular")
        popular = {"RELIANCE": "RELIANCE.NS", "TCS": "TCS.NS", "SBIN": "SBIN.NS", "INFY": "INFY.NS",
                   "HDFCBANK": "HDFCBANK.NS", "ITC": "ITC.NS", "SUZLON": "SUZLON.NS", "YESBANK": "YESBANK.NS",
                   "HAL": "HAL.NS", "RVNL": "RVNL.NS", "NTPC": "NTPC.NS", "ETERNAL": "ETERNAL.NS",
                   "TMPV": "TMPV.NS", "PNB": "PNB.NS", "VEDL": "VEDL.NS", "ZENSARTECH": "ZENSARTECH.NS"}
        pc = st.columns(4)
        for i, (nm, sym) in enumerate(popular.items()):
            with pc[i % 4]:
                if st.button(nm, key=f"pop_{nm}", **STRETCH):
                    ss.sym = sym; ss.stock_name = nm; ss.analyzed = True
                    st.success(f"✅ Open '📊 Analyze Stock' tab for {nm}")

    # ── TAB 5: JOURNAL ──
    with tab_journal:
        st.markdown("<div style='background:linear-gradient(135deg,#eff6ff,#dbeafe);border:2px solid #3b82f6;"
                    "border-radius:16px;padding:18px;margin-bottom:16px;'><div style='font-size:18px;font-weight:800;"
                    "color:#1d4ed8;'>📓 Trade Journal — did the plan work?</div><div style='color:#374151;"
                    "font-size:13px;margin-top:8px;'>Save any analysis (button under 'Save this analysis'), then come "
                    "back the next day. It checks what price actually did and marks each call a WIN, LOSS, or NO-FILL — "
                    "so you build a real accuracy record and see what to improve.</div></div>", unsafe_allow_html=True)
        entries = journal_load()
        if not entries:
            st.info("No saved analyses yet. Analyze a stock, then press '📓 Save to Journal'.")
        else:
            wins = losses = nofill = openp = 0
            rows = []
            for e in sorted(entries, key=lambda x: x.get("saved", ""), reverse=True):
                v = journal_verify(e)
                s = v["status"]
                if "WIN" in s: wins += 1
                elif "LOSS" in s: losses += 1
                elif "NO-FILL" in s: nofill += 1
                elif "OPEN" in s: openp += 1
                rows.append((e, v))
            resolved = wins + losses
            acc = round(wins / resolved * 100, 1) if resolved else 0
            k1, k2, k3, k4, k5 = st.columns(5)
            with k1: st.metric("Saved", len(entries))
            with k2: st.metric("Wins", wins)
            with k3: st.metric("Losses", losses)
            with k4: st.metric("No-fill (waited)", nofill)
            with k5: st.metric("Hit-rate", f"{acc}%", "of resolved")
            st.markdown("<div style='color:#6b7280;font-size:12px;margin:6px 0 14px;'>Hit-rate counts only "
                        "resolved trades (win vs loss). 'No-fill' means price never reached your buy level — "
                        "that's the plan correctly keeping you out.</div>", unsafe_allow_html=True)
            for e, v in rows:
                if not all(k in e for k in ("saved", "name", "buy_at", "sl", "t1", "t2", "price")):
                    continue   # skip damaged/incomplete entries instead of crashing
                with st.expander(f"{v['status']} · {e['name'][:26]} · saved {e['saved']} · "
                                 f"{e.get('trend','')} {e.get('signal','')}", expanded=False):
                    st.markdown(
                        f"<div style='background:white;border:1px solid #e0e7ff;border-left:5px solid {v['color']};"
                        f"border-radius:0 12px 12px 0;padding:14px 18px;'>"
                        f"<div style='font-size:15px;font-weight:800;color:{v['color']};'>{v['status']}</div>"
                        f"<div style='color:#374151;font-size:13px;margin-top:4px;'>{v['detail']}</div>"
                        f"<div style='color:#6b7280;font-size:12px;margin-top:8px;'>Planned: buy ₹{e['buy_at']:.2f} · "
                        f"stop ₹{e['sl']:.2f} · T1 ₹{e['t1']:.2f} · T2 ₹{e['t2']:.2f} · saved at ₹{e['price']:.2f} · "
                        f"stage {e.get('stage','')} · news: {e.get('news','—')}</div></div>", unsafe_allow_html=True)
                    if st.button("🗑️ Delete this entry", key=f"del_{e.get('id','')}"):
                        journal_delete(e.get("id", ""))
                        st.rerun()

    # ── TAB: EOD REVIEW (before vs after) ──
    with tab_eod:
        try:
            eod_review_tab(ss)
        except Exception as e:
            st.warning(f"⚠️ EOD review problem: {type(e).__name__}: {e}")

    # ── TAB 6: GUIDE ──
    with tab_guide:
        st.markdown("## 📚 Quick Trading Guide")
        st.markdown("""
        <div class='tr-g'><b style='color:#16a34a;'>Rule 1 — Target price hit = book &amp; stop.</b> Price reached your target? Take profit, don't get greedy.</div>
        <div class='tr-b'><b style='color:#dc2626;'>Rule 2 — Always use the stop loss.</b> No SL = no trade.</div>
        <div class='tr-b'><b style='color:#dc2626;'>Rule 3 — Never average down.</b> SL hit → exit, don't buy more.</div>
        <div class='tr-w'><b style='color:#b45309;'>Rule 4 — 9:15 AM is not for buying.</b> Wait for the 9:30 first candle.</div>
        <div class='tr-w'><b style='color:#b45309;'>Rule 5 — Gap up > 2%? Don't chase.</b> Wait for a pullback or skip.</div>
        <div class='tr-i'><b style='color:#2563eb;'>Rule 6 — Trend is your friend.</b> Uptrend → buy dips. Downtrend → avoid longs.</div>
        <div class='tr-i'><b style='color:#2563eb;'>🔴 Live Dashboard workflow:</b> START the board → watch the 🚀 UPTREND panel → open a stock's tab or click Analyze on a top card → confirm with pivots/news/ML in the full analysis → trade only with a stop. The board finds candidates; the deep analysis confirms.</div>
        <div class='tr-i'><b style='color:#2563eb;'>Pivots — which to use:</b> Standard = general S/R · Camarilla = tight intraday reversals (R3/S3) & breakouts (R4/S4) · Woodie = faster, momentum-weighted · Fibonacci = 38.2/50/61.8% pullback zones.</div>
        <div class='tr-i'><b style='color:#2563eb;'>Data source (current + upgrade path):</b> this app uses Yahoo Finance — free, no key, ~15-min delay possible and occasional throttling on huge boards. When you want true broker-grade real-time for NSE, the good options are: <b>Upstox API v2</b> (free with an account), <b>Angel One SmartAPI</b> (free), <b>Fyers API</b> (free), or <b>Zerodha Kite Connect</b> (paid, most popular). The app's logic stays the same — only the data fetch layer would change.</div>
        <div class='tr-i'><b style='color:#2563eb;'>🌙 EOD Review workflow:</b> run the 🔴 board during market hours (it saves ONE memory snapshot per day — that day's calculation) → end of day / next day open 🌙 EOD Review → pick the day → VERIFY → see exactly how many stocks moved as per the calculation, how many hit T1/T2 vs the stop, and your accuracy scorecard over time. On Streamlit Cloud, memory is wiped on redeploy — use 💾 Backup / ⬆️ Restore in the EOD tab around app updates.</div>
        <div class='tr-i'><b style='color:#2563eb;'>🧑‍🤝‍🧑 Sharing this app with a friend?</b> No clashing: in the 🌙 EOD tab,
        each of you types your <b>own memory name</b> once (e.g. ravi, arjun) and bookmarks the personal link
        (it carries <code>?u=ravi</code>). Snapshots, EOD scorecard and 📓 Journal are then kept completely separate —
        your friend can never see or overwrite your data, and live screens are always separate anyway.</div>
        <div class='tr-g'><b style='color:#16a34a;'>Honest truth:</b> no tool predicts price. These stack the odds and define your risk — they don't remove it. The stop loss is what actually protects your capital.</div>
        """, unsafe_allow_html=True)

    st.markdown("<div style='text-align:center;color:#9ca3af;font-size:10px;padding:16px;border-top:2px solid #e0e7ff;margin-top:20px;'>⚠️ EDUCATIONAL PURPOSE ONLY · NOT FINANCIAL ADVICE · ALWAYS USE STOP LOSS · TRADE AT YOUR OWN RISK · Past performance ≠ future results</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
